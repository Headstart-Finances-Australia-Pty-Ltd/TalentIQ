"""
TalentIQ - CandidateLens Router
Mirrors the original JobLens scoring logic:
  1. LLM extracts skills from JD (Groq instead of Ollama)
  2. Keyword match CV text against extracted skills
  3. Bonus for degree/experience mentions
  4. Generate interview questions via LLM
  5. Generate a 10-statement resume summary via LLM
  6. Send video-interview invite emails with a candidate-facing link
All persisted to PostgreSQL (tiq_joblens_* tables).
"""
import io
import re
import os
import json
import logging
import asyncio
import secrets
import smtplib
import requests
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from pathlib import Path
from typing import List, Optional
from datetime import datetime, timedelta

from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Form, BackgroundTasks, Request
from fastapi.responses import Response
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select, or_, func
from pydantic import BaseModel

from db.database import get_db, AsyncSessionLocal
from models.models import User, UserAPIKey, JobLensSession, JobLensCandidate
from utils.auth_utils import get_current_user
from utils.credentials import get_credential, get_all_credentials, get_groq_model, ollama_enabled, DEFAULT_GROQ_MODEL
from utils.sequencing import next_sequence_number
from utils.storage import (
    upload_video_and_get_key, get_video_bytes, get_video_presigned_url, delete_video,
    upload_file, get_file_bytes, get_presigned_url, delete_file,
)

router = APIRouter()
logger = logging.getLogger(__name__)


@router.get("/jd-options")
async def list_jd_options(
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """JD Management records for the 'New Analysis' JD-selection dropdown,
    including Client Name so it can be shown alongside the title."""
    from models.models import JDRecord, Client as ClientModel
    r = await db.execute(
        select(JDRecord).where(JDRecord.user_id == current_user.id).order_by(JDRecord.created_at.desc())
    )
    jds = r.scalars().all()
    out = []
    for jd in jds:
        client_name = ""
        if jd.client_id:
            cr = await db.execute(select(ClientModel).where(ClientModel.id == jd.client_id))
            client = cr.scalar_one_or_none()
            if client:
                client_name = client.name
        out.append({"id": jd.id, "jd_title": jd.title, "client_name": client_name, "status": jd.status})
    return out


@router.get("/requisition-options")
async def list_requisition_options(
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Open Requisitions for the 'New Analysis' Requisition-selection
    dropdown. Replaces the old JD Management-only source: once a
    requisition is picked, its JD (own attached file first, else the
    linked JD Management record) and its submitted candidates are pulled
    from the database automatically — see /run and /requisition-candidates."""
    from capabilities.requisition.models import Requisition
    from capabilities.acquisition import service as acquisition_service
    from models.models import Client as ClientModel
    org = await acquisition_service.get_or_create_default_organisation(db, current_user)
    r = await db.execute(
        select(Requisition).where(Requisition.organisation_id == org.id, Requisition.status == "Open")
        .order_by(Requisition.created_at.desc())
    )
    reqs = r.scalars().all()
    out = []
    for req in reqs:
        client_name = ""
        if req.client_id:
            cr = await db.execute(select(ClientModel).where(ClientModel.id == req.client_id))
            client = cr.scalar_one_or_none()
            if client:
                client_name = client.name
        out.append({
            "id": req.id,
            "title": req.title,
            "client_name": client_name,
            "has_jd_file": bool(req.jd_file_blob or req.jd_file_key),
            "has_jd_record": bool(req.jd_record_id),
            "has_jd": bool(req.jd_file_blob or req.jd_file_key or req.jd_record_id),
        })
    return out


@router.get("/requisition-candidates")
async def list_requisition_candidates(
    requisition_id: int,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Candidates actually SUBMITTED for this requisition — i.e. Application
    rows (Candidate <-> Requisition), not the older Vendor Management /
    JD Management TrackedCandidate flow. Resumes come from the Candidate
    Master's own resume_blob (the same file used everywhere else that
    candidate is referenced), so 'has_resume' below reflects what /run
    will actually be able to score."""
    from capabilities.requisition.models import Requisition
    from capabilities.acquisition.models import Application, Candidate
    from capabilities.acquisition import service as acquisition_service
    org = await acquisition_service.get_or_create_default_organisation(db, current_user)
    req = (await db.execute(
        select(Requisition).where(Requisition.id == requisition_id, Requisition.organisation_id == org.id)
    )).scalar_one_or_none()
    if not req:
        raise HTTPException(404, "Requisition not found")

    r = await db.execute(
        select(Application, Candidate)
        .join(Candidate, Candidate.id == Application.candidate_id)
        .where(Application.requisition_id == requisition_id, Application.organisation_id == org.id)
        .order_by(Application.applied_at.desc())
    )
    out = []
    for application, candidate in r.all():
        out.append({
            "id": application.id,
            "name": candidate.full_name,
            "email": candidate.email or "",
            "status": application.stage,
            "has_resume": bool(candidate.resume_blob),
        })
    return out


@router.get("/vendor-candidates")
async def list_vendor_candidates_for_jd(
    jd_id: int,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """TrackedCandidates (Vendor Management / Profile Management submissions)
    for a specific JD — dynamically filtered so 'New Analysis' only offers
    profiles actually relevant to a fresh analysis run: candidates already
    Shortlisted, Selected, or Offered are excluded, since re-running ATS
    scoring on a candidate who has already progressed past that stage isn't
    useful here."""
    from models.models import TrackedCandidate, Vendor as VendorModel
    EXCLUDED_STATUSES = {"Shortlisted", "Selected", "Offered"}
    r = await db.execute(
        select(TrackedCandidate).where(TrackedCandidate.jd_id == jd_id, TrackedCandidate.user_id == current_user.id)
        .order_by(TrackedCandidate.created_at.desc())
    )
    candidates = [c for c in r.scalars().all() if c.status not in EXCLUDED_STATUSES]
    out = []
    for c in candidates:
        vr = await db.execute(select(VendorModel).where(VendorModel.id == c.vendor_id))
        vendor = vr.scalar_one_or_none()
        out.append({
            "id": c.id,
            "name": c.name,
            "vendor_id": c.vendor_id,
            "vendor_name": vendor.name if vendor else "",
            "has_resume": bool(c.resume_blob),
            "status": c.status,
        })
    return out


# ── TEXT EXTRACTION ─────────────────────────────────────────────────────────

def extract_text(content: bytes, filename: str) -> str:
    fname = (filename or "").lower()
    if fname.endswith(".txt"):
        for enc in ("utf-8", "latin-1", "cp1252"):
            try: return content.decode(enc)
            except: continue
        return ""
    if fname.endswith(".pdf"):
        # Layout-aware pass first (multi-column + table detection) — see
        # utils/layout_parse.py. Shared with CVIntel so both modules
        # benefit from the same fix for scrambled multi-column resumes.
        try:
            from utils.layout_parse import extract_pdf_layout_aware
            layout_text = extract_pdf_layout_aware(content)
            if layout_text.strip():
                return layout_text
        except Exception:
            pass
        try:
            import pdfplumber
            with pdfplumber.open(io.BytesIO(content)) as pdf:
                t = "\n".join(p.extract_text() or "" for p in pdf.pages).strip()
                if t: return t
        except: pass
        try:
            import pypdf
            r = pypdf.PdfReader(io.BytesIO(content))
            t = "\n".join(p.extract_text() or "" for p in r.pages).strip()
            if t: return t
        except: pass
        return ""
    if fname.endswith((".docx", ".doc")):
        # Try python-docx first (works on .docx)
        try:
            import docx
            doc = docx.Document(io.BytesIO(content))

            # Headers often contain name/email/phone (especially in resumes
            # with a banner/letterhead layout). python-docx's section.header
            # only exposes ONE header per section by default, but a docx can
            # have up to 3 per section (default/first-page/even-page) stored
            # as header1.xml/header2.xml/header3.xml — read them all via the
            # raw XML so we never miss contact details placed there.
            header_footer_parts = []
            try:
                import re as _re2
                import zipfile as _zipfile
                with _zipfile.ZipFile(io.BytesIO(content)) as z:
                    for name in z.namelist():
                        if _re2.match(r"word/(header|footer)\d*\.xml$", name):
                            xml = z.read(name).decode("utf-8", errors="ignore")
                            texts = _re2.findall(r"<w:t[^>]*>([^<]*)</w:t>", xml)
                            joined = "".join(texts).strip()
                            if joined:
                                header_footer_parts.append(joined)
            except Exception:
                pass

            parts = list(header_footer_parts)  # put header/footer text first
            parts += [p.text for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip(): parts.append(cell.text.strip())
            t = "\n".join(parts).strip()
            if t: return t
        except: pass
        # Try docx2txt
        try:
            import docx2txt
            t = docx2txt.process(io.BytesIO(content))
            if t and t.strip(): return t.strip()
        except: pass
        # Fallback for old binary .doc (OLE2 format): extract ASCII text stream
        if fname.endswith(".doc"):
            try:
                import re as _re
                raw = content.decode("latin-1", errors="ignore")
                chunks = _re.findall(r"[\x20-\x7e\r\n\t]{3,}", raw)
                text = "\n".join(c.strip() for c in chunks if c.strip())
                # Remove common .doc binary artifacts
                text = _re.sub(r"bjbj[a-zA-Z0-9]+", "", text)
                text = _re.sub(r"WW8Num\w+", "", text)
                text = _re.sub(r'HYPERLINK\s+"[^"]+"', "", text)
                text = _re.sub(r"\\r", "\n", text)
                text = _re.sub(r"\s{4,}", "\n", text)
                text = _re.sub(r"\n{3,}", "\n\n", text).strip()
                if len(text) > 100:
                    return text
            except: pass
        return ""
    for enc in ("utf-8", "latin-1"):
        try: return content.decode(enc).strip()
        except: continue
    return ""


# ── CANDIDATE INFO EXTRACTION ────────────────────────────────────────────────

def extract_candidate_info(text: str, filename: str) -> dict:
    """Extract name, email, phone — mirrors original extractCandidateDetails."""
    # Word documents with letterhead-style headers often have run-together
    # text with no whitespace between fields (Word renders separate <w:t>
    # runs with different formatting as visually adjacent but textually
    # concatenated), e.g. "NSW 2145resume2@gmail.comLinkedIn:". Insert a
    # space at lowercase→uppercase and digit→letter transitions so regexes
    # below can find clean boundaries — this never touches genuine words.
    norm_text = re.sub(r'([a-z])([A-Z])', r'\1 \2', text)
    norm_text = re.sub(r'(\d)([a-zA-Z])', r'\1 \2', norm_text)

    # PDFs go the OTHER way too: a contact-icon glyph (the envelope/phone
    # icon many resume templates place right before the email/phone)
    # frequently gets extracted as stray whitespace, splitting
    # "name@gmail.com" into "name @gmail.com" or "name@ gmail.com" — a
    # space the email regex below won't tolerate right around the "@",
    # so the whole email silently failed to match even though it's right
    # there in the text. Collapse whitespace immediately touching an "@"
    # back together before matching, but only when it's actually part of
    # an email-shaped token (word chars on one side, a domain+TLD on the
    # other) — this never touches a genuine "meet @ 3pm"-style sentence.
    norm_text = re.sub(r'([\w.+-])\s+@\s*([\w.-]+\.[a-zA-Z]{2,6})', r'\1@\2', norm_text)
    norm_text = re.sub(r'([\w.+-])\s*@\s+([\w.-]+\.[a-zA-Z]{2,6})', r'\1@\2', norm_text)

    # Email
    email_m = re.search(r"[a-zA-Z][\w.+-]*@[\w.-]+\.[a-zA-Z]{2,6}", norm_text)
    email = email_m.group() if email_m else ""

    # Phone — original regex: (+?\d{1,4}[\s-]?\(?\d{1,4}\)?[\s-]?\d{3,4}[\s-]?\d{3,4})
    phone_m = re.search(r"(\+?\d{1,4}[\s\-]?\(?\d{1,4}\)?[\s\-]?\d{3,4}[\s\-]?\d{3,4})", text)
    phone = phone_m.group().replace(" ", "").strip() if phone_m else ""
    # Reject cert/reference numbers that happen to be digit runs with no
    # phone-style separator — UNLESS the run is exactly a plausible phone
    # length (10-13 digits, covering a bare local number up through a
    # full +countrycode number). Plenty of real resumes — this one
    # included — write a mobile number as one unbroken 10-digit run with
    # no space/dash/parens at all (very common for Indian mobile numbers
    # specifically), and the old version rejected every one of those as
    # if it were a certificate ID, discarding a perfectly valid phone
    # number. Genuine cert/reference numbers are the ones that fall
    # OUTSIDE this length band (shorter, or noticeably longer), so this
    # keeps the original guard for those while no longer punishing a
    # correctly-formatted-but-unspaced phone number.
    if phone and not re.search(r"[\s\-\+\(\)]", phone_m.group() if phone_m else ""):
        digits_only = re.sub(r"\D", "", phone)
        if not (10 <= len(digits_only) <= 13):
            phone = ""

    # Name — mirrors original: first non-email, non-4digit, <=5 word line
    name = ""
    NAME_SKIP_PATTERNS = (
        r"^page\s+\d+\s*(of\s+\d+)?$",
        r"^\d+\s*$",
        r"^confidential$",
        r"^curriculum vitae$",
        r"^resume$",
    )
    lines = [l.strip() for l in text.split("\n") if l.strip()]
    for line in lines[:25]:
        if "@" in line:
            continue
        if re.search(r"\d{4,}", line):
            continue
        if any(re.match(p, line, re.IGNORECASE) for p in NAME_SKIP_PATTERNS):
            continue
        if len(line.split()) <= 5 and len(line) > 2:
            candidate = re.sub(r"[^a-zA-Z\s\-\.]", "", line).strip()
            if candidate and len(candidate) > 2:
                name = candidate
                break

    if not name:
        name = Path(filename).stem.replace("_", " ").replace("-", " ").title()

    # ── Experience years — look for explicit "X years" mentions ───────
    exp_years = ""
    exp_matches = re.findall(
        r"(\d{1,2})\+?\s*(?:years?|yrs?)\s*(?:of\s+)?(?:experience|exp\b)",
        text, re.IGNORECASE,
    )
    if exp_matches:
        # Take the highest mentioned figure (usually the headline summary)
        exp_years = f"{max(int(y) for y in exp_matches)}+ years"
    else:
        # Fallback: count distinct 4-digit years mentioned in work history
        # to estimate a rough career span (e.g. 2009 ... 2024)
        years_found = sorted(set(int(y) for y in re.findall(r"\b(19[7-9]\d|20[0-2]\d)\b", text)))
        if len(years_found) >= 2:
            span = years_found[-1] - years_found[0]
            if 0 < span <= 45:
                exp_years = f"~{span} years"

    # ── Summary — first substantial paragraph (career objective / profile) ──
    summary = ""
    paragraphs = [p.strip() for p in re.split(r"\n\s*\n", text) if p.strip()]
    SUMMARY_SKIP = {"resume", "curriculum vitae", "cv", "page", "contact", "references"}
    for p in paragraphs[:15]:
        clean_p = re.sub(r"\s+", " ", p).strip()
        low = clean_p.lower()
        if len(clean_p) < 80 or len(clean_p) > 700:
            continue
        if any(s in low[:30] for s in SUMMARY_SKIP):
            continue
        if "@" in clean_p or re.search(r"^\s*[\u2022\-\*]", p):
            continue
        # Looks like prose (has multiple sentences / reasonable word count)
        if len(clean_p.split()) >= 12:
            summary = clean_p[:400]
            break

    return {
        "name": name, "email": email, "phone": phone,
        "experience_years": exp_years, "summary": summary,
    }


# ── SKILL EXTRACTION FROM JD (LLM) ──────────────────────────────────────────

_PLACEHOLDER_VALUES = {
    "nil", "n/a", "na", "none", "-", "--", "tbd", "tba", "blank", "n.a.",
    "not specified", "not applicable", "unknown", "null",
}


def _clean_extracted_field(value: Optional[str]) -> str:
    """Drops obviously-blank template placeholders (a JD table cell that
    literally says "Nil" or "N/A" is not a real role/location/company —
    treating it as one was the source of the "ROLE: Nil" bug)."""
    if not value:
        return ""
    v = value.strip()
    if not v or v.lower() in _PLACEHOLDER_VALUES or len(v) > 120:
        return ""
    return v


async def extract_jd_details(jd_text: str, groq_key: str, groq_model: str = DEFAULT_GROQ_MODEL) -> dict:
    """Extracts role title, location, company, and skills (categorized into
    Essential / Good to Have / Optional) from the JD in a single LLM call —
    used for both the CandidateLens "Job Description Summary" panel and the
    skill-matching logic below."""
    try:
        from langchain_groq import ChatGroq
        from langchain.schema import HumanMessage
        from utils.llm_extraction import _truncate_for_llm

        llm = ChatGroq(api_key=groq_key, model=groq_model, temperature=0.1, max_tokens=4000, reasoning_format="hidden", reasoning_effort="low", max_retries=0)
        prompt = f"""You are a recruitment AI assistant. Read the job description below and
extract these fields precisely. If a field genuinely isn't stated anywhere
in the text, return an empty string for it — do NOT guess, and do NOT
return placeholder text like "Nil", "N/A", or "TBD" as if it were a real
value.

Also categorize every required/desired skill or requirement into exactly
one of three tiers, based on how the JD phrases it:
- "essential": stated as required/must-have/mandatory
- "good_to_have": stated as preferred/desirable/advantageous but not mandatory
- "optional": mentioned only in passing, or a minor/bonus item

Job Description:
\"\"\"{_truncate_for_llm(jd_text, "JD text")}\"\"\"

Return ONLY valid JSON in this format:
{{
  "role": "<job title, or empty string if not stated>",
  "location": "<work location/city, or empty string if not stated>",
  "company": "<hiring company name, or empty string if not stated>",
  "essential": ["skill1", "skill2"],
  "good_to_have": ["skill3", "skill4"],
  "optional": ["skill5"]
}}"""

        from utils.llm_extraction import _parse_json_response
        resp = llm.invoke([HumanMessage(content=prompt)])
        data = _parse_json_response(resp.content)
        if data is None:
            raise ValueError(f"LLM returned unparseable/empty response (length {len(resp.content)})")
        essential = [s.lower().strip() for s in data.get("essential", []) if s]
        good_to_have = [s.lower().strip() for s in data.get("good_to_have", []) if s]
        optional = [s.lower().strip() for s in data.get("optional", []) if s]
        return {
            "role": _clean_extracted_field(data.get("role")),
            "location": _clean_extracted_field(data.get("location")),
            "company": _clean_extracted_field(data.get("company")),
            "essential": essential,
            "good_to_have": good_to_have,
            "optional": optional,
            "skills": list(dict.fromkeys(essential + good_to_have + optional)),  # flat list — existing scoring logic
        }
    except Exception as e:
        print(f"  WARNING: extract_jd_details LLM call failed, falling back to keyword heuristic — {type(e).__name__}: {str(e)[:300]}")
        return _heuristic_jd_details(jd_text)


async def extract_skills_from_jd(jd_text: str, groq_key: str) -> list:
    """Kept for any other call sites that only want the skill list."""
    details = await extract_jd_details(jd_text, groq_key)
    return details["skills"]


def _keyword_extract_jd(jd_text: str) -> list:
    """Fallback keyword extraction — returns only real skills, not JD prose."""
    DOMAIN_SKILLS = [
        "python","javascript","typescript","react","node","sql","postgresql","mongodb",
        "aws","azure","gcp","docker","kubernetes","git","agile","rest","api","graphql",
        "machine learning","ai","artificial intelligence","data science","excel","power bi","tableau","salesforce",
        "figma","django","flask","java","c#","c++","go","spark","kafka","dbt","snowflake",
        "databricks","redshift","bigquery","data architecture","data governance","etl",
        "data mesh","data fabric","data vault","dimensional modelling","dimensional modeling",
        "enterprise data warehouse","edw","lakehouse","master data management","mdm",
        "data quality","data catalog","collibra","alation","informatica","talend",
        "teradata","hive","hbase","adls","synapse","azure data factory","azure synapse",
        "event-driven architecture","real-time data","streaming","enterprise architecture",
        "solution architecture","zachman","basel","basel iii","banking","bfsi","insurance",
        "lending","regulatory compliance","risk management","governance framework",
        "xero","myob","quickbooks","sap","oracle","dynamics","netsuite",
        "cpa","ca","acca","cma","mba","cfa","phd","accounting","tax","audit","payroll",
        "financial reporting","budgeting","forecasting","reconciliation","ifrs","gaap",
        "leadership","communication","problem solving","scrum","project management",
        "togaf","pmp","csm","hadoop","datastage","tibco","react native",
        "devops","ci/cd","terraform","ansible","linux","bash","swift","kotlin",
        "microservices","restful","graphql","redis","elasticsearch","rabbitmq",
    ]
    jd_lower = jd_text.lower()
    found = [s for s in DOMAIN_SKILLS if s in jd_lower]
    return list(dict.fromkeys(found))[:30]


def _heuristic_jd_details(jd_text: str) -> dict:
    """Non-LLM fallback for role/location/company — same placeholder
    filtering as the LLM path, so a JD with a literal 'Position: Nil'
    template field doesn't get treated as a real job title."""
    role_m = re.search(r"(?:job\s*title|role|position\s*title)\s*[:\-]\s*(.+)", jd_text, re.IGNORECASE)
    loc_m = re.search(r"(?:location|based\s*in|located\s*in)\s*[:\-]\s*(.+)", jd_text, re.IGNORECASE)
    comp_m = re.search(r"(?:company|organisation|employer)\s*[:\-]\s*(.+)", jd_text, re.IGNORECASE)
    skills = _keyword_extract_jd(jd_text)
    return {
        "role": _clean_extracted_field(role_m.group(1).split("\n")[0] if role_m else None),
        "location": _clean_extracted_field(loc_m.group(1).split("\n")[0] if loc_m else None),
        "company": _clean_extracted_field(comp_m.group(1).split("\n")[0] if comp_m else None),
        "essential": skills,
        "good_to_have": [],
        "optional": [],
        "skills": skills,
    }


# ── SCORING (mirrors calculateScore exactly) ─────────────────────────────────

# Same synonym/abbreviation set used in CVAnalysis and JobHunter — plain
# substring matching alone produces false-negative "gaps"/"missing skills"
# for skills genuinely present but phrased/abbreviated/spelled differently
# than the JD's exact wording (e.g. resume says "ML", JD extraction says
# "Machine Learning"; or resume says "Dimensional Modelling" — a specific
# technique that IS a form of "Data Modeling" — which used to show as a
# false-negative gap since it's not a literal substring match at all).
_UK_TO_US_SPELLING = [
    (r"\bmodelling\b", "modeling"), (r"\blabelling\b", "labeling"),
    (r"\bcancelled\b", "canceled"), (r"\btravelling\b", "traveling"),
    (r"\borganisation", "organization"), (r"\bcolour", "color"),
    (r"\blicence", "license"), (r"\bcentre\b", "center"),
    (r"\bprogramme\b", "program"), (r"\banalyse", "analyze"),
    (r"\boptimise", "optimize"), (r"\bcategorise", "categorize"),
    (r"\bcustomise", "customize"), (r"\bfavour", "favor"),
    (r"\bbehaviour", "behavior"), (r"\bvisualise", "visualize"),
    (r"\bsummarise", "summarize"), (r"\bspecialise", "specialize"),
]


def _normalize_text(s: str) -> str:
    s = s.lower()
    for pattern, repl in _UK_TO_US_SPELLING:
        s = re.sub(pattern, repl, s)
    return s


# Two kinds of entries, both one-directional (key = the general/JD-style
# term; values = things that, if found in a resume, PROVE the general term
# is satisfied): true synonyms/abbreviations, and specific-technique ->
# general-skill-it's-a-form-of (curated, not fuzzy-string-guessed, so it
# doesn't cause false positives the way blind similarity matching would).
_SKILL_SYNONYMS = {
    "ai": ["artificial intelligence"], "artificial intelligence": ["ai"],
    "ml": ["machine learning"], "machine learning": ["ml",
        "regression", "classification", "neural network", "deep learning",
        "supervised learning", "unsupervised learning", "random forest",
        "gradient boosting", "xgboost", "scikit-learn", "tensorflow", "pytorch"],
    "bi": ["business intelligence"], "business intelligence": ["bi"],
    "power bi": ["powerbi", "power-bi"],
    "aws": ["amazon web services", "ec2", "s3", "redshift", "lambda",
        "aws glue", "amazon redshift", "cloudformation"],
    "amazon web services": ["aws"],
    "azure": ["microsoft azure", "azure data factory", "azure synapse",
        "azure synapse analytics", "adls", "adls gen2", "azure devops"],
    "gcp": ["google cloud platform", "google cloud", "bigquery", "gcp bigquery"],
    "google cloud platform": ["gcp"],
    "api": ["apis", "application programming interface", "rest api", "restful api", "graphql"],
    "apis": ["api"],
    "etl": ["extract transform load", "extract, transform, load", "elt",
        "data pipeline", "airflow", "dbt", "informatica", "talend", "ssis"],
    "elt": ["etl"],
    "data pipeline": ["etl", "elt", "airflow", "dbt", "data pipelines"],
    "sql": ["structured query language", "t-sql", "pl/sql", "mysql", "postgresql", "postgres"],
    "ci/cd": ["ci cd", "continuous integration", "continuous deployment", "jenkins", "github actions"],
    "devops": ["dev ops"],
    "nlp": ["natural language processing"], "natural language processing": ["nlp"],
    "llm": ["large language model", "large language models", "gpt", "generative ai"],
    "data governance": ["governance framework", "data governance framework",
        "data stewardship", "data catalog", "data cataloguing", "data lineage",
        "data quality framework", "collibra", "alation"],
    "edw": ["enterprise data warehouse"], "enterprise data warehouse": ["edw"],
    "mdm": ["master data management"], "master data management": ["mdm"],
    "data mesh": ["domain-oriented data", "data domain", "data products"],
    "kpi": ["key performance indicator"],
    "ux": ["user experience"], "ui": ["user interface"],
    "qa": ["quality assurance"],
    "pm": ["project management", "project manager"],
    "hr": ["human resources"],
    "crm": ["customer relationship management", "salesforce"],
    "erp": ["enterprise resource planning", "sap", "oracle erp", "netsuite"],
    "data modeling": [
        "dimensional modeling", "dimensional model", "data vault",
        "data vault 2.0", "star schema", "snowflake schema",
        "entity relationship modeling", "er modeling", "erd",
        "third normal form", "3nf modeling", "kimball", "inmon",
        "fsldm", "logical data modeling", "physical data modeling",
        "conceptual data modeling", "normalization", "denormalization",
    ],
    "data architecture": [
        "data mesh", "data fabric", "lakehouse", "data lakehouse",
        "enterprise data warehouse", "edw", "data lake", "data warehouse",
        "solution architecture", "enterprise architecture",
    ],
    "cloud architecture": ["aws", "azure", "gcp", "multi-cloud", "hybrid cloud"],
}


def _score_from_verdicts(strengths: dict, cv_text: str, essential_count: int, good_to_have_count: int) -> dict:
    """Builds the same score/matched/gap/bonus shape calculate_score() used
    to return, but driven by the LLM's per-item essential/good-to-have
    verdicts (utils.llm_extraction.extract_candidate_strengths) instead of
    deterministic keyword matching — same fix as CVAnalysis and JobHunter,
    for the same reason: exact/token matching can't judge long
    capability-statement requirements or requirements phrased differently
    than the resume (e.g. "Data Modeling" vs a resume that says
    "Dimensional Modeling", a specific technique that IS a form of it)."""
    matched_essential = strengths.get("essential_matched", [])
    missing_essential = strengths.get("essential_missing", [])
    matched_good = strengths.get("good_to_have_matched", [])

    essential_pct = (len(matched_essential) / essential_count * 100) if essential_count else 70
    good_bonus = min(15, len(matched_good) * 5) if good_to_have_count else 10

    cv_lower = _normalize_text(cv_text)
    bonus = 0
    reasons = []
    if re.search(r"bachelor|master|degree", cv_lower):
        bonus += 10
        reasons.append("Degree +10")
    if re.search(r"experience|\d+\s+(years|year)", cv_lower):
        bonus += 5
        reasons.append("Experience +5")

    score = min(100, round(essential_pct * 0.75 + good_bonus + bonus, 1))

    return {
        "score": score,
        "matched": matched_essential + matched_good,
        "gap": missing_essential,
        "bonus": bonus,
        "reasons": "; ".join(reasons),
    }


def _build_recommendation(
    matched: list, missing: list, matched_good_to_have: list,
    status: str, score: float, is_disqualified: bool, disqualify_reason: str,
) -> str:
    """A short, evidence-grounded recruiter recommendation for THIS
    candidate against THIS JD — composed here in plain Python from the
    essential/good-to-have verdicts the LLM already judged (see
    utils/llm_extraction.py's _verdict_chunk), never generated as a
    fresh, separate LLM narrative. Deliberately deterministic: the whole
    point is a recommendation that can never claim a strength or gap the
    underlying matched/missing lists don't already contain, so it can't
    reintroduce the exact "inferred a skill that isn't there" failure
    mode the matching prompt itself was just tightened to avoid.

    Used for the Screening (Resume Screening tab) and Screening Decision
    views — with multiple candidates applying to the same requisition,
    every one of them gets their OWN independently-computed matched/
    missing lists and their OWN recommendation text; nothing here is
    shared or averaged across candidates in the same batch."""
    essential_total = len(matched) + len(missing)

    def _list(items: list, n: int) -> str:
        return ", ".join(items[:n])

    if is_disqualified:
        reason = disqualify_reason or "a hard requirement (salary, notice period, or location) for this role"
        return (
            f"Not recommended to advance — disqualified on {reason}, independent of skills match "
            f"({len(matched)}/{essential_total} essential requirements met)."
        )

    if essential_total == 0:
        parts = [f"ATS score {round(score)}% — no essential requirements were specified for this JD to check against."]
    else:
        parts = [f"Meets {len(matched)} of {essential_total} essential requirements"]
        if matched:
            parts.append(f", including {_list(matched, 3)}")
        parts.append(".")
        if missing:
            parts.append(f" Gaps: {_list(missing, 3)}.")
        if matched_good_to_have:
            parts.append(f" Also brings {_list(matched_good_to_have, 2)} (good-to-have).")

    if status == "Qualified":
        verdict = " Recommend advancing to the next round (phone screening)."
    elif status == "Review":
        verdict = " Borderline — worth a phone screen to clarify the gaps above before deciding, rather than an outright pass."
    else:
        verdict = " Recommend not advancing — the gaps above cover core requirements for this role."

    return "".join(parts) + verdict


def calculate_score(cv_text: str, jd_skills: list) -> dict:
    """Direct port of the original JS calculateScore function, with matching
    made tolerant of synonyms/abbreviations/spelling and specific-technique
    -> general-skill relationships (same rationale as CVAnalysis's
    _skill_present — exact substring alone produces false-negative "gaps"
    for skills that are genuinely present but phrased differently)."""
    # Clean CV text same way as original, then spelling-normalize
    cv_lower = _normalize_text(cv_text)
    cv_lower = re.sub(r"\b(an|the|and|or|of|in|on|for|with|to|be|is|are|it|at|from|as|by|that|this|if|we|you)\b", "", cv_lower)
    cv_lower = re.sub(r"[^a-z0-9+#.\-]", " ", cv_lower)

    def _present(skill: str) -> bool:
        sk = _normalize_text(skill)
        if sk in cv_lower:
            return True
        for variant in _SKILL_SYNONYMS.get(sk, []):
            if _normalize_text(variant) in cv_lower:
                return True
        words = [w for w in sk.split() if len(w) > 2]
        if len(words) >= 2 and all(w in cv_lower for w in words):
            return True
        return False

    matched = [s for s in jd_skills if _present(s)]
    gap     = [s for s in jd_skills if s not in matched]

    score = (len(matched) / len(jd_skills) * 100) if jd_skills else 0
    bonus = 0
    reasons = []

    if re.search(r"bachelor|master|degree", cv_lower):
        bonus += 10
        reasons.append("Degree +10")

    if re.search(r"experience|\d+\s+(years|year)", cv_lower):
        bonus += 10
        reasons.append("Experience +10")

    score = min(100, score + bonus)

    return {
        "score":   round(score, 1),
        "matched": matched,
        "gap":     gap,
        "bonus":   bonus,
        "reasons": "; ".join(reasons),
    }


# ── QUESTION GENERATION ──────────────────────────────────────────────────────

async def generate_questions(
    jd_text: str, candidate_name: str, matched_skills: list, groq_key: str, groq_model: str = DEFAULT_GROQ_MODEL,
    resume_context: str = "",
) -> tuple[list, Optional[str]]:
    """Mirrors buildQuestionPrompt + callOllamaGenerate.

    resume_context is the candidate's OWN resume content (built from
    their already-extracted resume_summary/strengths/summary — see
    get_questions below), not just their JD's required-skills list. This
    is what makes a question like "Tell me about your work on the
    payments-processing API you led at Pinnacle Systems" possible,
    instead of only ever "Tell me about a project where you used
    Python" — the latter is answerable by anyone who's ever listed
    Python, the former only by someone whose actual resume says that.

    Returns (questions, error). error is None on a genuine LLM-generated
    set; otherwise it's a short reason the caller can surface — this used
    to silently swallow every failure and hand back _default_questions()
    with no indication anything had gone wrong, which is indistinguishable
    from a WORKING "regenerate" that just happens to produce the same
    generic questions every time (it's the same deterministic template
    based only on name+skills) — exactly what made clicking "Regenerate"
    look broken instead of surfacing the actual, fixable cause (e.g. no
    Groq key, an invalid one, or a rate limit)."""
    try:
        from langchain_groq import ChatGroq
        from langchain.schema import HumanMessage
        from utils.llm_extraction import _truncate_for_llm, _parse_json_response

        llm = ChatGroq(api_key=groq_key, model=groq_model, temperature=0.4, max_tokens=4000, reasoning_format="hidden", reasoning_effort="low", max_retries=0)
        skills_str = ", ".join(matched_skills[:8]) if matched_skills else "relevant skills"
        resume_block = (
            f"\n\nCandidate's Resume (their own experience — reference this specifically, not just their skills list):\n\"\"\"{_truncate_for_llm(resume_context, 'resume context', 4000)}\"\"\""
            if resume_context.strip() else ""
        )
        prompt = f"""You are a recruitment AI assistant preparing phone/video interview
questions for {candidate_name}.
Generate exactly 5 interview questions using BOTH the Job Description AND
the candidate's own resume below — at least 2-3 of the 5 questions should
reference something SPECIFIC from their resume (a named employer, project,
technology, or achievement), not just a generic skill from the JD's
required-skills list ({skills_str}). The goal is questions only this
specific candidate could meaningfully answer, not questions any candidate
with the same skill list would get.

Job Description:
\"\"\"{_truncate_for_llm(jd_text, "JD text", 8000)}\"\"\"{resume_block}

Return ONLY valid JSON:
{{
  "questions": ["Question 1","Question 2","Question 3","Question 4","Question 5"]
}}"""

        resp = llm.invoke([HumanMessage(content=prompt)])
        data = _parse_json_response(resp.content)
        if data is None:
            raise ValueError(f"LLM returned unparseable/empty response (length {len(resp.content)})")
        questions = data.get("questions", [])[:5]
        if not questions:
            raise ValueError("LLM response had no questions in it")
        return questions, None
    except Exception as e:
        logger.warning(f"generate_questions: LLM generation failed, falling back to default questions — {e}")
        return _default_questions(candidate_name, matched_skills), f"AI question generation failed ({str(e)[:150]}) — showing default questions instead."


def _default_questions(name: str, skills: list) -> list:
    qs = []
    if skills:
        qs.append(f"Tell me about a project where you used {skills[0]}.")
        if len(skills) > 1:
            qs.append(f"Rate your proficiency in {skills[1]} and give a real example.")
        if len(skills) > 2:
            qs.append(f"What challenges have you faced with {skills[2]}?")
    qs.append(f"Why are you the right candidate for this role, {name}?")
    qs.append("Where do you see yourself in 3 years?")
    return qs[:5]


# ── RESUME SUMMARY (10 statements) ──────────────────────────────────────────

async def generate_resume_summary(cv_text: str, groq_key: Optional[str], groq_model: str = DEFAULT_GROQ_MODEL) -> dict:
    """Produce a categorized resume summary — multiple specific bullet
    points grouped under Experience, Skills, Education, Achievements, and
    Availability & Work Rights — rather than one flat list of generic
    sentences. Each bullet should surface a genuinely relevant, specific
    detail (a role, a scale, a result, a named skill), not a filler
    restatement of the section heading. Uses Groq LLM when a key is
    available, otherwise falls back to heuristic keyword extraction."""
    if groq_key:
        try:
            from langchain_groq import ChatGroq
            from langchain.schema import HumanMessage
            from utils.llm_extraction import _truncate_for_llm, _parse_json_response

            llm = ChatGroq(api_key=groq_key, model=groq_model, temperature=0.2, max_tokens=4000, reasoning_format="hidden", reasoning_effort="low", max_retries=0)
            prompt = f"""You are a recruitment analyst producing a sharp, specific candidate
summary for a recruiter who is short on time. Read the resume below and
extract the MOST relevant and important points — prioritize specifics
(role titles, years, scale, named technologies, quantified results) over
generic statements. Do not invent facts the resume doesn't support.

Resume:
\"\"\"{_truncate_for_llm(cv_text, "resume text")}\"\"\"

Return ONLY valid JSON, no markdown, no commentary, in this exact format:
{{
  "experience": ["<specific, substantive bullet about a role/scope/achievement>", "..."],
  "skills": ["<specific bullet grouping related skills or naming a standout one>", "..."],
  "education": ["<specific bullet — degree, institution, year if stated>", "..."],
  "achievements": ["<specific, quantified accomplishment if the resume supports one>", "..."],
  "availability_work_rights": ["<specific bullet if stated>", "..."]
}}

Rules:
- experience: 3-5 bullets, each about a distinct role, project, or scope of responsibility — not one bullet per job title, but the most IMPORTANT/relevant parts of their experience
- skills: 2-4 bullets, grouping related skills together rather than one skill per bullet
- education: 1-2 bullets
- achievements: 1-3 bullets — only include if the resume actually states a concrete result/metric; omit entirely (empty list) rather than inventing one
- availability_work_rights: 0-2 bullets — only include if the resume actually mentions notice period, availability, citizenship, or work rights; omit entirely if not mentioned
- Every bullet must be a full, specific sentence a recruiter could act on — not a category label restated as a sentence"""

            resp = llm.invoke([HumanMessage(content=prompt)])
            data = _parse_json_response(resp.content)
            if data is None:
                raise ValueError(f"LLM returned unparseable/empty response (length {len(resp.content)})")
            result = {
                "experience": [s for s in data.get("experience", []) if s][:5],
                "skills": [s for s in data.get("skills", []) if s][:4],
                "education": [s for s in data.get("education", []) if s][:2],
                "achievements": [s for s in data.get("achievements", []) if s][:3],
                "availability_work_rights": [s for s in data.get("availability_work_rights", []) if s][:2],
            }
            if any(result.values()):
                return result
        except Exception:
            pass
    return _fallback_resume_summary(cv_text)


def _fallback_resume_summary(cv_text: str) -> dict:
    """Heuristic categorized resume summary used when no Groq key is configured."""
    low = cv_text.lower()
    result: dict = {"experience": [], "skills": [], "education": [], "achievements": [], "availability_work_rights": []}

    edu_m = re.search(r"(bachelor|master|phd|doctorate|diploma|degree)[^\n.]{0,90}", low)
    if edu_m:
        result["education"].append(edu_m.group().strip().capitalize() + ".")

    keyword_bank = [
        "python", "java", "javascript", "sql", "excel", "power bi", "tableau",
        "aws", "azure", "gcp", "docker", "kubernetes", "project management",
        "accounting", "communication", "leadership", "react", "node",
        "salesforce", "sap", "financial reporting", "data analysis",
    ]
    found_skills = [s for s in keyword_bank if s in low]
    if found_skills:
        result["skills"].append(f"Resume indicates experience with {', '.join(found_skills[:6])}.")

    exp_m = re.search(r"(\d{1,2})\+?\s*(?:years?|yrs?)\s*(?:of\s+)?experience", low)
    if exp_m:
        result["experience"].append(f"Approximately {exp_m.group(1)}+ years of relevant experience.")
    if re.search(r"present|current(ly)?\s+work", low):
        result["experience"].append("Resume references a current or recent professional role.")
    if re.search(r"project", low):
        result["experience"].append("Resume highlights project-based work experience.")
    if not result["experience"]:
        result["experience"].append("Years and scope of experience not explicitly stated in resume.")

    if re.search(r"immediate(ly)?\s+available|available\s+immediately|notice\s+period", low):
        avail_m = re.search(r"([a-z0-9 ]{0,10}notice\s+period[a-z0-9 ]{0,15}|immediate(ly)?\s+available)", low)
        result["availability_work_rights"].append((avail_m.group().strip().capitalize() + ".") if avail_m else "Availability mentioned in resume.")
    if re.search(r"citizen|permanent resident|\bpr\b|work visa|work rights|unrestricted work rights|485 visa|482 visa|sponsorship", low):
        cit_m = re.search(r"([a-z0-9 ]{0,20}(citizen|permanent resident|work visa|work rights|sponsorship)[a-z0-9 ]{0,20})", low)
        result["availability_work_rights"].append((cit_m.group().strip().capitalize() + ".") if cit_m else "Citizenship/work rights mentioned in resume.")

    if re.search(r"certificat", low):
        result["achievements"].append("Resume mentions professional certification(s).")

    return result


# ── SMTP EMAIL SENDING ────────────────────────────────────────────────────────
# Moved to utils/email_send.py so the Interview Management capability can
# send Calendly-link emails through the exact same SMTP plumbing (Phone
# Interview's "Send Calendly Link" / Interview Scheduling's "Email
# Calendly Link"). Re-exported under the original private names here so
# nothing else in this file has to change.
from utils.email_send import get_smtp_config as _get_smtp_config, send_email as _send_email


class SendInviteRequest(BaseModel):
    to_email: str
    subject: str
    body_html: str


# ── FORMAT CANDIDATE ─────────────────────────────────────────────────────────

def _fmt(c: JobLensCandidate, phone_schedule: Optional[dict] = None) -> dict:
    # phone_schedule: this candidate's entry from _phone_schedule_map()
    # below (Interview Scheduling row for the Phone Interview round) —
    # None/missing when no such row exists yet (candidate never
    # contacted, texted, or Calendly-linked).
    ps = phone_schedule or {}
    return {
        "id": c.id,
        "name": c.name,
        "email": c.email,
        "phone": c.phone,
        "filename": c.filename,
        "ats_score": round(c.ats_score, 1),
        "status": c.status,
        "screening_recommendation": c.screening_recommendation or "",
        "matched_skills": c.matched_skills or [],
        "missing_skills": c.missing_skills or [],
        "bonus": c.bonus,
        "bonus_reasons": c.bonus_reasons,
        "experience_years": c.experience_years or "",
        "summary": c.summary or "",
        "resume_summary": c.resume_summary or [],
        "interview_token": c.interview_token,
        "contacted": bool(c.contacted),
        "video_status": c.video_status,
        "shortlisted": c.shortlisted,
        "interview_questions": c.interview_questions or [],
        "emotion_happy": c.emotion_happy,
        "emotion_neutral": c.emotion_neutral,
        "emotion_sad": c.emotion_sad,
        "emotion_angry": c.emotion_angry,
        "emotion_fear": c.emotion_fear,
        "emotion_disgust": c.emotion_disgust,
        "emotion_surprise": c.emotion_surprise,
        "dominant_emotion": c.dominant_emotion,
        "has_resume_file": bool(c.resume_file_blob),
        # BUG FIX: this used to be bool(c.video_blob) only — but
        # upload_interview_video() clears video_blob to None whenever S3
        # storage succeeds (the preferred path; see upload_video_and_get_key),
        # storing video_key instead. Any candidate whose video actually went
        # to S3 therefore had has_video come back False here — which hides
        # the "View recorded video" link AND every analysis-status/results
        # block below (all gated on has_video), even though the video and
        # its analysis are both there. Whichever storage path was used,
        # having either field set means a video exists.
        "has_video": bool(c.video_blob) or bool(c.video_key),
        "video_transcript": c.video_transcript or "",
        "video_analysis": c.video_analysis or None,
        "video_analysis_status": c.video_analysis_status or "Pending",
        "qa_evaluation": c.qa_evaluation or None,
        "qa_evaluation_score": c.qa_evaluation_score,
        "source_vendor_id": c.source_vendor_id,
        "source_vendor_name": c.source_vendor_name or "",
        "strengths_breakdown": c.strengths_breakdown or None,
        # ── Dual-track scoring (RevaMatrix-AI parity) ──────────────────
        "technical_score": round(c.technical_score, 1) if c.technical_score is not None else None,
        "non_technical_score": round(c.non_technical_score, 1) if c.non_technical_score is not None else None,
        "logistics": c.logistics or None,
        "hard_disqualified": bool(c.hard_disqualified),
        "disqualify_reason": c.disqualify_reason,
        # ── Phone Interview (split-out stage) ──────────────────────────
        "phone_screening_status": c.phone_screening_status or "Not Started",
        "phone_screening_recommendation": c.phone_screening_recommendation or "",
        "phone_screening_notes": c.phone_screening_notes or "",
        "phone_screening_at": c.phone_screening_at.isoformat() if c.phone_screening_at else None,
        # ── Video Interview — Decision & Comments ───────────────────────
        "video_screening_recommendation": c.video_screening_recommendation or "",
        "video_screening_notes": c.video_screening_notes or "",
        "video_screening_at": c.video_screening_at.isoformat() if c.video_screening_at else None,
        # ── Phone Interview: telephony (click-to-call + SMS scheduling) ──
        # Pulled from this candidate's Interview Scheduling row (Phone
        # Interview round) — see _phone_schedule_map below. Absent
        # (None/"") until that row exists.
        "phone_interview_scheduled_at": ps.get("scheduled_at"),
        "phone_interview_status": ps.get("status", ""),
        "phone_call_status": ps.get("phone_call_status", ""),
        "phone_called_at": ps.get("phone_called_at"),
        "call_sms_sent_at": ps.get("call_sms_sent_at"),
        "phone_transcript": ps.get("phone_transcript", ""),
        "phone_transcript_status": ps.get("phone_transcript_status", ""),
        # ── Screening Decision's "Send Rejection Email" bulk action ──────
        # Was being set correctly by send_rejection_emails() but never
        # actually returned here, so the Screening Decision table always
        # showed "Not sent" no matter how many times an email went out —
        # the send genuinely worked, this was purely a serialization gap.
        "rejection_email_sent_at": c.rejection_email_sent_at.isoformat() if c.rejection_email_sent_at else None,
    }


async def _phone_schedule_map(db: AsyncSession, candidate_ids: list) -> dict:
    """Bulk-fetches each candidate's Interview Scheduling row for the
    Phone Interview round in ONE query, keyed by joblens_candidate_id —
    used to enrich _fmt() output for a page of candidates without an
    N+1 query per row. Returns {} entries are simply omitted (caller
    treats a missing key the same as "no row yet")."""
    if not candidate_ids:
        return {}
    from capabilities.interview.models import Interview
    rows = (await db.execute(
        select(Interview).where(
            Interview.joblens_candidate_id.in_(candidate_ids),
            Interview.interview_type == "Phone Interview",
        )
    )).scalars().all()
    return {
        i.joblens_candidate_id: {
            "scheduled_at": i.scheduled_at.isoformat() if i.scheduled_at else None,
            "status": i.status,
            "phone_call_status": i.phone_call_status or "",
            "phone_called_at": i.phone_called_at.isoformat() if i.phone_called_at else None,
            "call_sms_sent_at": i.call_sms_sent_at.isoformat() if i.call_sms_sent_at else None,
            "phone_transcript": i.phone_transcript or "",
            "phone_transcript_status": i.phone_transcript_status or "",
        }
        for i in rows
    }


# ══════════════════════════════════════════════════════════════════════════════
# ENDPOINTS
# ══════════════════════════════════════════════════════════════════════════════

@router.post("/fetch-jd-url")
async def fetch_jd_url(
    payload: dict,
    current_user: User = Depends(get_current_user),
):
    """Identical to CVIntel's /api/cvintel/fetch-jd-url (same shared
    utils/jd_url_fetch.py) — kept as a same-shaped endpoint under this
    router too so the CandidateLens frontend doesn't need to cross-call
    CVIntel's API. Converts a job-posting URL into JD text for the
    frontend to drop into the same "paste JD text" field used everywhere
    else — the bulk analysis that follows is unaffected by how the text
    arrived.

    Body: {"url": "https://..."}
    """
    from utils.jd_url_fetch import fetch_jd_from_url, JDFetchError

    url = (payload.get("url") or "").strip()
    if not url:
        raise HTTPException(400, "A URL is required.")

    try:
        result = await fetch_jd_from_url(url)
    except JDFetchError as e:
        raise HTTPException(422, str(e))

    return result


async def _check_candidate_quota(db: AsyncSession, user: User, requested_count: int) -> None:
    """Enforces the current plan's Max Candidates cap — checked BEFORE
    any resume processing starts (JD extraction, LLM calls, etc.), so a
    batch that would blow through the limit is rejected up front rather
    than partway through, after quota-consuming work has already run.

    Counted as "candidates processed this calendar month" across ALL of
    this user's JobLens sessions combined — the plan's quota is a
    monthly allowance for the whole account, not a per-JD/per-session
    limit, so uploading against three different JDs in the same month
    still draws from one shared pool.

    No plan on record, or a plan with max_candidates=0 (unlimited),
    both skip the check entirely rather than blocking — this is an
    additive safety net, not the only gate on who can use the feature at
    all."""
    from models.billing_models import Subscription, PricingPlan
    sub = (await db.execute(select(Subscription).where(Subscription.user_id == user.id))).scalar_one_or_none()
    if not sub or not sub.plan_slug:
        return
    plan = (await db.execute(select(PricingPlan).where(PricingPlan.slug == sub.plan_slug))).scalar_one_or_none()
    if not plan or not plan.max_candidates:
        return

    now = datetime.utcnow()
    month_start = datetime(now.year, now.month, 1)
    processed_this_month = (await db.execute(
        select(func.count(JobLensCandidate.id))
        .join(JobLensSession, JobLensCandidate.session_id == JobLensSession.id)
        .where(JobLensSession.user_id == user.id, JobLensCandidate.created_at >= month_start)
    )).scalar() or 0

    if processed_this_month + requested_count > plan.max_candidates:
        remaining = max(0, plan.max_candidates - processed_this_month)
        raise HTTPException(
            403,
            f"This batch of {requested_count} would exceed your {plan.name} plan's monthly limit of "
            f"{plan.max_candidates} candidates ({processed_this_month} already processed this month, "
            f"{remaining} remaining). Upgrade your plan to continue, or reduce this batch to {remaining} or fewer.",
        )


@router.post("/run")
async def run_joblens(
    jd_text: str = Form(""),
    low_threshold: int = Form(40),
    high_threshold: int = Form(70),
    jd_file: Optional[UploadFile] = File(None),
    cv_files: List[UploadFile] = File(default=[]),
    jd_record_id: Optional[int] = Form(None),          # NEW: pull JD from JD Management instead of text/file
    source_candidate_ids: str = Form(""),               # NEW: comma-separated TrackedCandidate ids from Vendor Management
    requisition_id: Optional[int] = Form(None),          # NEW: pull JD + submitted candidates from a Requisition
    source_application_ids: str = Form(""),              # NEW: comma-separated Application ids submitted for that requisition
    # ── Dual-track scoring: logistics constraints + dynamic weights ─────
    # If jd_record_id is set AND that JD Management record already has
    # logistics constraints saved, those take precedence (single source of
    # truth); otherwise these form fields are used directly — lets a
    # recruiter run CandidateLens straight off pasted/uploaded JD text
    # without first creating a JD Management record.
    salary_budget_min: int = Form(0),
    salary_budget_max: int = Form(0),
    max_notice_days: int = Form(0),
    remote_allowed: bool = Form(False),
    weights: Optional[str] = Form(None),         # JSON — see utils/scoring.DEFAULT_WEIGHTS
    disqualifiers: Optional[str] = Form(None),   # JSON — see utils/scoring.DEFAULT_DISQUALIFIERS
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    from utils.scoring import merge_weights, merge_disqualifiers
    import json as _json
    try:
        weight_overrides = _json.loads(weights) if weights else None
    except (_json.JSONDecodeError, TypeError):
        raise HTTPException(400, "weights must be a valid JSON object")
    try:
        disqualifier_overrides = _json.loads(disqualifiers) if disqualifiers else None
    except (_json.JSONDecodeError, TypeError):
        raise HTTPException(400, "disqualifiers must be a valid JSON object")
    session_weights = merge_weights(weight_overrides)
    session_disqualifiers = merge_disqualifiers(disqualifier_overrides)

    # Plan quota — checked here, before ANY resume processing (JD
    # extraction, LLM calls) starts, using an upper-bound count of every
    # possible candidate source in this request (uploaded files +
    # sourced tracked-candidates + submitted applications combined).
    requested_count = len(cv_files or [])
    requested_count += len([x for x in source_candidate_ids.split(",") if x.strip()])
    requested_count += len([x for x in source_application_ids.split(",") if x.strip()])
    if requested_count > 0:
        await _check_candidate_quota(db, current_user, requested_count)

    # ── Requisition source: resolves requisition_id into whichever of
    # jd_record_id / a direct JD file it carries, so everything below
    # (skills extraction, logistics constraints, candidate sourcing) can
    # keep working off the existing jd_record_id-based plumbing without
    # duplicating it. Requisition's OWN attached JD file (Text/Word/PDF)
    # takes precedence over its linked JD Management record when both
    # are present — it's the more specific, requisition-level document.
    requisition_jd_text: Optional[str] = None
    requisition = None
    if requisition_id:
        from capabilities.requisition.models import Requisition
        from capabilities.acquisition import service as acquisition_service
        org = await acquisition_service.get_or_create_default_organisation(db, current_user)
        requisition = (await db.execute(
            select(Requisition).where(Requisition.id == requisition_id, Requisition.organisation_id == org.id)
        )).scalar_one_or_none()
        if not requisition:
            raise HTTPException(404, "Selected requisition not found.")
        # Requisition's own JD file can live in either place depending on
        # when it was uploaded: jd_file_blob (legacy, stored directly in
        # the DB row) or jd_file_key (current — an S3/R2 object key,
        # fetched via get_file_bytes). This previously only checked
        # jd_file_blob, so a requisition whose JD was uploaded AFTER the
        # move to S3/R2 storage (jd_file_key set, jd_file_blob left null)
        # was incorrectly reported as having no JD attached at all, even
        # though one was clearly sitting in the bucket — see the same
        # jd_file_key-first pattern already used in
        # capabilities/requisition/router.py and candidatetrack.py.
        if requisition.jd_file_key or requisition.jd_file_blob:
            jd_bytes = None
            if requisition.jd_file_key:
                jd_bytes = await get_file_bytes(db, requisition.jd_file_key)
            if not jd_bytes:
                jd_bytes = requisition.jd_file_blob
            if jd_bytes:
                extracted = extract_text(jd_bytes, requisition.jd_file_filename or "jd.txt")
                if extracted.strip():
                    requisition_jd_text = extracted
        if not requisition_jd_text and not requisition.jd_record_id:
            raise HTTPException(400, "This requisition has no JD attached yet — attach one on the Requisitions page first.")
        # Falls through to the jd_record_id branch below when this
        # requisition has no own JD file but does have a linked JD
        # Management record.
        if not requisition_jd_text:
            jd_record_id = requisition.jd_record_id

    # ── Extract JD — either from JD Management (new) or text/file upload (existing) ──
    final_jd = jd_text.strip()
    jd_client_name = ""
    jd_record = None
    if requisition_jd_text:
        final_jd = requisition_jd_text
        if requisition.client_id:
            from models.models import Client as ClientModel
            cr = await db.execute(select(ClientModel).where(ClientModel.id == requisition.client_id))
            client = cr.scalar_one_or_none()
            if client:
                jd_client_name = client.name
    elif jd_record_id:
        from models.models import JDRecord, Client as ClientModel
        jr = await db.execute(select(JDRecord).where(JDRecord.id == jd_record_id, JDRecord.user_id == current_user.id))
        jd_record = jr.scalar_one_or_none()
        if not jd_record:
            raise HTTPException(404, "Selected JD not found.")
        # Use the JD's stored description as the JD text; title/client used for the summary panel
        final_jd = jd_record.description or jd_record.title
        jd_client_name = ""
        if jd_record.client_id:
            cr = await db.execute(select(ClientModel).where(ClientModel.id == jd_record.client_id))
            client = cr.scalar_one_or_none()
            if client:
                jd_client_name = client.name
    elif jd_file and jd_file.filename:
        raw = await jd_file.read()
        extracted = extract_text(raw, jd_file.filename)
        if extracted.strip():
            final_jd = extracted

    # ── Resolve candidate sources: uploaded files (existing) + Vendor
    #    Management submissions (existing) + Applications submitted for a
    #    Requisition (new) — all can be used together.
    from models.models import TrackedCandidate, Vendor as VendorModel
    source_candidates = []
    if source_candidate_ids.strip():
        ids = [int(x) for x in source_candidate_ids.split(",") if x.strip().isdigit()]
        if ids:
            tcr = await db.execute(
                select(TrackedCandidate).where(TrackedCandidate.id.in_(ids), TrackedCandidate.user_id == current_user.id)
            )
            source_candidates = tcr.scalars().all()

    # Candidates submitted for a Requisition — Application (Candidate <->
    # Requisition), NOT TrackedCandidate. Resume comes from the Candidate
    # Master's own resume_blob, same record used everywhere else.
    from capabilities.acquisition.models import Application, Candidate
    source_applications = []
    if source_application_ids.strip():
        ids = [int(x) for x in source_application_ids.split(",") if x.strip().isdigit()]
        if ids:
            acr = await db.execute(
                select(Application, Candidate)
                .join(Candidate, Candidate.id == Application.candidate_id)
                .where(Application.id.in_(ids))
            )
            source_applications = acr.all()

    if not final_jd:
        raise HTTPException(400, "Job description is required.")
    if not cv_files and not source_candidates and not source_applications:
        raise HTTPException(400, "At least one CV is required (upload files, or select candidates submitted for the requisition).")

    # ── Resolve logistics constraints: linked JD Management record wins
    # (single source of truth) if it has them set; otherwise use whatever
    # was passed directly on this request.
    session_salary_min, session_salary_max = salary_budget_min, salary_budget_max
    session_max_notice, session_remote = max_notice_days, remote_allowed
    if jd_record_id and jd_record and (jd_record.salary_budget_max or jd_record.max_notice_days):
        session_salary_min = jd_record.salary_budget_min or session_salary_min
        session_salary_max = jd_record.salary_budget_max or session_salary_max
        session_max_notice = jd_record.max_notice_days or session_max_notice
        session_remote = jd_record.remote_allowed or session_remote

    # ── Get Groq key (own key first, else the healthiest key in the shared
    # pool — see utils/groq_pool.py) ──
    from utils.groq_pool import resolve_groq_key
    key_resolution = await resolve_groq_key(db, current_user.id)
    groq_key = key_resolution["groq_key"]
    groq_model = await get_groq_model(db, current_user.id)
    ollama_creds = await get_all_credentials(db, current_user.id, "ollama") if ollama_enabled() else {}
    ollama_base_url = ollama_creds.get("base_url")
    ollama_model = ollama_creds.get("model")
    from utils.llm_extraction import (
        get_taxonomy_hint, get_semantic_taxonomy_hint, enrich_skill_taxonomy,
        extract_jd_requirements_categorized,
    )
    known_terms = await get_semantic_taxonomy_hint(db, final_jd)

    # ── Extract JD details: role, location, company, categorized skills ──
    # Uses the SAME extract_jd_requirements_categorized function CVIntel
    # uses (see utils/llm_extraction.py) instead of a separate CandidateLens-
    # only prompt — CVIntel and CandidateLens used to run two DIFFERENT LLM
    # prompts to categorize the same JD's essential/good-to-have
    # requirements, which alone could produce different essential-skill
    # lists (and therefore different match percentages) for the identical
    # JD text between the two modules, before either one scored anything.
    # This also gives CandidateLens the same min_years_experience/
    # education_requirement/logistics fields CVIntel already had.
    #
    # If this session was started from an existing JD Management record
    # that already has persisted categorized requirements, reuse those
    # directly — re-extracting the same JD's requirements via LLM on every
    # single analysis run would be wasteful and could drift from what's
    # shown in JD Management itself.
    if jd_record_id and jd_record and (jd_record.essential_skills or jd_record.good_to_have_skills):
        jd_req = {
            "role": jd_record.title,
            "location": "",
            "company": jd_client_name,
            "essential": jd_record.essential_skills or [],
            "good_to_have": jd_record.good_to_have_skills or [],
            "optional": jd_record.optional_skills or [],
            "min_years_experience": jd_record.min_years_experience or 0,
            "education_requirement": jd_record.education_requirement or "",
            "salary_budget_min": jd_record.salary_budget_min or 0,
            "salary_budget_max": jd_record.salary_budget_max or 0,
            "max_notice_days": jd_record.max_notice_days or 0,
            "remote_allowed": bool(jd_record.remote_allowed),
        }
    else:
        jd_req = await extract_jd_requirements_categorized(
            final_jd, groq_key, groq_model, ollama_base_url, ollama_model, known_terms,
            db=db, user_id=current_user.id,
        )
    # jd_details kept as an alias — the rest of this function (JD summary
    # panel, interview-question context, etc.) reads from jd_details, and
    # "skills" (a flat, deduped list) is only this module's own convenience
    # field, not part of the shared jd_req shape.
    jd_details = {
        **jd_req,
        "skills": list(dict.fromkeys(
            (jd_req.get("essential") or []) + (jd_req.get("good_to_have") or []) + (jd_req.get("optional") or [])
        )),
    }
    jd_skills = jd_details["skills"]

    # ── Create session ─────────────────────────────────────────────────
    try:
        seq_num = await next_sequence_number(db, JobLensSession, current_user.id)
        session = JobLensSession(
            sequence_number=seq_num,
            user_id=current_user.id,
            jd_text=final_jd,
            jd_skills=jd_skills,
            jd_role=jd_details["role"],
            jd_location=jd_details["location"],
            jd_company=jd_details["company"],
            jd_record_id=jd_record_id,
            jd_client_name=jd_client_name,
            jd_essential_skills=jd_details.get("essential", []),
            jd_good_to_have_skills=jd_details.get("good_to_have", []),
            jd_optional_skills=jd_details.get("optional", []),
            low_threshold=low_threshold,
            high_threshold=high_threshold,
            status="completed",
            cv_count=len(cv_files) + len(source_candidates) + len(source_applications),
            created_at=datetime.utcnow(),
            salary_budget_min=session_salary_min,
            salary_budget_max=session_salary_max,
            max_notice_days=session_max_notice,
            jd_remote_allowed=session_remote,
            weights=session_weights,
            disqualifiers=session_disqualifiers,
        )
        db.add(session)
        await db.commit()
        await db.refresh(session)
    except Exception as e:
        raise HTTPException(500, f"DB error creating session: {str(e)}")

    # Bounds actual Groq API calls directly — the real rate-limit concern.
    # Previously there were two separate, uncoordinated concurrency limits
    # (how many candidates process at once, and — as of the last fix — how
    # many calls run per candidate) that could MULTIPLY against each other:
    # N candidates x 3 concurrent calls each = up to 3N simultaneous Groq
    # requests for one batch, an easy way to trip rate limits that neither
    # limit alone would suggest. One shared semaphore around every actual
    # Groq call means peak concurrency is bounded by this single number,
    # no matter how many candidates or calls-per-candidate are involved.
    #
    # Was raised from 4 to 8 to cut wall-clock time on larger batches of
    # longer, realistic resumes. In practice, on this deployment's
    # container resources, that coincided with a hard crash mid-batch
    # (container restart logged, request came back as a 503) — 8
    # concurrent Groq calls each holding a full multi-page resume's text
    # in memory, layered on top of concurrently parsing multiple .docx
    # files, pushed peak memory usage past what this container has
    # available. Dialed back down to 5 as a safer middle ground: still
    # faster than the original 4, without doubling peak concurrent
    # memory usage on a container that's already shown it can't absorb
    # that. If crashes persist even at 5, the real fix is more memory
    # for this service on Northflank, not a lower number here — this
    # semaphore trades speed for memory pressure, it can't create memory
    # that isn't there.
    _groq_semaphore = asyncio.Semaphore(5)

    async def _with_groq_limit(coro):
        async with _groq_semaphore:
            return await coro

    # ── Score each CV ──────────────────────────────────────────────────
    async def _score_and_build_candidate(
        content: bytes, filename: str,
        source_vendor_id=None, source_vendor_name=None, source_tracked_candidate_id=None,
        source_application_id=None,
        call_groq_key=None, call_groq_model=None,
    ):
        # Falls back to the shared groq_key/groq_model if no per-candidate
        # key was resolved for this call — keeps this function safe to
        # call exactly as before if a caller doesn't opt into per-candidate
        # pool draws.
        _groq_key = call_groq_key if call_groq_key is not None else groq_key
        _groq_model = call_groq_model if call_groq_model is not None else groq_model

        cv_text = extract_text(content, filename)
        if not cv_text.strip():
            return None

        info = extract_candidate_info(cv_text, filename)

        # These three calls are independent of each other's OUTPUTS (questions
        # generation previously waited on the matched-skills result from
        # strengths extraction purely for a "focus on these skills" hint —
        # the JD's own essential/good_to_have list serves that just as well
        # and is known immediately, with no need to wait). Running them
        # concurrently instead of sequentially cuts per-candidate latency
        # from ~3x a single Groq call to ~1x the slowest of the three —
        # on top of the existing between-candidate concurrency, this is
        # where the real "3 resumes taking forever" time was going.
        #
        # NOTE: db/user_id are deliberately NOT passed to
        # extract_candidate_strengths here — this function is already
        # called concurrently across MULTIPLE candidates sharing the same
        # DB session (see _score_with_limit below), so enabling its
        # internal per-CHUNK pool resolution here too would mean multiple
        # candidates' internal DB calls racing on the same session, which
        # SQLAlchemy's AsyncSession doesn't allow. Instead, each CANDIDATE
        # gets its own pre-resolved key up front (see the dispatch loop
        # below) — multi-key parallelism happens at the candidate level
        # here, not the chunk level within one candidate.
        from utils.llm_extraction import extract_candidate_strengths
        jd_focus_skills = (jd_details.get("essential", []) + jd_details.get("good_to_have", []))[:8]

        strengths_task = _with_groq_limit(extract_candidate_strengths(
            cv_text,
            {"essential": jd_details.get("essential", []), "good_to_have": jd_details.get("good_to_have", [])},
            _groq_key, _groq_model,
            ollama_base_url=ollama_base_url, ollama_model=ollama_model, known_terms_hint=known_terms,
            # db (not user_id) enables the extraction cache — see
            # models.ExtractionCache's docstring. This directly fixes the
            # "same resume+JD scores differently on repeat runs" bug: a
            # genuine repeat now returns the cached result instead of
            # re-racing LLM providers. user_id stays None deliberately —
            # multiple candidates are scored concurrently here, and passing
            # user_id too would also re-enable this function's internal
            # API-key-pool resolution, which DOES use the shared `db`
            # session directly and isn't safe under that concurrency (the
            # cache itself uses its own isolated session, so it has no such
            # restriction).
            db=db,
        ))
        questions_task = (
            _with_groq_limit(generate_questions(final_jd, info["name"], jd_focus_skills, _groq_key, _groq_model, resume_context=cv_text))
            if _groq_key else asyncio.sleep(0, result=None)
        )
        summary_task = _with_groq_limit(generate_resume_summary(cv_text, _groq_key, _groq_model))

        strengths_breakdown, questions_result, resume_summary = await asyncio.gather(
            strengths_task, questions_task, summary_task,
        )

        # Same shared technical-scoring engine CVIntel uses (see
        # utils/technical_scoring.py's docstring for why this replaced the
        # old _score_from_verdicts, which awarded a flat +15 regex bonus
        # for merely mentioning "degree"/"N years experience" ANYWHERE in
        # the resume, independent of whether that matched THIS JD's actual
        # requirements — the main reason CandidateLens scores ran higher
        # than CVIntel's for identical resume/JD pairs).
        from utils.technical_scoring import compute_technical_score
        result = compute_technical_score(jd_req, strengths_breakdown, cv_text, session_weights)

        # ── Dual-track composite: technical (above) + non-technical/
        # logistics (salary/notice/location), combined via this session's
        # weights, with hard disqualifiers checked independently of score.
        from utils.scoring import compute_non_technical_score, check_hard_disqualifiers, compute_composite_score
        technical_score = result["overall"]
        candidate_logistics = {
            "expected_salary": strengths_breakdown.get("expected_salary") or 0,
            "notice_period_days": strengths_breakdown.get("notice_period_days", -1),
            "current_location": strengths_breakdown.get("current_location") or "",
            "salary_budget_min": session_salary_min,
            "salary_budget_max": session_salary_max,
            "max_notice_days": session_max_notice,
            "jd_location": jd_details.get("location") or "",
            "remote_allowed": session_remote,
        }
        non_technical = compute_non_technical_score(candidate_logistics, session_weights)
        is_disqualified, disqualify_reason = check_hard_disqualifiers(candidate_logistics, session_disqualifiers)
        score = compute_composite_score(technical_score, non_technical, session_weights)

        status = "Not Qualified"
        if score >= high_threshold:
            status = "Qualified"
        elif score >= low_threshold:
            status = "Review"
        if is_disqualified:
            # Hard disqualifiers override the score-based tier entirely —
            # a technically strong candidate who fails a hard business
            # constraint (notice period, salary overrun) should never
            # show as "Qualified"/"Review", mirroring how a recruiter
            # would actually triage.
            status = "Not Qualified"

        recommendation = _build_recommendation(
            result["matched"], result["missing"], result.get("matched_good_to_have", []),
            status, score, is_disqualified, disqualify_reason,
        )

        questions = questions_result[0] if groq_key and questions_result else _default_questions(info["name"], result["matched"])

        fname_lower = (filename or "").lower()
        if fname_lower.endswith(".pdf"):
            resume_mimetype = "application/pdf"
        elif fname_lower.endswith(".docx"):
            resume_mimetype = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        elif fname_lower.endswith(".doc"):
            resume_mimetype = "application/msword"
        else:
            resume_mimetype = "text/plain"

        resume_ext = {"application/pdf": "pdf",
                      "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
                      "application/msword": "doc"}.get(resume_mimetype, "txt")
        resume_uploaded = await upload_file(
            db, "resumes", current_user.id, session.id, content, resume_mimetype, resume_ext,
        )

        return JobLensCandidate(
            session_id=session.id,
            name=info["name"],
            email=info["email"],
            phone=info["phone"],
            filename=filename,
            ats_score=score,
            status=status,
            screening_recommendation=recommendation,
            matched_skills=result["matched"],
            missing_skills=result["missing"],
            bonus=result["matched_good_to_have"] and min(15, len(result["matched_good_to_have"]) * 5) or 0,
            bonus_reasons=(
                f"Skills {result['skills_pct']}% · Experience {result['experience_pct']}% · Education {result['education_pct']}%"
            ),
            experience_years=info.get("experience_years", ""),
            summary=info.get("summary", ""),
            resume_summary=resume_summary,
            technical_score=technical_score,
            non_technical_score=non_technical["score"],
            logistics={
                **candidate_logistics,
                "salary_score": non_technical["salary_score"],
                "notice_score": non_technical["notice_score"],
                "location_score": non_technical["location_score"],
                "applicable": non_technical["applicable"],
            },
            hard_disqualified=is_disqualified,
            disqualify_reason=disqualify_reason,
            interview_questions=questions,
            video_status="Pending",
            shortlisted=False,
            resume_file_blob=None if resume_uploaded else content,
            resume_file_mimetype=resume_mimetype,
            resume_key=resume_uploaded["key"] if resume_uploaded else None,
            resume_size_bytes=resume_uploaded["size_bytes"] if resume_uploaded else len(content),
            source_vendor_id=source_vendor_id,
            source_vendor_name=source_vendor_name,
            source_tracked_candidate_id=source_tracked_candidate_id,
            source_application_id=source_application_id,
            strengths_breakdown={
                "essentialMatched": strengths_breakdown.get("essential_matched", []),
                "technicalSkills": strengths_breakdown.get("technical_skills", []),
                "businessSkills": strengths_breakdown.get("business_skills", []),
                "softSkills": strengths_breakdown.get("soft_skills", []),
                "significantExperience": strengths_breakdown.get("significant_experience", []),
                "certificationsDegrees": strengths_breakdown.get("certifications_degrees", []),
                "yearsExperience": strengths_breakdown.get("years_experience", 0),
                "education": strengths_breakdown.get("education", ""),
                "aiPowered": strengths_breakdown.get("ai_powered", False),
                # ── Full tier + type score breakdown (Essential/Good-to-
                # Have/Qualification/Technical/Tools/Domain/Soft Skills +
                # final ATS) — same shape CVIntel exposes as
                # "scoreBreakdown", so both modules' frontends can share
                # one display component.
                "scoreBreakdown": {
                    "essential": {"pct": result["essential_pct"], "label": "Essential Requirements"},
                    "goodToHave": {"pct": result["good_to_have_pct"], "label": "Good to Have"},
                    "qualification": {"pct": result["qualification_pct"], "label": "Qualification / Education"},
                    "technical": {**result["category_breakdown"]["technical"], "label": "Technical Skills"},
                    "tools": {**result["category_breakdown"]["tool"], "label": "Tools & Platforms"},
                    "domain": {**result["category_breakdown"]["domain"], "label": "Domain Knowledge"},
                    "softSkills": {**result["category_breakdown"]["soft_skill"], "label": "Soft Skills"},
                    "finalATS": score,
                },
            },
        )

    candidates = []
    # This bounds how many candidates are processed "in flight" at once —
    # local work (file parsing, regex extraction) is cheap, so this can be
    # looser than the real rate-limit concern. Actual Groq API calls are
    # bounded separately by _groq_semaphore above, regardless of how many
    # candidates are in flight here; their calls simply queue for a slot.
    _score_semaphore = asyncio.Semaphore(8)

    async def _score_with_limit(content: bytes, filename: str, call_groq_key=None, call_groq_model=None):
        async with _score_semaphore:
            try:
                return await _score_and_build_candidate(content, filename, call_groq_key=call_groq_key, call_groq_model=call_groq_model)
            except Exception as e:
                print(f"Error processing {filename}: {e}")
                return None

    upload_payloads = [(await upload.read(), upload.filename) for upload in cv_files]

    # Resolve a key for EACH candidate SEQUENTIALLY before any concurrent
    # work starts — key resolution is a DB call, and these candidates are
    # about to be scored CONCURRENTLY sharing this same session, so every
    # DB touch for key selection has to happen up front, one at a time,
    # never inside the concurrent gather itself. This is what gives a
    # single CandidateLens batch genuine multi-key parallelism: different
    # candidates draw different pool keys, all processed at the same time.
    candidate_key_assignments = []
    for _ in upload_payloads:
        from utils.groq_pool import resolve_groq_key
        kr = await resolve_groq_key(db, current_user.id)
        candidate_key_assignments.append((kr["groq_key"], kr["model"] or groq_model, kr["pool_id"] if kr["source"] == "pool" else None))

    scored = await asyncio.gather(*[
        _score_with_limit(content, filename, call_groq_key=candidate_key_assignments[i][0], call_groq_model=candidate_key_assignments[i][1])
        for i, (content, filename) in enumerate(upload_payloads)
    ])

    # Report each candidate's key outcome sequentially, now that all
    # concurrent scoring is done — same reasoning as above, this is a DB
    # write and must not race with other DB access on this session.
    for i, candidate in enumerate(scored):
        _, _, pool_id = candidate_key_assignments[i]
        if pool_id is None:
            continue
        from utils.groq_pool import record_key_outcome
        succeeded = bool(candidate and (candidate.strengths_breakdown or {}).get("aiPowered"))
        await record_key_outcome(db, pool_id, success=succeeded)

    for candidate in scored:
        if candidate:
            db.add(candidate)
            candidates.append(candidate)
            # Enrichment does its own db.commit() — done here, sequentially,
            # after all concurrent scoring has finished, rather than inside
            # _score_and_build_candidate itself (which runs concurrently
            # across candidates — committing from multiple coroutines on
            # the same session at once is not safe with AsyncSession).
            sb = candidate.strengths_breakdown or {}
            if sb.get("aiPowered"):
                await enrich_skill_taxonomy(db, {
                    "technical": sb.get("technicalSkills", []),
                    "business": sb.get("businessSkills", []),
                    "soft": sb.get("softSkills", []),
                })

    # ── Score candidates sourced from Vendor Management submissions ──────
    for tc in source_candidates:
        try:
            if not tc.resume_blob:
                continue
            vendor_row = await db.execute(select(VendorModel).where(VendorModel.id == tc.vendor_id))
            vendor = vendor_row.scalar_one_or_none()
            candidate = await _score_and_build_candidate(
                tc.resume_blob, tc.resume_filename or f"{tc.name}.pdf",
                source_vendor_id=tc.vendor_id,
                source_vendor_name=vendor.name if vendor else "",
                source_tracked_candidate_id=tc.id,
            )
            if candidate:
                # Prefer the vendor-submitted contact details if the resume
                # parse didn't find them
                candidate.name = tc.name or candidate.name
                candidate.email = candidate.email or tc.email or ""
                candidate.phone = candidate.phone or tc.phone or ""
                db.add(candidate)
                candidates.append(candidate)
                sb = candidate.strengths_breakdown or {}
                if sb.get("aiPowered"):
                    await enrich_skill_taxonomy(db, {
                        "technical": sb.get("technicalSkills", []),
                        "business": sb.get("businessSkills", []),
                        "soft": sb.get("softSkills", []),
                    })
        except Exception as e:
            print(f"Error processing vendor candidate {tc.id}: {e}")
            continue

    # ── Score candidates sourced from a Requisition's submitted Applications ──
    for application, cand in source_applications:
        try:
            if not cand.resume_blob:
                continue
            candidate = await _score_and_build_candidate(
                cand.resume_blob, cand.resume_filename or f"{cand.full_name}.pdf",
                source_application_id=application.id,
            )
            if candidate:
                # Prefer the Candidate Master's own contact details if the
                # resume parse didn't find them.
                candidate.name = cand.full_name or candidate.name
                candidate.email = candidate.email or cand.email or ""
                candidate.phone = candidate.phone or cand.phone or ""
                db.add(candidate)
                candidates.append(candidate)
                sb = candidate.strengths_breakdown or {}
                if sb.get("aiPowered"):
                    await enrich_skill_taxonomy(db, {
                        "technical": sb.get("technicalSkills", []),
                        "business": sb.get("businessSkills", []),
                        "soft": sb.get("softSkills", []),
                    })
        except Exception as e:
            print(f"Error processing requisition candidate (application {application.id}): {e}")
            continue

    await db.commit()
    for c in candidates:
        await db.refresh(c)

    # Auto-log each candidate's Resume Screening completion into Interview
    # Scheduling — see _log_joblens_interview's docstring. Best-effort per
    # candidate (never raises), so a logging hiccup can't fail the whole
    # scoring run.
    for c in candidates:
        await _log_joblens_interview(db, current_user, c, round_name="Resume Screening", interview_type="Resume Screening",
                                      status="Completed", notes="Auto-logged: resume screened and scored.")

    candidates.sort(key=lambda c: c.ats_score, reverse=True)

    # Whether a Groq key EXISTS says nothing about whether extraction
    # actually SUCCEEDED via it — a bad model/expired key/unreachable
    # Ollama fallback all silently degrade to weak keyword-only matching
    # while a key is still technically configured. Report real per-candidate
    # outcomes instead, so a silent degradation is never invisible again.
    ai_powered_flags = [bool((c.strengths_breakdown or {}).get("aiPowered")) for c in candidates]

    # Report the outcome back to the pool — if at least one candidate in
    # this session succeeded via AI, the key is healthy; if none did while
    # relying on the shared pool, that's a signal to route around it for a
    # while (see utils/groq_pool.py for the cooldown/auto-recovery logic).
    if groq_key and key_resolution["source"] == "pool":
        from utils.groq_pool import record_key_outcome
        await record_key_outcome(db, key_resolution["pool_id"], success=any(ai_powered_flags) if ai_powered_flags else False)

    _phone_schedule = await _phone_schedule_map(db, [c.id for c in candidates])
    return {
        "session_id": session.id,
        "jd_skills":  jd_skills[:30],
        "ai_powered": all(ai_powered_flags) if ai_powered_flags else (groq_key is not None),
        "ai_powered_partial": 0 < sum(ai_powered_flags) < len(ai_powered_flags),
        "total":      len(candidates),
        "candidates": [_fmt(c, _phone_schedule.get(c.id)) for c in candidates],
    }


@router.get("/sessions")
async def list_sessions(
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    r = await db.execute(
        select(JobLensSession)
        .where(JobLensSession.user_id == current_user.id)
        .order_by(JobLensSession.created_at.desc())
    )
    return [
        {
            "id": s.id, "sequence_number": s.sequence_number or s.id, "cv_count": s.cv_count,
            "low_threshold": s.low_threshold, "high_threshold": s.high_threshold,
            "status": s.status,
            "created_at": s.created_at.isoformat() if s.created_at else None,
            "jd_preview": (s.jd_text or "")[:120] + "...",
            "ai_powered": False,
        }
        for s in r.scalars().all()
    ]


@router.get("/sessions/{session_id}")
async def get_session(
    session_id: int,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    sr = await db.execute(
        select(JobLensSession).where(
            JobLensSession.id == session_id,
            JobLensSession.user_id == current_user.id,
        )
    )
    session = sr.scalar_one_or_none()
    if not session:
        raise HTTPException(404, "Session not found")

    cr = await db.execute(
        select(JobLensCandidate)
        .where(JobLensCandidate.session_id == session_id)
        .order_by(JobLensCandidate.ats_score.desc())
    )
    candidates = cr.scalars().all()

    _phone_schedule = await _phone_schedule_map(db, [c.id for c in candidates])
    return {
        "id": session.id,
        "sequence_number": session.sequence_number or session.id,
        "jd_text": session.jd_text,
        "jd_skills": session.jd_skills,
        "jd_role": session.jd_role or "",
        "jd_location": session.jd_location or "",
        "jd_company": session.jd_company or "",
        "jd_record_id": session.jd_record_id,
        "jd_client_name": session.jd_client_name or session.jd_company or "",
        "jd_essential_skills": session.jd_essential_skills or [],
        "jd_good_to_have_skills": session.jd_good_to_have_skills or [],
        "jd_optional_skills": session.jd_optional_skills or [],
        "low_threshold": session.low_threshold,
        "high_threshold": session.high_threshold,
        "status": session.status,
        "cv_count": session.cv_count,
        "created_at": session.created_at.isoformat() if session.created_at else None,
        "candidates": [_fmt(c, _phone_schedule.get(c.id)) for c in candidates],
        # ── Dual-track scoring config used for this session — lets the UI
        # pre-fill weight sliders / logistics fields to "what was actually
        # used" when reopening a session ──────────────────────────────────
        "salary_budget_min": session.salary_budget_min or 0,
        "salary_budget_max": session.salary_budget_max or 0,
        "max_notice_days": session.max_notice_days or 0,
        "jd_remote_allowed": bool(session.jd_remote_allowed),
        "weights": session.weights or {},
        "disqualifiers": session.disqualifiers or {},
    }


@router.post("/sessions/{session_id}/reweight")
async def reweight_session(
    session_id: int,
    payload: dict,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Re-composite and re-rank an ENTIRE session's candidates with new
    weights/disqualifiers/logistics constraints — purely in Python, no LLM
    calls, no re-parsing resumes. Each candidate's technical_score and raw
    logistics facts (expected_salary, notice_period_days, current_location)
    are already stored from the original /run — only the JD-side
    constraints and the weighting formula can change here, which is
    exactly what a recruiter dragging a weight slider needs.

    Body (all optional — omitted fields keep the session's current value):
    {
      "weights": {...overrides...},
      "disqualifiers": {...overrides...},
      "salary_budget_min": int, "salary_budget_max": int,
      "max_notice_days": int, "remote_allowed": bool
    }
    """
    from utils.scoring import merge_weights, merge_disqualifiers, compute_non_technical_score, check_hard_disqualifiers, compute_composite_score

    sr = await db.execute(
        select(JobLensSession).where(
            JobLensSession.id == session_id,
            JobLensSession.user_id == current_user.id,
        )
    )
    session = sr.scalar_one_or_none()
    if not session:
        raise HTTPException(404, "Session not found")

    weights = merge_weights({**(session.weights or {}), **(payload.get("weights") or {})})
    disqualifiers = merge_disqualifiers({**(session.disqualifiers or {}), **(payload.get("disqualifiers") or {})})
    salary_min = payload.get("salary_budget_min", session.salary_budget_min or 0)
    salary_max = payload.get("salary_budget_max", session.salary_budget_max or 0)
    max_notice = payload.get("max_notice_days", session.max_notice_days or 0)
    remote = payload.get("remote_allowed", session.jd_remote_allowed or False)

    cr = await db.execute(select(JobLensCandidate).where(JobLensCandidate.session_id == session_id))
    candidates = cr.scalars().all()

    for c in candidates:
        stored_logistics = c.logistics or {}
        logistics = {
            "expected_salary": stored_logistics.get("expected_salary") or 0,
            "notice_period_days": stored_logistics.get("notice_period_days", -1),
            "current_location": stored_logistics.get("current_location") or "",
            "salary_budget_min": salary_min,
            "salary_budget_max": salary_max,
            "max_notice_days": max_notice,
            "jd_location": stored_logistics.get("jd_location") or session.jd_location or "",
            "remote_allowed": remote,
        }
        non_technical = compute_non_technical_score(logistics, weights)
        is_disqualified, disqualify_reason = check_hard_disqualifiers(logistics, disqualifiers)
        composite = compute_composite_score(c.technical_score or 0, non_technical, weights)

        status = "Not Qualified"
        if composite >= session.high_threshold:
            status = "Qualified"
        elif composite >= session.low_threshold:
            status = "Review"
        if is_disqualified:
            status = "Not Qualified"

        c.ats_score = composite
        c.non_technical_score = non_technical["score"]
        c.status = status
        c.hard_disqualified = is_disqualified
        c.disqualify_reason = disqualify_reason
        c.logistics = {
            **logistics,
            "salary_score": non_technical["salary_score"],
            "notice_score": non_technical["notice_score"],
            "location_score": non_technical["location_score"],
            "applicable": non_technical["applicable"],
        }
        db.add(c)

    session.weights = weights
    session.disqualifiers = disqualifiers
    session.salary_budget_min = salary_min
    session.salary_budget_max = salary_max
    session.max_notice_days = max_notice
    session.jd_remote_allowed = remote
    db.add(session)

    await db.commit()
    for c in candidates:
        await db.refresh(c)

    candidates = sorted(candidates, key=lambda c: c.ats_score, reverse=True)
    _phone_schedule = await _phone_schedule_map(db, [c.id for c in candidates])
    return {
        "session_id": session.id,
        "weights": weights,
        "disqualifiers": disqualifiers,
        "candidates": [_fmt(c, _phone_schedule.get(c.id)) for c in candidates],
    }


@router.get("/sessions/{session_id}/video-decision-settings")
async def get_video_decision_settings(
    session_id: int, current_user: User = Depends(get_current_user), db: AsyncSession = Depends(get_db),
):
    from utils.scoring import merge_video_decision_weights, merge_video_decision_thresholds
    sr = await db.execute(select(JobLensSession).where(JobLensSession.id == session_id, JobLensSession.user_id == current_user.id))
    session = sr.scalar_one_or_none()
    if not session:
        raise HTTPException(404, "Session not found")
    return {
        "weights": merge_video_decision_weights(session.video_decision_weights),
        "thresholds": merge_video_decision_thresholds(session.video_decision_thresholds),
    }


@router.post("/sessions/{session_id}/video-decision-settings")
async def set_video_decision_settings(
    session_id: int, payload: dict,
    current_user: User = Depends(get_current_user), db: AsyncSession = Depends(get_db),
):
    """Same 'reweight and re-apply instantly' idea as /reweight above, for
    Video Interview's auto-decision instead of Resume Screening's ATS
    score. Unlike the per-candidate auto-decision in _run_video_analysis
    (which never overwrites a recommendation that's already set), THIS
    endpoint recomputes and overwrites the recommendation for every
    already-analyzed candidate in the session — an explicit "apply these
    new settings now" action, not a background side-effect, so
    overwriting an existing value here is the whole point rather than
    something to protect against.

    Body: {"weights": {...overrides...}, "thresholds": {...overrides...}}
    """
    from utils.scoring import (
        merge_video_decision_weights, merge_video_decision_thresholds,
        compute_video_composite_score, compute_video_decision,
    )
    sr = await db.execute(select(JobLensSession).where(JobLensSession.id == session_id, JobLensSession.user_id == current_user.id))
    session = sr.scalar_one_or_none()
    if not session:
        raise HTTPException(404, "Session not found")

    weights = merge_video_decision_weights({**(session.video_decision_weights or {}), **(payload.get("weights") or {})})
    thresholds = merge_video_decision_thresholds({**(session.video_decision_thresholds or {}), **(payload.get("thresholds") or {})})

    cr = await db.execute(select(JobLensCandidate).where(JobLensCandidate.session_id == session_id))
    candidates = cr.scalars().all()
    updated = 0
    for c in candidates:
        if not c.video_analysis or "error" in (c.video_analysis or {}):
            continue
        composite = compute_video_composite_score(c.video_analysis, weights)
        if composite is None:
            continue
        c.video_screening_recommendation = compute_video_decision(composite, thresholds)
        analysis = dict(c.video_analysis)
        analysis["auto_decision_score"] = composite
        analysis["auto_decision"] = c.video_screening_recommendation
        c.video_analysis = analysis
        updated += 1

    session.video_decision_weights = weights
    session.video_decision_thresholds = thresholds
    await db.commit()
    return {"weights": weights, "thresholds": thresholds, "updated": updated}


@router.post("/sessions/{session_id}/candidates/{candidate_id}/questions")
async def get_questions(
    session_id: int,
    candidate_id: int,
    regenerate: bool = False,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    sr = await db.execute(
        select(JobLensSession).where(
            JobLensSession.id == session_id,
            JobLensSession.user_id == current_user.id,
        )
    )
    session = sr.scalar_one_or_none()
    if not session:
        raise HTTPException(404, "Session not found")

    cr = await db.execute(
        select(JobLensCandidate).where(JobLensCandidate.id == candidate_id)
    )
    candidate = cr.scalar_one_or_none()
    if not candidate:
        raise HTTPException(404, "Candidate not found")

    # Return existing questions or regenerate. regenerate=True skips this
    # cache deliberately — candidates scored before this endpoint started
    # using resume content (see generate_questions's resume_context
    # param) have OLD, JD-only questions cached here; a recruiter
    # clicking "Generate" on one of them expects a fresh, resume-aware
    # set, not the stale cached ones silently returned again.
    if candidate.interview_questions and not regenerate:
        return {"questions": candidate.interview_questions, "ai_powered": False}

    # BUG FIX: this used to call get_credential(db, user_id, "groq",
    # "api_key") directly, which only ever checks a user's own key or a
    # single legacy is_global=True UserAPIKey row — it has no idea the
    # Groq Key Pool (GroqKeyPool table, utils/groq_pool.py) exists at
    # all. The main resume-scoring pipeline already resolves Groq keys
    # through resolve_groq_key() (which checks personal -> pool ->
    # legacy-global, in that order), so any admin who set Groq up via
    # the Key Pool — the primary supported way, per Settings' own
    # Admin Console UI — had scoring work fine while THIS endpoint kept
    # reporting "No Groq API key configured" and silently regenerating
    # the exact same default questions every time.
    from utils.groq_pool import resolve_groq_key
    key_resolution = await resolve_groq_key(db, current_user.id)
    groq_key = key_resolution["groq_key"]
    groq_model = key_resolution["model"] or await get_groq_model(db, current_user.id)

    if groq_key:
        # Built from data already extracted from THIS candidate's own
        # resume during scoring (resume_summary's categorized bullets +
        # the free-text summary + their strengths breakdown) — not the
        # original file re-read from scratch, and not just their
        # matched-skills list, which is JD-derived rather than
        # resume-derived and identical for every candidate who happens
        # to share those skills.
        resume_parts = []
        if candidate.resume_summary:
            for section, bullets in candidate.resume_summary.items():
                if bullets:
                    resume_parts.append(f"{section.replace('_', ' ').title()}: " + "; ".join(bullets))
        if candidate.summary:
            resume_parts.append(f"Summary: {candidate.summary}")
        resume_context = "\n".join(resume_parts)

        questions, gen_error = await generate_questions(
            session.jd_text or "", candidate.name,
            candidate.matched_skills or [], groq_key, groq_model,
            resume_context=resume_context,
        )
    else:
        questions = _default_questions(candidate.name, candidate.matched_skills or [])
        gen_error = "No Groq API key configured (Settings -> API Keys) — showing default questions instead."

    candidate.interview_questions = questions
    await db.commit()
    return {"questions": questions, "ai_powered": groq_key is not None and gen_error is None, "error": gen_error}


@router.put("/candidates/{candidate_id}/shortlist")
async def toggle_shortlist(
    candidate_id: int,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    cr = await db.execute(
        select(JobLensCandidate).where(JobLensCandidate.id == candidate_id)
    )
    c = cr.scalar_one_or_none()
    if not c:
        raise HTTPException(404, "Candidate not found")
    c.shortlisted = not c.shortlisted
    await db.commit()
    return {"shortlisted": c.shortlisted}


# ── CANDIDATE CONTACT / VIDEO INTERVIEW INVITE ───────────────────────────────

@router.post("/candidates/{candidate_id}/prepare-invite")
async def prepare_invite(
    candidate_id: int,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Ensures the candidate has a unique interview token and returns it so
    the frontend can build a shareable, no-login video-interview link."""
    cr = await db.execute(
        select(JobLensCandidate).where(JobLensCandidate.id == candidate_id)
    )
    c = cr.scalar_one_or_none()
    if not c:
        raise HTTPException(404, "Candidate not found")

    if not c.interview_token:
        c.interview_token = secrets.token_urlsafe(24)
        await db.commit()
        await db.refresh(c)

    return {
        "token": c.interview_token,
        "candidate_name": c.name,
        "candidate_email": c.email,
    }


@router.post("/candidates/{candidate_id}/send-invite")
async def send_invite(
    candidate_id: int,
    payload: SendInviteRequest,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    cr = await db.execute(
        select(JobLensCandidate).where(JobLensCandidate.id == candidate_id)
    )
    c = cr.scalar_one_or_none()
    if not c:
        raise HTTPException(404, "Candidate not found")

    smtp_cfg = await _get_smtp_config(current_user.id, db)
    _send_email(smtp_cfg, payload.to_email, payload.subject, payload.body_html)

    from capabilities.acquisition.service import get_or_create_default_organisation
    from capabilities.communication import service as comm_service
    org = await get_or_create_default_organisation(db, current_user)
    await comm_service.log_manual_send(
        db, org.id, "video_invite", payload.subject, payload.body_html,
        sent_by_user_id=current_user.id, joblens_candidate_id=c.id,
    )
    await db.commit()

    return {"sent": True}


class RejectionEmailRequest(BaseModel):
    candidate_ids: List[int]
    subject: str
    # {name} is the only placeholder — replaced with each candidate's own
    # first name right before that candidate's individual email is sent.
    body_html_template: str


@router.post("/candidates/reject-email")
async def send_rejection_emails(
    payload: RejectionEmailRequest,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Screening Decision's bulk 'Send Rejection Email' action — sends
    ONE separate, individually-addressed email per candidate (a loop of
    single-recipient send_email() calls, never one email with everyone
    in To/CC), so no candidate can ever see another candidate's name or
    email address. Each send is independent: one candidate's SMTP/email
    failure doesn't stop the rest, and the response reports exactly who
    succeeded and who didn't so the recruiter can retry just the
    failures instead of guessing whether "Send" actually worked for
    everyone. Every successful send stamps rejection_email_sent_at,
    which the Screening Decision table shows directly (see
    JobLensPage.tsx's "candidateContact"-style column)."""
    if not payload.candidate_ids:
        raise HTTPException(400, "No candidates selected.")

    smtp_cfg = await _get_smtp_config(current_user.id, db)

    from capabilities.acquisition.service import get_or_create_default_organisation
    from capabilities.communication import service as comm_service
    org = await get_or_create_default_organisation(db, current_user)

    rows = (await db.execute(
        select(JobLensCandidate).where(JobLensCandidate.id.in_(payload.candidate_ids))
    )).scalars().all()
    by_id = {c.id: c for c in rows}

    sent, failed = [], []
    for cid in payload.candidate_ids:
        c = by_id.get(cid)
        if not c:
            failed.append({"candidate_id": cid, "name": None, "error": "Candidate not found."})
            continue
        if not c.email:
            failed.append({"candidate_id": cid, "name": c.name, "error": "No email on file."})
            continue
        first_name = (c.name or "").strip().split(" ")[0] or "there"
        body_html = payload.body_html_template.replace("{name}", first_name)
        try:
            _send_email(smtp_cfg, c.email, payload.subject, body_html)
            c.rejection_email_sent_at = datetime.utcnow()
            sent.append({"candidate_id": cid, "name": c.name, "email": c.email})
            await comm_service.log_manual_send(
                db, org.id, "screening_rejection", payload.subject, body_html,
                sent_by_user_id=current_user.id, joblens_candidate_id=c.id,
            )
        except HTTPException as e:
            # SMTP misconfiguration (bad credentials, host unreachable)
            # affects every remaining send too — stop the loop early
            # instead of repeating the same failure for every candidate.
            failed.append({"candidate_id": cid, "name": c.name, "error": e.detail})
            await comm_service.log_manual_send(
                db, org.id, "screening_rejection", payload.subject, body_html,
                sent_by_user_id=current_user.id, joblens_candidate_id=c.id,
                status="Failed", failure_reason=str(e.detail)[:500],
            )
            if "not configured" in str(e.detail).lower() or "rejected these smtp credentials" in str(e.detail).lower():
                break
        except Exception as e:
            failed.append({"candidate_id": cid, "name": c.name, "error": str(e)[:200]})
            await comm_service.log_manual_send(
                db, org.id, "screening_rejection", payload.subject, body_html,
                sent_by_user_id=current_user.id, joblens_candidate_id=c.id,
                status="Failed", failure_reason=str(e)[:500],
            )

    await db.commit()
    return {"sent": sent, "failed": failed}


async def _get_or_create_joblens_interview(
    db: AsyncSession, current_user: User, candidate: JobLensCandidate,
    round_name: str, interview_type: str,
):
    """Finds this candidate's existing Interview Scheduling row for a
    given round (Phone Interview / Video Interview), or creates one.
    Deliberately ONE row per (joblens_candidate_id, interview_type) —
    not a fresh row per action — so that "screening done", "phone link
    sent", "phone status", "video link sent", "video status", etc. all
    accumulate as state on a single trackable record instead of a
    scattered activity log. Returns the Interview row (uncommitted
    changes are the caller's to commit).
    """
    from capabilities.acquisition.service import get_or_create_default_organisation
    from capabilities.interview.models import Interview
    from capabilities.interview import service as interview_service

    org = await get_or_create_default_organisation(db, current_user)
    existing = (await db.execute(
        select(Interview).where(
            Interview.joblens_candidate_id == candidate.id,
            Interview.interview_type == interview_type,
        )
    )).scalar_one_or_none()
    if existing:
        return existing

    interview = Interview(
        organisation_id=org.id, owner_user_id=current_user.id,
        sequence_number=await interview_service.get_next_sequence(db, org.id),
        candidate_id=None, joblens_candidate_id=candidate.id,
        round_name=round_name, interview_type=interview_type,
        status="Requested",
    )
    db.add(interview)
    return interview


async def _log_joblens_interview(
    db: AsyncSession, current_user: User, candidate: JobLensCandidate,
    round_name: str, interview_type: str, notes: str = "", status: str = "",
    completed_at: Optional[datetime] = None,
) -> None:
    """Registers/updates a row in Interview Scheduling for a JobLens/
    CandidateLens action — Video Interview's 'Send Interview Invite' and
    Phone Interview's 'Candidate reached by phone' both call this, as do
    Resume Screening / Phone Interview / Video Interview's actual
    COMPLETION points (status="Completed") so a candidate finishing any
    of the three stages shows up in Interview Scheduling with a real
    completion date, not just "an invite/call happened at some point".
    JobLens candidates live in tiq_joblens_candidates, not the Talent
    Pool's tiq_candidates that Interview.candidate_id has always pointed
    at, so this sets joblens_candidate_id instead (see that column's
    docstring in capabilities/interview/models.py) and skips
    create_interview's full endpoint — that path validates a Talent Pool
    candidate/requisition, builds an interviewers list, and can fire a
    self-schedule/approval flow, none of which applies here.

    Upserts (see _get_or_create_joblens_interview) rather than always
    inserting, so repeated actions against the same candidate/round
    update one tracking row instead of piling up duplicates.

    status, when given, is written directly (e.g. "Completed") — and for
    "Completed" specifically, completed_at is set to now WITHOUT
    touching scheduled_at, so a real booked time (e.g. from a Calendly
    booking) already on the row survives a later completion instead of
    being overwritten by the completion timestamp.

    Best-effort: never raises. A failure here shouldn't block the actual
    invite/contact/completion action it's just a side-effect record of —
    but it IS logged (not silently swallowed with zero trace), since a
    failure here previously left no evidence anywhere that an Interview
    Scheduling row should have existed but doesn't.
    """
    try:
        interview = await _get_or_create_joblens_interview(db, current_user, candidate, round_name, interview_type)
        if notes:
            interview.notes = notes
        if status:
            interview.status = status
            if status == "Completed":
                interview.completed_at = completed_at or datetime.utcnow()
        interview.updated_at = datetime.utcnow()
        await db.commit()
    except Exception as e:
        await db.rollback()
        logger.warning(f"_log_joblens_interview failed for candidate_id={candidate.id}, round={interview_type}: {e}")


@router.post("/candidates/{candidate_id}/mark-contacted")
async def mark_contacted(
    candidate_id: int,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Flips the 'contacted' flag once the recruiter sends the interview
    invite letter (via mailto handoff to their own mail client)."""
    cr = await db.execute(
        select(JobLensCandidate).where(JobLensCandidate.id == candidate_id)
    )
    c = cr.scalar_one_or_none()
    if not c:
        raise HTTPException(404, "Candidate not found")
    c.contacted = True
    await db.commit()
    interview = await _get_or_create_joblens_interview(db, current_user, c, round_name="Video Interview", interview_type="Video Interview")
    interview.video_invite_sent_at = datetime.utcnow()
    interview.notes = "Auto-logged: interview invite sent via Video Interview's Candidate Contact."
    interview.updated_at = datetime.utcnow()
    await db.commit()
    return {"contacted": True}


PHONE_RECOMMENDATIONS = ["Proceed", "Hold", "Reject"]


@router.post("/candidates/{candidate_id}/phone-contacted")
async def mark_phone_contacted(
    candidate_id: int,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Marks a candidate as reached for their phone screening call —
    separate from the video-interview `contacted` flag above, since a
    candidate can be phone-contacted long before (or without) ever being
    invited to a video round."""
    cr = await db.execute(select(JobLensCandidate).where(JobLensCandidate.id == candidate_id))
    c = cr.scalar_one_or_none()
    if not c:
        raise HTTPException(404, "Candidate not found")
    if c.phone_screening_status == "Not Started":
        c.phone_screening_status = "Contacted"
        c.phone_screening_at = datetime.utcnow()
        await db.commit()
        await _log_joblens_interview(db, current_user, c, round_name="Phone Screening", interview_type="Phone Interview",
                                      notes="Auto-logged: marked reached by phone in Phone Interview.")
    else:
        await db.commit()
    return {"phone_screening_status": c.phone_screening_status}


class SendCalendlyLinkRequest(BaseModel):
    to_email: str = ""   # defaults to the candidate's own email if blank
    subject: str = "Schedule your phone screening interview"
    body_html: str = ""  # defaults to a standard message wrapping the link if blank
    # Set by the frontend's compose-modal flow (prepare-calendly-link
    # fetched this once when the modal opened, and the editable message
    # body already has it embedded) — reusing it here avoids minting a
    # SECOND single-use Calendly link that never actually gets emailed.
    # Left blank, this endpoint generates its own exactly as before.
    booking_url: str = ""


@router.post("/candidates/{candidate_id}/phone-interview/prepare-calendly-link")
async def prepare_phone_calendly_link(
    candidate_id: int,
    current_user: User = Depends(get_current_user), db: AsyncSession = Depends(get_db),
):
    """Resolves (and, for the single-use-link Calendly setup, actually
    mints) the booking link WITHOUT sending anything — lets the frontend
    show a real, working link in the compose modal's editable message
    body before the recruiter hits Send, the same way Video Interview's
    invite modal shows a real interview link (from /prepare-invite)
    before it's emailed."""
    cr = await db.execute(select(JobLensCandidate).where(JobLensCandidate.id == candidate_id))
    c = cr.scalar_one_or_none()
    if not c:
        raise HTTPException(404, "Candidate not found")

    from capabilities.interview import service as interview_service

    creds = await interview_service.get_calendly_credentials(db, current_user.id)
    if creds["booking_url"]:
        booking_url = creds["booking_url"]
    elif creds["api_key"] and creds["event_type_uri"]:
        booking_url = await interview_service.create_calendly_single_use_link(creds["api_key"], creds["event_type_uri"])
    else:
        raise HTTPException(400, "Set up Calendly under Settings -> API Keys first — either paste your Calendly booking link, or a Personal Access Token + Event Type.")

    return {"booking_url": booking_url, "candidate_name": c.name, "candidate_email": c.email}


@router.post("/candidates/{candidate_id}/phone-interview/send-calendly-link")
async def send_phone_calendly_link(
    candidate_id: int, payload: SendCalendlyLinkRequest,
    current_user: User = Depends(get_current_user), db: AsyncSession = Depends(get_db),
):
    """Phone Interview page action: generates the recruiter's Calendly
    booking link (same credentials as Interview Scheduling's Calendly
    integration — Settings -> API Keys -> Calendly) and emails it to the
    candidate so they can book their own initial HR screening call.

    Registers/updates the same Interview Scheduling row this candidate's
    other Phone Interview actions use (see _get_or_create_joblens_interview)
    so "phone interview schedule link sent" and the resulting booking are
    both visible there, not just on this page.
    """
    cr = await db.execute(select(JobLensCandidate).where(JobLensCandidate.id == candidate_id))
    c = cr.scalar_one_or_none()
    if not c:
        raise HTTPException(404, "Candidate not found")

    to_email = (payload.to_email or c.email or "").strip()
    if not to_email:
        raise HTTPException(400, "This candidate has no email on file — nowhere to send the link.")

    from capabilities.interview import service as interview_service

    if payload.booking_url.strip():
        booking_url = payload.booking_url.strip()
    else:
        creds = await interview_service.get_calendly_credentials(db, current_user.id)
        if creds["booking_url"]:
            booking_url = creds["booking_url"]
        elif creds["api_key"] and creds["event_type_uri"]:
            booking_url = await interview_service.create_calendly_single_use_link(creds["api_key"], creds["event_type_uri"])
        else:
            raise HTTPException(400, "Set up Calendly under Settings -> API Keys first — either paste your Calendly booking link, or a Personal Access Token + Event Type.")

    body_html = payload.body_html.strip() or (
        f"<p>Hi {c.name or 'there'},</p>"
        f"<p>Thanks for your interest — we'd like to set up a quick initial phone screening interview with you.</p>"
        f"<p>Please use the link below to pick a time that works for you:</p>"
        f"<p><a href=\"{booking_url}\">{booking_url}</a></p>"
        f"<p>Looking forward to speaking with you.</p>"
    )
    subject = payload.subject.strip() or "Schedule your phone screening interview"
    smtp_cfg = await _get_smtp_config(current_user.id, db)
    _send_email(smtp_cfg, to_email, subject, body_html)

    interview = await _get_or_create_joblens_interview(db, current_user, c, round_name="Phone Screening", interview_type="Phone Interview")
    interview.calendly_scheduling_url = booking_url
    interview.calendly_link_sent_at = datetime.utcnow()
    if interview.status not in ("Scheduled", "Completed"):
        interview.status = "Requested"
    interview.updated_at = datetime.utcnow()

    from capabilities.acquisition.service import get_or_create_default_organisation
    from capabilities.communication import service as comm_service
    org = await get_or_create_default_organisation(db, current_user)
    await comm_service.log_manual_send(
        db, org.id, "phone_calendly_link", subject, body_html,
        sent_by_user_id=current_user.id, joblens_candidate_id=c.id,
    )
    await db.commit()

    return {"sent": True, "calendly_scheduling_url": booking_url, "interview_id": interview.id}


# ── TELEPHONY (click-to-call + SMS scheduling — see utils/telephony.py) ──
# Phone Interview page's counterpart to Interview Scheduling's /call and
# /sms-schedule endpoints (capabilities/interview/router.py) — same
# underlying Twilio plumbing, but operating on a JobLensCandidate
# directly rather than requiring an Interview row to already exist, same
# bridge pattern _get_or_create_joblens_interview/_log_joblens_interview
# already use for "Send Calendly Link" above.

@router.post("/candidates/{candidate_id}/phone-interview/call")
async def call_phone_candidate(
    candidate_id: int, current_user: User = Depends(get_current_user), db: AsyncSession = Depends(get_db),
):
    """Click-to-call from the Phone Interview page: bridges the
    recruiter's own configured caller number (Settings -> API Keys ->
    Telephony) to this candidate's phone. Registers/updates the same
    Interview Scheduling row this candidate's other Phone Interview
    actions use, so the call shows up there too."""
    from utils.telephony import get_telephony_config, place_click_to_call

    cr = await db.execute(select(JobLensCandidate).where(JobLensCandidate.id == candidate_id))
    c = cr.scalar_one_or_none()
    if not c:
        raise HTTPException(404, "Candidate not found")

    config = await get_telephony_config(db, current_user.id)
    result = await place_click_to_call(config, c.phone or "", record=True)

    interview = await _get_or_create_joblens_interview(db, current_user, c, round_name="Phone Screening", interview_type="Phone Interview")
    interview.phone_call_sid = result["sid"]
    interview.phone_call_status = result["status"]
    interview.phone_called_at = datetime.utcnow()
    # Fresh call — any transcript from a PREVIOUS call to this candidate
    # no longer corresponds to what's about to happen, so clear it rather
    # than leave a stale transcript sitting under a new call_sid.
    interview.phone_recording_sid = None
    interview.phone_transcript = None
    interview.phone_transcript_status = None
    interview.updated_at = datetime.utcnow()
    await db.commit()

    return {
        "call_sid": result["sid"], "status": result["status"],
        "caller_number": result["from"], "candidate_number": result["to"],
    }


@router.post("/candidates/{candidate_id}/phone-interview/fetch-transcript")
async def fetch_phone_transcript(
    candidate_id: int, current_user: User = Depends(get_current_user), db: AsyncSession = Depends(get_db),
):
    """Pulls the Twilio call recording for this candidate's most recent
    Phone Interview call (see call_phone_candidate's record=True) and
    transcribes it via Groq Whisper — same transcription engine Video
    Interview already uses (_transcribe_video), reused here rather than
    a second copy of the same Groq audio-transcription call.

    On-demand rather than automatic: recordings typically take a few
    seconds to a minute to become available after the call ends, and
    this app deliberately avoids requiring a public webhook URL just to
    know when that's happened (see utils/telephony.place_click_to_call's
    docstring) — the recruiter clicks "Fetch Call Transcript" once
    they're ready to check.

    Note: this only works for calls placed via Twilio click-to-call
    (Settings -> API Keys -> Telephony). A call placed via the Windows/
    Android Caller dials out on the recruiter's own phone's native SIM —
    a real cellular call TalentIQ has no access to the audio of at all,
    so there is nothing here to record or transcribe for that path.
    """
    from capabilities.interview.models import Interview
    from utils.telephony import get_telephony_config, fetch_call_recordings, download_recording_audio

    cr = await db.execute(select(JobLensCandidate).where(JobLensCandidate.id == candidate_id))
    c = cr.scalar_one_or_none()
    if not c:
        raise HTTPException(404, "Candidate not found")

    interview = (await db.execute(
        select(Interview).where(
            Interview.joblens_candidate_id == candidate_id,
            Interview.interview_type == "Phone Interview",
        )
    )).scalar_one_or_none()
    if not interview or not interview.phone_call_sid:
        raise HTTPException(400, "No call has been placed to this candidate yet — use Call Candidate first.")

    config = await get_telephony_config(db, current_user.id)
    recordings = await fetch_call_recordings(config, interview.phone_call_sid)
    if not recordings:
        raise HTTPException(404, "No recording available yet for that call — Twilio usually takes a few seconds to a minute after the call ends. Try again shortly.")

    # Most recent recording for this call (a call can technically produce
    # more than one, e.g. if re-answered) — sorted explicitly rather than
    # trusting Twilio's response ordering.
    recording = sorted(recordings, key=lambda r: r.get("date_created", ""), reverse=True)[0]
    recording_sid = recording.get("sid")

    interview.phone_transcript_status = "Processing"
    await db.commit()

    from utils.groq_pool import resolve_groq_key, record_key_outcome
    key_resolution = await resolve_groq_key(db, current_user.id)
    groq_key = key_resolution["groq_key"]
    if not groq_key:
        interview.phone_transcript_status = "Failed"
        await db.commit()
        raise HTTPException(400, "No Groq API key configured (own or admin-shared) — required for transcription.")

    try:
        audio_bytes = await download_recording_audio(config, recording_sid)
        transcript = _transcribe_video(audio_bytes, "audio/mpeg", groq_key)
        if not transcript:
            interview.phone_transcript_status = "Failed"
            await db.commit()
            raise HTTPException(400, "Transcription returned no speech content — the recording may be silent or too short.")

        interview.phone_recording_sid = recording_sid
        interview.phone_transcript = transcript
        interview.phone_transcript_status = "Completed"
        await db.commit()
        if key_resolution["pool_id"] is not None:
            await record_key_outcome(db, key_resolution["pool_id"], success=True)
        return {"transcript": transcript, "status": "Completed"}
    except HTTPException:
        raise
    except Exception as e:
        interview.phone_transcript_status = "Failed"
        await db.commit()
        raise HTTPException(500, f"Failed to transcribe the call recording: {str(e)[:200]}")


class SendScheduleSmsRequest(BaseModel):
    scheduled_at: str    # ISO datetime — when the candidate will be called
    message: str = ""    # defaults to a standard "you'll be called at <time>" text if blank


@router.post("/candidates/{candidate_id}/phone-interview/send-sms-schedule")
async def send_phone_schedule_sms(
    candidate_id: int, payload: SendScheduleSmsRequest,
    current_user: User = Depends(get_current_user), db: AsyncSession = Depends(get_db),
):
    """Texts the candidate the time they'll be called, and sets that
    time on this candidate's Interview Scheduling row (scheduled_at +
    status="Scheduled") so it's reflected in both the calendar and the
    table there — the interviewer-picks-the-time alternative to
    'Send Calendly Link' above, for when the candidate would rather get
    a text than self-schedule."""
    from utils.telephony import get_telephony_config, send_sms

    cr = await db.execute(select(JobLensCandidate).where(JobLensCandidate.id == candidate_id))
    c = cr.scalar_one_or_none()
    if not c:
        raise HTTPException(404, "Candidate not found")

    try:
        scheduled_dt = datetime.fromisoformat(payload.scheduled_at.replace("Z", "+00:00"))
    except ValueError:
        raise HTTPException(400, "scheduled_at must be a valid ISO datetime string.")

    when_display = scheduled_dt.strftime("%A, %B %d at %I:%M %p").replace(" 0", " ")
    body = payload.message.strip() or (
        f"Hi {c.name or 'there'}, this is a heads-up that we'll be calling you for your "
        f"phone interview on {when_display}. Talk soon!"
    )

    config = await get_telephony_config(db, current_user.id)
    await send_sms(config, c.phone or "", body)

    interview = await _get_or_create_joblens_interview(db, current_user, c, round_name="Phone Screening", interview_type="Phone Interview")
    interview.scheduled_at = scheduled_dt
    interview.status = "Scheduled"
    interview.call_sms_sent_at = datetime.utcnow()
    interview.updated_at = datetime.utcnow()
    await db.commit()

    return {"sent": True, "scheduled_at": interview.scheduled_at.isoformat(), "status": interview.status}


CANDIDATE_STATUSES = ["Qualified", "Review", "Not Qualified"]


class StatusUpdateRequest(BaseModel):
    status: str


@router.put("/candidates/{candidate_id}/status")
async def update_candidate_status(
    candidate_id: int,
    payload: StatusUpdateRequest,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Manual override of the AI-computed status — Resume Screening's
    Status column is a dropdown, not just a read-only badge, since a
    recruiter reading the actual resume can reasonably disagree with the
    score (e.g. a borderline Review case they've decided is worth
    treating as Qualified). Overwrites status directly rather than
    touching ats_score/threshold — the score and its breakdown stay as
    the AI computed them; only the recruiter's final call changes."""
    if payload.status not in CANDIDATE_STATUSES:
        raise HTTPException(400, f"status must be one of: {', '.join(CANDIDATE_STATUSES)}")
    cr = await db.execute(select(JobLensCandidate).where(JobLensCandidate.id == candidate_id))
    c = cr.scalar_one_or_none()
    if not c:
        raise HTTPException(404, "Candidate not found")
    c.status = payload.status
    await db.commit()
    return {"status": c.status}


class PhoneResultRequest(BaseModel):
    recommendation: str
    notes: str = ""


@router.post("/candidates/{candidate_id}/phone-result")
async def save_phone_result(
    candidate_id: int, payload: PhoneResultRequest,
    current_user: User = Depends(get_current_user), db: AsyncSession = Depends(get_db),
):
    """Records the outcome of a completed phone screening call — the
    Phone Interview stage's equivalent of the video stage's
    interview-result endpoint above, but simpler: no webcam/emotion data,
    just a recruiter's own recommendation and notes after the call."""
    if payload.recommendation not in PHONE_RECOMMENDATIONS:
        raise HTTPException(400, f"recommendation must be one of: {', '.join(PHONE_RECOMMENDATIONS)}")
    cr = await db.execute(select(JobLensCandidate).where(JobLensCandidate.id == candidate_id))
    c = cr.scalar_one_or_none()
    if not c:
        raise HTTPException(404, "Candidate not found")
    c.phone_screening_status = "Completed"
    c.phone_screening_recommendation = payload.recommendation
    c.phone_screening_notes = payload.notes.strip()
    c.phone_screening_at = datetime.utcnow()
    await db.commit()
    await _log_joblens_interview(db, current_user, c, round_name="Phone Screening", interview_type="Phone Interview",
                                  status="Completed", notes="Auto-logged: phone screening completed.")
    return {
        "phone_screening_status": c.phone_screening_status,
        "phone_screening_recommendation": c.phone_screening_recommendation,
        "phone_screening_notes": c.phone_screening_notes,
    }


class VideoResultRequest(BaseModel):
    recommendation: str
    notes: str = ""


@router.post("/candidates/{candidate_id}/video-result")
async def save_video_result(
    candidate_id: int, payload: VideoResultRequest,
    current_user: User = Depends(get_current_user), db: AsyncSession = Depends(get_db),
):
    """Records the recruiter's Decision & Comments for the Video Interview
    stage — the video-stage equivalent of save_phone_result above, kept
    separate from the AI video_analysis (which scores the interview
    itself, not the recruiter's own call on it)."""
    if payload.recommendation not in PHONE_RECOMMENDATIONS:
        raise HTTPException(400, f"recommendation must be one of: {', '.join(PHONE_RECOMMENDATIONS)}")
    cr = await db.execute(select(JobLensCandidate).where(JobLensCandidate.id == candidate_id))
    c = cr.scalar_one_or_none()
    if not c:
        raise HTTPException(404, "Candidate not found")
    c.video_screening_recommendation = payload.recommendation
    c.video_screening_notes = payload.notes.strip()
    c.video_screening_at = datetime.utcnow()
    await db.commit()
    return {
        "video_screening_recommendation": c.video_screening_recommendation,
        "video_screening_notes": c.video_screening_notes,
    }


@router.get("/morphcast-key")
async def get_morphcast_key(
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Returns the recruiter's saved MorphCast license key (Settings > API
    Keys, service: morphcast, key: license_key), used client-side by the
    video interview modal for facial emotion analysis. Not a sensitive
    secret in the traditional sense — MorphCast's SDK runs entirely in the
    browser and the key is visible in that SDK's own network calls anyway."""
    key = await get_credential(db, current_user.id, "morphcast", "license_key")
    return {"license_key": key or ""}


# ══════════════════════════════════════════════════════════════════════════
# INTERVIEW SETTINGS — answer time per question + TTS voice, configured
# platform-wide by an admin (Settings > Admin Console), read by both the
# recruiter-side interview modal and the candidate's public interview page.
# ══════════════════════════════════════════════════════════════════════════

DEFAULT_ANSWER_SECONDS = 30


async def _resolve_interview_settings(db: AsyncSession, user_id: int) -> dict:
    from utils.tts import DEFAULT_EDGE_VOICE
    creds = await get_all_credentials(db, user_id, "interview")
    try:
        answer_seconds = int(creds.get("answer_seconds") or DEFAULT_ANSWER_SECONDS)
    except (TypeError, ValueError):
        answer_seconds = DEFAULT_ANSWER_SECONDS
    answer_seconds = max(10, min(answer_seconds, 600))  # sane guardrails either side
    tts_engine = creds.get("tts_engine") or "edge"  # "edge" (default/first choice) | "browser"
    tts_voice = creds.get("tts_voice") or DEFAULT_EDGE_VOICE
    return {"answer_seconds": answer_seconds, "tts_voice": tts_voice, "tts_engine": tts_engine}


@router.get("/interview-settings")
async def get_interview_settings(
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Answer time (seconds) + TTS voice for CandidateLens interviews.
    Admin-configured platform-wide (Settings > Admin Console); every
    recruiter/candidate reads the same values via get_all_credentials'
    global fallback (service='interview' is in SHAREABLE_SERVICES)."""
    settings = await _resolve_interview_settings(db, current_user.id)
    from utils.tts import EDGE_VOICES
    settings["edge_voices"] = EDGE_VOICES
    return settings


@router.post("/tts")
async def synthesize_interview_question(
    payload: dict,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Speaks interview question text using the admin-configured engine
    (Microsoft Edge neural voices) instead of the browser's built-in,
    robotic SpeechSynthesis voice. Returns 503 if that engine
    isn't available right now — the frontend falls back to the browser
    voice on any non-200 response, so an interview never gets stuck."""
    text = (payload or {}).get("text", "").strip()
    if not text:
        raise HTTPException(400, "text is required")
    settings = await _resolve_interview_settings(db, current_user.id)
    if settings["tts_engine"] == "browser":
        raise HTTPException(503, "Server-side TTS is turned off — using the browser voice.")
    from utils.tts import synthesize
    audio, media_type = await synthesize(text, engine=settings["tts_engine"], voice=settings["tts_voice"])
    if audio is None:
        raise HTTPException(503, f"{settings['tts_engine']} TTS is not available right now — falling back to the browser voice.")
    return Response(content=audio, media_type=media_type)


class InterviewResult(BaseModel):
    happy: int = 0
    neutral: int = 0
    sad: int = 0
    angry: int = 0
    fear: int = 0
    disgust: int = 0
    surprise: int = 0
    dominant: str = "Neutral"


@router.post("/candidates/{candidate_id}/interview-result")
async def save_interview_result(
    candidate_id: int,
    result: InterviewResult,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    cr = await db.execute(
        select(JobLensCandidate).where(JobLensCandidate.id == candidate_id)
    )
    c = cr.scalar_one_or_none()
    if not c:
        raise HTTPException(404, "Candidate not found")
    c.emotion_happy    = result.happy
    c.emotion_neutral  = result.neutral
    c.emotion_sad      = result.sad
    c.emotion_angry    = result.angry
    c.emotion_fear     = result.fear
    c.emotion_disgust  = result.disgust
    c.emotion_surprise = result.surprise
    c.dominant_emotion = result.dominant
    c.video_status     = "Completed"
    await db.commit()
    await _log_joblens_interview(db, current_user, c, round_name="Video Interview", interview_type="Video Interview",
                                  status="Completed", notes="Auto-logged: video interview completed.")
    return {"status": "saved"}


# ══════════════════════════════════════════════════════════════════════════════
# AUTOMATIC VIDEO ANALYSIS — runs once the video blob is stored.
# Transcribes via Groq's Whisper endpoint (accepts .webm directly, no ffmpeg
# needed), then scores the transcript against the interview questions with
# an LLM. Result is written back onto the SAME candidate row. Runs as a
# FastAPI background task so the upload request returns immediately rather
# than blocking on transcription + LLM latency.
# ══════════════════════════════════════════════════════════════════════════════

def _transcribe_video(video_bytes: bytes, mimetype: str, groq_key: str) -> str:
    """Groq's /audio/translations endpoint (Whisper's "translate" task,
    not "transcribe") — same API shape as /audio/transcriptions, but
    ALWAYS outputs English text regardless of what language was actually
    spoken, instead of transcribing faithfully in the spoken language
    (which is what /audio/transcriptions did before this fix — a
    candidate answering in Urdu produced an Urdu transcript, not an
    English one). Recruiters reviewing transcripts need them in one
    consistent language regardless of which language a candidate
    answered in. Accepts webm/mp4/mp3/wav/m4a/ogg directly — no local
    audio extraction/conversion needed."""
    resp = requests.post(
        "https://api.groq.com/openai/v1/audio/translations",
        headers={"Authorization": f"Bearer {groq_key}"},
        files={"file": ("interview.webm", video_bytes, mimetype or "video/webm")},
        data={"model": "whisper-large-v3", "response_format": "text"},
        timeout=180,
        proxies={"http": None, "https": None},  # avoid any system proxy intercepting this call
    )
    resp.raise_for_status()
    return resp.text.strip()


async def _analyze_transcript(
    transcript: str, questions: list, candidate_name: str, groq_key: str, groq_model: str = DEFAULT_GROQ_MODEL
) -> dict:
    from langchain_groq import ChatGroq
    from langchain.schema import HumanMessage

    llm = ChatGroq(api_key=groq_key, model=groq_model, temperature=0.2, max_tokens=4000, reasoning_format="hidden", reasoning_effort="low", max_retries=0)
    questions_block = "\n".join(f"{i+1}. {q}" for i, q in enumerate(questions)) or "(not recorded)"
    prompt = f"""You are an experienced hiring manager reviewing a recorded video
interview transcript for {candidate_name}. Be fair and evidence-based — only
comment on what the transcript actually supports.

QUESTIONS ASKED:
{questions_block}

TRANSCRIPT (auto-generated, translated to English if the candidate answered
in another language — may contain minor recognition errors):
\"\"\"{transcript[:6000]}\"\"\"

Assess the candidate's spoken interview performance. Also split the
transcript into the individual question/answer pairs it actually contains,
in the order asked — the transcript is one continuous recording with no
built-in boundaries between questions, so use the QUESTIONS ASKED list
above plus the natural content/topic shifts in the transcript to work out
where each answer starts and ends. If a question doesn't appear to have
been answered (skipped, cut off, inaudible), still include it with
answer_transcript set to an empty string rather than omitting it — the
reader needs to see that gap, not have it silently disappear.

Return ONLY valid JSON, no markdown, no commentary:
{{
  "communication_score": <0-100, clarity/structure of spoken answers>,
  "relevance_score": <0-100, how directly answers addressed the questions asked>,
  "confidence_score": <0-100, based on language used — decisiveness, specificity, hedging>,
  "overall_score": <0-100, holistic>,
  "strengths": ["<specific, evidence-based>", "..."],
  "concerns": ["<specific, evidence-based>", "..."],
  "key_observations": ["<notable moment or answer>", "..."],
  "summary": "<3-4 sentence overall assessment>",
  "qa_pairs": [
    {{"question": "<question text, verbatim from QUESTIONS ASKED>", "answer_transcript": "<the portion of the transcript that answers it, verbatim or lightly trimmed — empty string if not answered>"}}
  ]
}}"""
    resp = llm.invoke([HumanMessage(content=prompt)])
    from utils.llm_extraction import _parse_json_response
    data = _parse_json_response(resp.content)
    if data is None:
        raise ValueError(f"LLM returned unparseable/empty response (length {len(resp.content)})")
    return data


# Strong references for fire-and-forget asyncio.create_task() calls (see
# analyze_unanalyzed_videos below) — asyncio only holds a WEAK reference
# to a task otherwise, so without this a task can be garbage-collected
# mid-flight with no warning, silently dropping a queued analysis.
_BACKGROUND_TASK_REFS: set = set()


async def _run_video_analysis(candidate_id: int):
    """Background task — opens its OWN DB session since the request-scoped
    one is already closed by the time this runs after the response returns."""
    async with AsyncSessionLocal() as db:
        key_resolution = None
        c = None
        try:
            cr = await db.execute(
                select(JobLensCandidate, JobLensSession)
                .join(JobLensSession, JobLensCandidate.session_id == JobLensSession.id)
                .where(JobLensCandidate.id == candidate_id)
            )
            row = cr.first()
            if not row:
                return
            c, session = row

            c.video_analysis_status = "Processing"
            await db.commit()

            # Same bug as get_questions() had before this session's fix:
            # get_credential() alone doesn't know the Groq Key Pool
            # exists — resolve_groq_key checks personal -> pool -> legacy
            # global, matching what the scoring pipeline already does.
            from utils.groq_pool import resolve_groq_key, record_key_outcome
            key_resolution = await resolve_groq_key(db, session.user_id)
            groq_key = key_resolution["groq_key"]
            groq_model = key_resolution["model"] or await get_groq_model(db, session.user_id)
            if not groq_key:
                c.video_analysis_status = "Failed"
                c.video_analysis = {"error": "No Groq API key configured (own or admin-shared) — required for transcription and analysis."}
                await db.commit()
                return

            if not c.video_key and not c.video_blob:
                c.video_analysis_status = "Failed"
                c.video_analysis = {"error": "No video stored for this candidate."}
                await db.commit()
                return

            # S3-backed video (video_key) takes priority; video_blob is
            # only read for legacy rows or if S3 isn't configured.
            video_bytes = None
            if c.video_key:
                video_bytes = await get_video_bytes(db, c.video_key)
            if video_bytes is None:
                video_bytes = c.video_blob
            if not video_bytes:
                c.video_analysis_status = "Failed"
                c.video_analysis = {"error": "Video could not be retrieved from storage."}
                await db.commit()
                return

            transcript = _transcribe_video(video_bytes, c.video_mimetype, groq_key)
            if not transcript:
                c.video_analysis_status = "Failed"
                c.video_analysis = {"error": "Transcription returned no speech content."}
                await db.commit()
                return

            # Save the transcript as soon as it exists, BEFORE attempting
            # the LLM scoring step below — previously both were only ever
            # written together after a successful analysis, so a
            # transcription that worked fine followed by a scoring
            # failure (LLM error, rate limit, bad JSON) silently threw
            # the transcript away too, leaving nothing to show even
            # though the hard part (transcribing the actual audio)
            # had already succeeded.
            c.video_transcript = transcript
            await db.commit()

            analysis = await _analyze_transcript(
                transcript, c.interview_questions or [], c.name or "the candidate", groq_key, groq_model
            )

            c.video_analysis = analysis
            c.video_analysis_status = "Completed"

            # Auto-decision: only ever sets the recommendation, never
            # overwrites one — if a recruiter already made a manual call
            # (e.g. between an earlier failed analysis and this retry),
            # that choice sticks. Explicit re-weighting via
            # /sessions/{id}/video-decision-settings is the only path
            # that recomputes decisions that already exist.
            if not c.video_screening_recommendation:
                from utils.scoring import (
                    merge_video_decision_weights, merge_video_decision_thresholds,
                    compute_video_composite_score, compute_video_decision,
                )
                vd_weights = merge_video_decision_weights(session.video_decision_weights)
                vd_thresholds = merge_video_decision_thresholds(session.video_decision_thresholds)
                composite = compute_video_composite_score(analysis, vd_weights)
                if composite is not None:
                    c.video_screening_recommendation = compute_video_decision(composite, vd_thresholds)
                    analysis["auto_decision_score"] = composite
                    analysis["auto_decision"] = c.video_screening_recommendation
                    c.video_analysis = analysis  # re-assign so JSON column picks up the added keys

            await db.commit()
            if key_resolution["pool_id"] is not None:
                await record_key_outcome(db, key_resolution["pool_id"], success=True)
        except Exception as e:
            if key_resolution is not None and key_resolution.get("pool_id") is not None:
                try:
                    from utils.groq_pool import record_key_outcome
                    await record_key_outcome(db, key_resolution["pool_id"], success=False)
                except Exception:
                    pass
            try:
                if c is not None:
                    c.video_analysis_status = "Failed"
                    c.video_analysis = {"error": str(e)[:300]}
                    await db.commit()
            except Exception:
                pass


@router.post("/candidates/{candidate_id}/video")
async def upload_interview_video(
    candidate_id: int,
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...),
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Stores the recorded interview video as a blob on the candidate row,
    alongside their resume — then kicks off automatic transcription +
    performance analysis in the background, written back onto this same row."""
    cr = await db.execute(
        select(JobLensCandidate).where(JobLensCandidate.id == candidate_id)
    )
    c = cr.scalar_one_or_none()
    if not c:
        raise HTTPException(404, "Candidate not found")
    content = await file.read()

    # Prefer S3 (transcoded + uploaded); fall back to the old in-Postgres
    # blob if S3 isn't configured yet, so upload never just breaks.
    uploaded = await upload_video_and_get_key(db, current_user.id, candidate_id, content, file.content_type)
    if uploaded:
        c.video_key = uploaded["key"]
        c.video_mimetype = uploaded["mimetype"]
        c.video_size_bytes = uploaded["size_bytes"]
        c.video_blob = None
    else:
        c.video_blob = content
        c.video_mimetype = file.content_type or "video/webm"
        c.video_size_bytes = len(content)
    c.video_analysis_status = "Pending"
    await db.commit()
    background_tasks.add_task(_run_video_analysis, candidate_id)
    return {"status": "saved", "size_bytes": c.video_size_bytes, "storage": "s3" if uploaded else "database"}


@router.post("/candidates/{candidate_id}/reanalyze-video")
async def reanalyze_video(
    candidate_id: int,
    background_tasks: BackgroundTasks,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Re-runs transcription + performance analysis on the ALREADY-STORED
    video for this candidate — no re-recording or re-upload needed. Useful
    after adding a Groq key, or to re-check with updated interview
    questions/context."""
    cr = await db.execute(
        select(JobLensCandidate).where(JobLensCandidate.id == candidate_id)
    )
    c = cr.scalar_one_or_none()
    if not c:
        raise HTTPException(404, "Candidate not found")
    if not c.video_key and not c.video_blob:
        raise HTTPException(400, "No video stored for this candidate yet.")
    c.video_analysis_status = "Pending"
    await db.commit()
    background_tasks.add_task(_run_video_analysis, candidate_id)
    return {"status": "queued"}


@router.post("/candidates/analyze-unanalyzed-videos")
async def analyze_unanalyzed_videos(
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Backfill action: finds every one of THIS recruiter's candidates who
    has a stored interview video (video_key or the legacy video_blob) but
    no completed analysis on file — status is NULL/"Pending"/"Failed", or
    a status of "Completed" with no actual video_analysis payload (the
    has_video bug fixed earlier this session could leave rows in exactly
    that inconsistent state) — and queues _run_video_analysis for each.
    Explicitly excludes candidates already "Processing" so a second click
    doesn't pile on duplicate concurrent runs of the same video.

    Scoped to the calling recruiter's own sessions only, same ownership
    boundary as every other candidate-level action in this router — this
    is not an admin-wide "reprocess everything" tool."""
    rows = (await db.execute(
        select(JobLensCandidate.id)
        .join(JobLensSession, JobLensCandidate.session_id == JobLensSession.id)
        .where(
            JobLensSession.user_id == current_user.id,
            or_(JobLensCandidate.video_key.isnot(None), JobLensCandidate.video_blob.isnot(None)),
            or_(
                JobLensCandidate.video_analysis_status.is_(None),
                JobLensCandidate.video_analysis_status.in_(["Pending", "Failed"]),
                JobLensCandidate.video_analysis.is_(None),
            ),
        )
    )).scalars().all()

    # Fire all of these CONCURRENTLY via asyncio.create_task rather than
    # background_tasks.add_task — FastAPI's BackgroundTasks runs its
    # tasks one at a time, in sequence, awaiting each fully before
    # starting the next. With N candidates queued that meant N videos
    # got transcribed+analysed strictly one after another regardless of
    # how many Groq keys were sitting in the pool ready to be used in
    # parallel — the whole point of having multiple keys is exactly this
    # kind of concurrent throughput, and BackgroundTasks was silently
    # throwing that away. _BACKGROUND_TASK_REFS holds a strong reference
    # to each task so it can't be garbage-collected mid-flight (asyncio
    # only holds a weak reference otherwise) — cleared via the task's own
    # done-callback once it finishes, success or failure either way.
    for candidate_id in rows:
        task = asyncio.create_task(_run_video_analysis(candidate_id))
        _BACKGROUND_TASK_REFS.add(task)
        task.add_done_callback(_BACKGROUND_TASK_REFS.discard)

    if rows:
        await db.execute(
            JobLensCandidate.__table__.update()
            .where(JobLensCandidate.id.in_(rows))
            .values(video_analysis_status="Pending")
        )
        await db.commit()

    return {"queued": len(rows), "candidate_ids": rows}


@router.post("/candidates/backfill-interview-scheduling")
async def backfill_interview_scheduling(
    current_user: User = Depends(get_current_user), db: AsyncSession = Depends(get_db),
):
    """One-time catch-up action for Interview Scheduling: creates/updates
    rows for candidates who completed Resume Screening / Phone Interview /
    Video Interview BEFORE the auto-logging in the /run, save_phone_result,
    and save_interview_result endpoints existed — those completions
    happened with no Interview Scheduling row created at the time (the
    feature didn't exist yet), and there's no automatic retroactive
    mechanism for that other than running this once.

    Exact historical completion timestamps for Resume Screening and Video
    Interview were never actually stored anywhere (this app didn't track
    them until the auto-logging feature existed), so the best available
    approximation is used instead of leaving completed_at blank:
      - Resume Screening: the candidate's SESSION creation time (scoring
        happens synchronously at upload, so this is normally exact or
        within seconds of the truth).
      - Video Interview: video_screening_at (the Decision & Comments
        timestamp) if set — normally made shortly after watching the
        recording, so a reasonable proxy — otherwise this candidate's
        session creation time as a last resort.
      - Phone Interview: phone_screening_at IS the actual, exact
        completion time already (that field has existed all along), so
        no approximation needed there.

    Scoped to the calling recruiter's own sessions only, same ownership
    boundary as analyze_unanalyzed_videos above.
    """
    sessions = (await db.execute(
        select(JobLensSession).where(JobLensSession.user_id == current_user.id)
    )).scalars().all()
    if not sessions:
        return {"resume": 0, "phone": 0, "video": 0}

    session_ids = [s.id for s in sessions]
    session_created_at = {s.id: s.created_at for s in sessions}

    candidates = (await db.execute(
        select(JobLensCandidate).where(JobLensCandidate.session_id.in_(session_ids))
    )).scalars().all()

    counts = {"resume": 0, "phone": 0, "video": 0}
    for c in candidates:
        session_time = session_created_at.get(c.session_id) or datetime.utcnow()

        # Every candidate that exists in a session was, by definition,
        # already resume-screened — that's what creates the row at all.
        await _log_joblens_interview(
            db, current_user, c, round_name="Resume Screening", interview_type="Resume Screening",
            status="Completed", notes="Backfilled: resume screened and scored (pre-existing candidate).",
            completed_at=session_time,
        )
        counts["resume"] += 1

        if c.phone_screening_status == "Completed":
            await _log_joblens_interview(
                db, current_user, c, round_name="Phone Screening", interview_type="Phone Interview",
                status="Completed", notes="Backfilled: phone screening completed.",
                completed_at=c.phone_screening_at or session_time,
            )
            counts["phone"] += 1

        if c.video_status == "Completed":
            await _log_joblens_interview(
                db, current_user, c, round_name="Video Interview", interview_type="Video Interview",
                status="Completed", notes="Backfilled: video interview completed.",
                completed_at=c.video_screening_at or session_time,
            )
            counts["video"] += 1

    return counts


@router.get("/candidates/{candidate_id}/video")
async def download_interview_video(
    candidate_id: int,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    cr = await db.execute(
        select(JobLensCandidate, JobLensSession)
        .join(JobLensSession, JobLensCandidate.session_id == JobLensSession.id)
        .where(JobLensCandidate.id == candidate_id)
    )
    row = cr.first()
    if not row:
        raise HTTPException(404, "Candidate not found")
    c, session = row
    # Row-level access control: only the owning recruiter or an admin may
    # view this candidate's video — never any other user.
    if session.user_id != current_user.id and current_user.role != "admin":
        raise HTTPException(404, "Candidate not found")
    if not c.video_key and not c.video_blob:
        raise HTTPException(404, "No video recorded for this candidate")
    data = c.video_blob
    if c.video_key:
        data = await get_video_bytes(db, c.video_key) or c.video_blob
    if not data:
        raise HTTPException(404, "No video recorded for this candidate")
    return Response(content=data, media_type=c.video_mimetype or "video/webm")


@router.post("/candidates/{candidate_id}/video-view-token")
async def create_video_view_token(
    candidate_id: int,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Mints a short-lived (5 minute), video-scoped token so the browser's
    native <video> element can stream this candidate's recording directly
    from its `src` — with HTTP Range support for near-instant playback and
    seeking — without needing to attach an Authorization header (which a
    plain <video src> can't do). The one-time ownership/access check
    happens HERE, not on the streaming endpoint itself; the token is what
    proves that check already passed."""
    from jose import jwt as _jwt
    from utils.auth_utils import SECRET_KEY, ALGORITHM
    cr = await db.execute(
        select(JobLensCandidate, JobLensSession)
        .join(JobLensSession, JobLensCandidate.session_id == JobLensSession.id)
        .where(JobLensCandidate.id == candidate_id)
    )
    row = cr.first()
    if not row:
        raise HTTPException(404, "Candidate not found")
    c, session = row
    if session.user_id != current_user.id and current_user.role != "admin":
        raise HTTPException(404, "Candidate not found")
    if not c.video_key and not c.video_blob:
        raise HTTPException(404, "No video recorded for this candidate")

    # S3-backed video: hand back a presigned URL directly. Range requests
    # and seeking are then handled natively by R2/S3, not proxied through
    # our own server at all — no need for our own token/stream dance.
    if c.video_key:
        url = await get_video_presigned_url(db, c.video_key, expires_in=300)
        if url:
            return {"url": url}
        # Presign failed (e.g. S3 briefly unreachable) — fall through to
        # the legacy blob path below only if a blob also exists.
        if not c.video_blob:
            raise HTTPException(502, "Video storage is temporarily unavailable — try again shortly.")

    token = _jwt.encode(
        {"vcid": candidate_id, "exp": datetime.utcnow() + timedelta(minutes=5)},
        SECRET_KEY, algorithm=ALGORITHM,
    )
    return {"url": f"/api/joblens/candidates/{candidate_id}/video-stream?vt={token}"}


@router.get("/candidates/{candidate_id}/video-stream")
async def stream_interview_video(
    candidate_id: int,
    vt: str,
    request: Request,
    db: AsyncSession = Depends(get_db),
):
    """Range-aware video streaming for the popup player — the token from
    /video-view-token stands in for a normal Authorization header (a plain
    <video src="..."> can't send one), and IS the access check: it's
    short-lived and bound to this exact candidate_id. Supporting Range
    requests here is what makes the popup player start playing almost
    immediately instead of waiting for the whole file, and lets the
    scrubber seek without re-downloading everything before that point."""
    from jose import jwt as _jwt, JWTError
    from utils.auth_utils import SECRET_KEY, ALGORITHM
    try:
        payload = _jwt.decode(vt, SECRET_KEY, algorithms=[ALGORITHM])
    except JWTError:
        raise HTTPException(401, "This video link has expired — reopen the video to get a new one.")
    if payload.get("vcid") != candidate_id:
        raise HTTPException(403, "This link does not grant access to this video.")

    cr = await db.execute(select(JobLensCandidate).where(JobLensCandidate.id == candidate_id))
    c = cr.scalar_one_or_none()
    if not c:
        raise HTTPException(404, "No video recorded for this candidate")

    # This proxy endpoint only serves the legacy blob-in-Postgres path.
    # S3-backed videos (video_key) are served via a presigned URL handed
    # out directly by /video-view-token and never reach here in normal
    # operation — but if an old cached link does land here, redirect to
    # a fresh presigned URL instead of 404ing.
    if c.video_key and not c.video_blob:
        url = await get_video_presigned_url(db, c.video_key, expires_in=300)
        if url:
            from fastapi.responses import RedirectResponse
            return RedirectResponse(url, status_code=302)
        raise HTTPException(502, "Video storage is temporarily unavailable — try again shortly.")

    if not c.video_blob:
        raise HTTPException(404, "No video recorded for this candidate")

    data = c.video_blob
    total = len(data)
    media_type = c.video_mimetype or "video/webm"
    range_header = request.headers.get("range")

    if range_header:
        try:
            range_spec = range_header.strip().lower().replace("bytes=", "")
            start_s, _, end_s = range_spec.partition("-")
            start = int(start_s) if start_s else 0
            end = int(end_s) if end_s else total - 1
            end = min(end, total - 1)
        except ValueError:
            start, end = 0, total - 1
        if start >= total or start > end:
            raise HTTPException(416, "Requested range not satisfiable", headers={"Content-Range": f"bytes */{total}"})
        chunk = data[start:end + 1]
        return Response(
            content=chunk, status_code=206, media_type=media_type,
            headers={
                "Content-Range": f"bytes {start}-{end}/{total}",
                "Accept-Ranges": "bytes",
                "Content-Length": str(len(chunk)),
            },
        )

    return Response(
        content=data, status_code=200, media_type=media_type,
        headers={"Accept-Ranges": "bytes", "Content-Length": str(total)},
    )


@router.get("/candidates/{candidate_id}/resume-file")
async def download_resume_file(
    candidate_id: int,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    cr = await db.execute(
        select(JobLensCandidate, JobLensSession)
        .join(JobLensSession, JobLensCandidate.session_id == JobLensSession.id)
        .where(JobLensCandidate.id == candidate_id)
    )
    row = cr.first()
    if not row:
        raise HTTPException(404, "Candidate not found")
    c, session = row
    if session.user_id != current_user.id and current_user.role != "admin":
        raise HTTPException(404, "Candidate not found")
    if not c.resume_key and not c.resume_file_blob:
        raise HTTPException(404, "No resume file stored for this candidate")
    data = c.resume_file_blob
    if c.resume_key:
        data = await get_file_bytes(db, c.resume_key) or c.resume_file_blob
    if not data:
        raise HTTPException(404, "No resume file stored for this candidate")
    return Response(
        content=data,
        media_type=c.resume_file_mimetype or "application/octet-stream",
        headers={"Content-Disposition": f'inline; filename="{c.filename or "resume"}"'},
    )
# These power the emailed video-interview link so a candidate can complete
# the interview without a TalentIQ login. Access is gated by the unguessable
# token, not a session/auth check.

@router.get("/public/interview/{token}")
async def public_get_interview(token: str, db: AsyncSession = Depends(get_db)):
    cr = await db.execute(
        select(JobLensCandidate, JobLensSession)
        .join(JobLensSession, JobLensCandidate.session_id == JobLensSession.id)
        .where(JobLensCandidate.interview_token == token)
    )
    row = cr.first()
    if not row:
        raise HTTPException(404, "This interview link is invalid or has expired.")
    c, session = row
    settings = await _resolve_interview_settings(db, session.user_id)
    return {
        "candidate_name": c.name,
        "questions": c.interview_questions or [],
        "video_status": c.video_status,
        "privacy_accepted": bool(c.privacy_accepted_at),
        "answer_seconds": settings["answer_seconds"],
    }


@router.post("/public/interview/{token}/accept-privacy")
async def public_accept_interview_privacy(token: str, db: AsyncSession = Depends(get_db)):
    """The candidate ticking 'I understand and agree' on the pre-interview
    notice (recording, storage, and review by decision-makers). Camera
    access is never requested by the frontend before this succeeds."""
    cr = await db.execute(
        select(JobLensCandidate).where(JobLensCandidate.interview_token == token)
    )
    c = cr.scalar_one_or_none()
    if not c:
        raise HTTPException(404, "This interview link is invalid or has expired.")
    c.privacy_accepted_at = datetime.utcnow()
    await db.commit()
    return {"accepted": True, "accepted_at": c.privacy_accepted_at.isoformat()}


@router.post("/public/interview/{token}/tts")
async def public_synthesize_interview_question(token: str, payload: dict, db: AsyncSession = Depends(get_db)):
    """Same TTS as the authenticated /tts endpoint, resolved via the
    interview token (no login) — engine/voice settings belong to whichever
    recruiter's account generated this candidate's interview."""
    text = (payload or {}).get("text", "").strip()
    if not text:
        raise HTTPException(400, "text is required")
    cr = await db.execute(
        select(JobLensCandidate, JobLensSession)
        .join(JobLensSession, JobLensCandidate.session_id == JobLensSession.id)
        .where(JobLensCandidate.interview_token == token)
    )
    row = cr.first()
    if not row:
        raise HTTPException(404, "This interview link is invalid or has expired.")
    _, session = row
    settings = await _resolve_interview_settings(db, session.user_id)
    if settings["tts_engine"] == "browser":
        raise HTTPException(503, "Server-side TTS is turned off — using the browser voice.")
    from utils.tts import synthesize
    audio, media_type = await synthesize(text, engine=settings["tts_engine"], voice=settings["tts_voice"])
    if audio is None:
        raise HTTPException(503, f"{settings['tts_engine']} TTS is not available right now — falling back to the browser voice.")
    return Response(content=audio, media_type=media_type)


@router.get("/public/interview/{token}/morphcast-key")
async def public_get_morphcast_key(token: str, db: AsyncSession = Depends(get_db)):
    """Same MorphCast license key as the authenticated endpoint above, but
    resolved via the interview token (no login) — belongs to whichever
    recruiter's account generated this candidate's interview."""
    cr = await db.execute(
        select(JobLensCandidate, JobLensSession)
        .join(JobLensSession, JobLensCandidate.session_id == JobLensSession.id)
        .where(JobLensCandidate.interview_token == token)
    )
    row = cr.first()
    if not row:
        raise HTTPException(404, "This interview link is invalid or has expired.")
    _, session = row

    key = await get_credential(db, session.user_id, "morphcast", "license_key")
    return {"license_key": key or ""}


@router.post("/public/interview/{token}/result")
async def public_save_interview_result(
    token: str,
    result: InterviewResult,
    db: AsyncSession = Depends(get_db),
):
    cr = await db.execute(
        select(JobLensCandidate).where(JobLensCandidate.interview_token == token)
    )
    c = cr.scalar_one_or_none()
    if not c:
        raise HTTPException(404, "This interview link is invalid or has expired.")
    if not c.privacy_accepted_at:
        raise HTTPException(403, "The privacy notice must be accepted before submitting interview results.")
    c.emotion_happy    = result.happy
    c.emotion_neutral  = result.neutral
    c.emotion_sad      = result.sad
    c.emotion_angry    = result.angry
    c.emotion_fear     = result.fear
    c.emotion_disgust  = result.disgust
    c.emotion_surprise = result.surprise
    c.dominant_emotion = result.dominant
    c.video_status     = "Completed"
    await db.commit()
    # No current_user on this public/token-based path (the candidate,
    # not the recruiter, is calling it) — _log_joblens_interview needs
    # the actual owner User row (get_or_create_default_organisation
    # takes a User, not just an id), so fetch it via the candidate's own
    # session.user_id.
    sr = await db.execute(select(JobLensSession).where(JobLensSession.id == c.session_id))
    owning_session = sr.scalar_one_or_none()
    if owning_session:
        ur = await db.execute(select(User).where(User.id == owning_session.user_id))
        owner = ur.scalar_one_or_none()
        if owner:
            await _log_joblens_interview(db, owner, c, round_name="Video Interview", interview_type="Video Interview",
                                          status="Completed", notes="Auto-logged: video interview completed (candidate self-serve link).")
    return {"status": "saved"}


@router.post("/public/interview/{token}/video")
async def public_upload_interview_video(
    token: str,
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...),
    db: AsyncSession = Depends(get_db),
):
    """Same as the authenticated video-upload endpoint, but reached via the
    interview token so the candidate (no login) can submit their recorded
    video directly from the public interview page."""
    cr = await db.execute(
        select(JobLensCandidate, JobLensSession)
        .join(JobLensSession, JobLensCandidate.session_id == JobLensSession.id)
        .where(JobLensCandidate.interview_token == token)
    )
    row = cr.first()
    if not row:
        raise HTTPException(404, "This interview link is invalid or has expired.")
    c, session = row
    if not c.privacy_accepted_at:
        raise HTTPException(403, "The privacy notice must be accepted before submitting interview video.")
    content = await file.read()

    uploaded = await upload_video_and_get_key(db, session.user_id, c.id, content, file.content_type)
    if uploaded:
        c.video_key = uploaded["key"]
        c.video_mimetype = uploaded["mimetype"]
        c.video_size_bytes = uploaded["size_bytes"]
        c.video_blob = None
    else:
        c.video_blob = content
        c.video_mimetype = file.content_type or "video/webm"
        c.video_size_bytes = len(content)
    c.video_analysis_status = "Pending"
    await db.commit()
    background_tasks.add_task(_run_video_analysis, c.id)
    return {"status": "saved", "size_bytes": c.video_size_bytes, "storage": "s3" if uploaded else "database"}


@router.get("/sessions/{session_id}/export")
async def export_session(
    session_id: int,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    import pandas as pd

    sr = await db.execute(
        select(JobLensSession).where(
            JobLensSession.id == session_id,
            JobLensSession.user_id == current_user.id,
        )
    )
    session = sr.scalar_one_or_none()
    if not session:
        raise HTTPException(404, "Session not found")

    cr = await db.execute(
        select(JobLensCandidate)
        .where(JobLensCandidate.session_id == session_id)
        .order_by(JobLensCandidate.ats_score.desc())
    )
    candidates = cr.scalars().all()

    rows = []
    for i, c in enumerate(candidates, 1):
        sb = c.strengths_breakdown or {}
        score_bd = sb.get("scoreBreakdown") or {}
        va = c.video_analysis or {}
        rows.append({
            "Rank":              i,
            "Name":              c.name,
            "Email":             c.email,
            "Phone":             c.phone,
            "Resume Summary":    " | ".join(c.resume_summary or []),
            "ATS Score":         f"{c.ats_score:.1f}%",
            "Technical Score":   c.technical_score,
            "Non-Technical (Logistics) Score": c.non_technical_score,
            "Key Strength":      ", ".join(c.matched_skills or []),
            "Considerations":    ", ".join(c.missing_skills or []),
            "Status":            c.status,
            # ── Full score breakdown (tier + skill-type %s that make up the ATS score) ──
            "Score: Essential Requirements %":     (score_bd.get("essential") or {}).get("pct"),
            "Score: Good to Have %":                (score_bd.get("goodToHave") or {}).get("pct"),
            "Score: Qualification / Education %":   (score_bd.get("qualification") or {}).get("pct"),
            "Score: Technical Skills %":            (score_bd.get("technical") or {}).get("pct"),
            "Score: Tools & Platforms %":           (score_bd.get("tools") or {}).get("pct"),
            "Score: Domain Knowledge %":            (score_bd.get("domain") or {}).get("pct"),
            "Score: Soft Skills %":                 (score_bd.get("softSkills") or {}).get("pct"),
            "Score: Final ATS":                     score_bd.get("finalATS"),
            # ── Strengths breakdown (categorized, from resume parsing) ──
            "Strengths: Essential Matched":         ", ".join(sb.get("essentialMatched") or []),
            "Strengths: Technical Skills":          ", ".join(sb.get("technicalSkills") or []),
            "Strengths: Business Skills":           ", ".join(sb.get("businessSkills") or []),
            "Strengths: Soft Skills":                ", ".join(sb.get("softSkills") or []),
            "Strengths: Significant Experience":    ", ".join(sb.get("significantExperience") or []),
            "Strengths: Certifications & Degrees":  ", ".join(sb.get("certificationsDegrees") or []),
            "Strengths: Years Experience":          sb.get("yearsExperience"),
            "Strengths: Education":                 sb.get("education"),
            "Bonus Points":      c.bonus,
            "Bonus Reasons":     c.bonus_reasons,
            "Summary":           c.summary or "",
            # ── Phone Interview — Decision & Comments ──
            "Phone Screening Status":         c.phone_screening_status,
            "Phone Screening Decision":       c.phone_screening_recommendation,
            "Phone Screening Comments":       c.phone_screening_notes,
            "Phone Screening Date":           c.phone_screening_at.isoformat() if c.phone_screening_at else "",
            # ── Video Interview — Decision & Comments ──
            "Video Status":      c.video_status,
            "Video Analysis Status":          c.video_analysis_status,
            "Video Analysis: Overall Score":       va.get("overall_score"),
            "Video Analysis: Communication Score": va.get("communication_score"),
            "Video Analysis: Relevance Score":     va.get("relevance_score"),
            "Video Analysis: Confidence Score":    va.get("confidence_score"),
            "Video Analysis: Summary":             va.get("summary"),
            "Video Analysis: Strengths":           ", ".join(va.get("strengths") or []),
            "Video Analysis: Concerns":            ", ".join(va.get("concerns") or []),
            "Happy %":           c.emotion_happy or 0,
            "Neutral %":         c.emotion_neutral or 0,
            "Sad %":             c.emotion_sad or 0,
            "Angry %":           c.emotion_angry or 0,
            "Fear %":            c.emotion_fear or 0,
            "Disgust %":         c.emotion_disgust or 0,
            "Surprise %":        c.emotion_surprise or 0,
            "Dominant Emotion":  c.dominant_emotion or "Neutral",
            "Video Screening Decision":       c.video_screening_recommendation,
            "Video Screening Comments":       c.video_screening_notes,
            "Video Screening Date":           c.video_screening_at.isoformat() if c.video_screening_at else "",
            "Shortlisted":       "Yes" if c.shortlisted else "No",
        })

    buf = io.BytesIO()
    pd.DataFrame(rows).to_excel(buf, index=False)
    buf.seek(0)

    return Response(
        content=buf.read(),
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f"attachment; filename=candidatelens_{session_id}.xlsx"},
    )


@router.delete("/sessions/{session_id}")
async def delete_session(
    session_id: int,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Delete a CandidateLens session and all its candidates.

    Deleting the candidates isn't always safe on its own: a candidate can
    have been carried into Interview Scheduling (Video/Phone Interview's
    "Start Interview" or "Candidate reached by phone" create a
    tiq_interviews row via joblens_candidate_id) or into an Avatar
    Interview session, BOTH of which reference tiq_joblens_candidates.id
    with no ON DELETE behavior set on the foreign key — Postgres's default
    there is RESTRICT. Deleting a candidate that's still referenced used
    to raise an IntegrityError straight out of this endpoint as a bare
    500, which the frontend's mutation didn't surface either — from the
    UI it just looked like "Delete" silently did nothing. Detaching those
    references first (nulling the now-nullable joblens_candidate_id column
    rather than deleting the interview/avatar rows themselves) keeps that
    interview history intact while unblocking the delete, and this now
    also fails loudly with a clear message on any OTHER integrity issue
    instead of a bare 500.
    """
    from sqlalchemy import delete as sql_delete, update as sql_update
    from sqlalchemy.exc import IntegrityError
    result = await db.execute(
        select(JobLensSession).where(
            JobLensSession.id == session_id,
            JobLensSession.user_id == current_user.id,
        )
    )
    session = result.scalar_one_or_none()
    if not session:
        raise HTTPException(404, "Session not found")

    candidate_ids_r = await db.execute(
        select(JobLensCandidate.id).where(JobLensCandidate.session_id == session_id)
    )
    candidate_ids = [r[0] for r in candidate_ids_r.all()]

    if candidate_ids:
        from capabilities.interview.models import Interview
        from capabilities.avatarinterview.models import AvatarInterviewSession
        await db.execute(
            sql_update(Interview).where(Interview.joblens_candidate_id.in_(candidate_ids))
            .values(joblens_candidate_id=None)
        )
        await db.execute(
            sql_update(AvatarInterviewSession).where(AvatarInterviewSession.joblens_candidate_id.in_(candidate_ids))
            .values(joblens_candidate_id=None)
        )

    await db.execute(sql_delete(JobLensCandidate).where(JobLensCandidate.session_id == session_id))
    await db.delete(session)
    try:
        await db.commit()
    except IntegrityError as e:
        await db.rollback()
        raise HTTPException(409, f"Could not delete this session — it's still referenced elsewhere: {e}")
    return {"message": "Deleted"}


@router.delete("/sessions")
async def delete_all_sessions(
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Delete ALL sessions for the current user. Same joblens_candidate_id
    detach step as delete_session above, and for the same reason — see
    that docstring."""
    from sqlalchemy import delete as sql_delete, update as sql_update
    from sqlalchemy.exc import IntegrityError
    ids_r = await db.execute(
        select(JobLensSession.id).where(JobLensSession.user_id == current_user.id)
    )
    ids = [r[0] for r in ids_r.all()]
    if ids:
        candidate_ids_r = await db.execute(
            select(JobLensCandidate.id).where(JobLensCandidate.session_id.in_(ids))
        )
        candidate_ids = [r[0] for r in candidate_ids_r.all()]
        if candidate_ids:
            from capabilities.interview.models import Interview
            from capabilities.avatarinterview.models import AvatarInterviewSession
            await db.execute(
                sql_update(Interview).where(Interview.joblens_candidate_id.in_(candidate_ids))
                .values(joblens_candidate_id=None)
            )
            await db.execute(
                sql_update(AvatarInterviewSession).where(AvatarInterviewSession.joblens_candidate_id.in_(candidate_ids))
                .values(joblens_candidate_id=None)
            )
        await db.execute(sql_delete(JobLensCandidate).where(JobLensCandidate.session_id.in_(ids)))
        await db.execute(sql_delete(JobLensSession).where(JobLensSession.user_id == current_user.id))
        try:
            await db.commit()
        except IntegrityError as e:
            await db.rollback()
            raise HTTPException(409, f"Could not delete all sessions — some candidates are still referenced elsewhere: {e}")
    return {"message": f"Deleted {len(ids)} sessions"}