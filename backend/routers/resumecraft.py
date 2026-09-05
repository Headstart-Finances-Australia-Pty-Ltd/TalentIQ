"""
TalentIQ - ResumeCraft Router
=============================
Job-seeker module: build a tailored, industry-formatted resume + cover
letter for one specific application.

Two entry points into the same storage/download/edit machinery:
  1. /generate     - AI path. Pulls the matched/missing skills and JD
                      requirements CVIntel already computed (either from
                      an existing saved CVAnalysisRecord, or by running
                      the SAME analysis CVIntel uses if raw resume/JD
                      text is given instead), then calls Groq
                      (agents/resumecraft_agent.py) to draft both
                      documents.
  2. /manual       - resume.io-style path. No AI, no CVIntel link
                      required — the frontend's structured form posts
                      resume_data + cover_letter_text directly.

Both land in the same tiq_application_documents row, which supports full
manual editing afterwards (PUT) and .docx download for either document,
regardless of which path created it.
"""
import io
import json
from datetime import datetime
from typing import Optional

from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import Response
from pydantic import BaseModel
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select

from db.database import get_db
from models.models import User, ApplicationDocument, CVAnalysisRecord, Resume
from utils.auth_utils import get_current_user
from utils.credentials import get_groq_model
from utils.groq_pool import resolve_groq_key, record_key_outcome
from utils.sequencing import next_sequence_number
from agents.resumecraft_agent import (
    generate_tailored_resume, generate_tailored_cover_letter, EMPTY_RESUME_DATA,
)

router = APIRouter()


# ── Request/response schemas ────────────────────────────────────────────

class GenerateRequest(BaseModel):
    job_title: str = ""
    company_name: str = ""
    # Path A: reuse an existing CVIntel analysis (preferred — this is what
    # "takes gaps/strengths/requirements from CVIntel" means in practice).
    cvanalysis_record_id: Optional[int] = None
    # Path B: no saved analysis yet — paste resume/JD text fresh and this
    # endpoint runs the SAME CVIntel scoring internally before generating,
    # so the two are never out of sync even for a brand-new pair.
    resume_text: Optional[str] = None
    jd_text: Optional[str] = None
    source_resume_id: Optional[int] = None


class ManualCreateRequest(BaseModel):
    job_title: str = ""
    company_name: str = ""
    resume_data: dict
    cover_letter_text: str = ""
    resume_template: str = "modern"
    cover_letter_template: str = "standard"
    source_resume_id: Optional[int] = None


class UpdateRequest(BaseModel):
    job_title: Optional[str] = None
    company_name: Optional[str] = None
    resume_data: Optional[dict] = None
    cover_letter_text: Optional[str] = None
    resume_template: Optional[str] = None
    cover_letter_template: Optional[str] = None


def _fmt(doc: ApplicationDocument) -> dict:
    return {
        "id": doc.id,
        "sequenceNumber": doc.sequence_number or doc.id,
        "jobTitle": doc.job_title or "",
        "companyName": doc.company_name or "",
        "sourceResumeId": doc.source_resume_id,
        "cvanalysisRecordId": doc.cvanalysis_record_id,
        "resumeData": doc.resume_data or dict(EMPTY_RESUME_DATA),
        "resumeTemplate": doc.resume_template or "modern",
        "coverLetterText": doc.cover_letter_text or "",
        "coverLetterTemplate": doc.cover_letter_template or "standard",
        "aiPowered": bool(doc.ai_powered),
        "groqModel": doc.groq_model,
        "createdAt": doc.created_at.isoformat() if doc.created_at else None,
        "updatedAt": doc.updated_at.isoformat() if doc.updated_at else None,
    }


async def _get_owned_doc(db: AsyncSession, doc_id: int, user_id: int) -> ApplicationDocument:
    r = await db.execute(
        select(ApplicationDocument).where(
            ApplicationDocument.id == doc_id, ApplicationDocument.user_id == user_id,
        )
    )
    doc = r.scalar_one_or_none()
    if not doc:
        raise HTTPException(404, "Application document not found")
    return doc


# ── AI generation ────────────────────────────────────────────────────────

@router.post("/generate")
async def generate_documents(
    payload: GenerateRequest,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    # Same key resolution CVAnalysis itself uses (routers/cvintel.py's
    # _score_resume): personal key -> shared Groq key pool -> legacy
    # single global key -> none. Admin Console's "Groq Key Pool" (under
    # API Keys) is exactly this pool — a plain get_credential() lookup
    # would miss it entirely and wrongly report "no key configured" even
    # when the pool has healthy keys in it, which is why this uses the
    # same resolver CVAnalysis does rather than a simpler one.
    default_model = await get_groq_model(db, current_user.id)
    key_resolution = await resolve_groq_key(db, current_user.id)
    groq_key = key_resolution["groq_key"]
    groq_model = key_resolution["model"] or default_model

    resume_text = (payload.resume_text or "").strip()
    jd_text = (payload.jd_text or "").strip()
    cv_result: dict = {}
    cvanalysis_record_id = payload.cvanalysis_record_id
    candidate_name = current_user.name or ""

    if payload.cvanalysis_record_id:
        r = await db.execute(
            select(CVAnalysisRecord).where(
                CVAnalysisRecord.id == payload.cvanalysis_record_id,
                CVAnalysisRecord.user_id == current_user.id,
            )
        )
        record = r.scalar_one_or_none()
        if not record:
            raise HTTPException(404, "CVAnalysis record not found")
        cv_result = record.result or {}
        candidate_info = record.candidate_info or {}
        jd_info = record.jd_info or {}
        # candidate_info/jd_info are the raw text CVIntel scored against —
        # reused here so the resume/letter prompt sees exactly what CVIntel
        # saw, not a re-typed copy of it.
        resume_text = resume_text or candidate_info.get("rawText") or candidate_info.get("resumeText") or ""
        jd_text = jd_text or jd_info.get("rawText") or jd_info.get("description") or ""
        candidate_name = candidate_name or (cv_result.get("candidateProfile") or {}).get("name") or ""
    else:
        if not resume_text or not jd_text:
            raise HTTPException(
                400,
                "Provide either cvanalysis_record_id (an existing CVAnalysis record), "
                "or both resume_text and jd_text.",
            )
        # No saved analysis picked — run the SAME scoring CVIntel uses so
        # this resume/letter is generated from a real match analysis
        # rather than skipping straight past CVIntel. Also persisted as a
        # normal CVIntel history row so it shows up in CVIntel too, and
        # this ApplicationDocument can link to it like any other.
        from routers.cvintel import _score_resume
        cv_result = await _score_resume(resume_text, jd_text, groq_key, groq_model, db=db, user_id=current_user.id)
        seq_num = await next_sequence_number(db, CVAnalysisRecord, current_user.id)
        record = CVAnalysisRecord(
            user_id=current_user.id,
            sequence_number=seq_num,
            source_name=payload.job_title or "ResumeCraft",
            overall_score=cv_result.get("overallScore", 0),
            result=cv_result,
            candidate_info={"rawText": resume_text},
            jd_info={"rawText": jd_text},
            created_at=datetime.utcnow(),
        )
        db.add(record)
        await db.flush()
        cvanalysis_record_id = record.id

    resume_data = await generate_tailored_resume(
        resume_text, jd_text, payload.job_title, payload.company_name, cv_result, groq_key, groq_model,
    )
    cover_letter = await generate_tailored_cover_letter(
        resume_text, jd_text, payload.job_title, payload.company_name, candidate_name, cv_result, groq_key, groq_model,
    )

    ai_powered = bool(resume_data.get("ai_powered")) or bool(cover_letter.get("ai_powered"))
    # Feeds the pool's adaptive routing (utils/groq_pool.record_key_outcome):
    # a success clears this key's cooldown immediately; a real failure
    # while a key WAS available starts one, so the next request
    # automatically routes to a different pool key rather than retrying
    # the same struggling one. No-ops when pool_id is None (personal/
    # legacy/no key), so this is always safe to call.
    await record_key_outcome(db, key_resolution["pool_id"], success=ai_powered or key_resolution["groq_key"] is None)
    clean_resume_data = {k: v for k, v in resume_data.items() if k in EMPTY_RESUME_DATA}

    seq_num = await next_sequence_number(db, ApplicationDocument, current_user.id)
    doc = ApplicationDocument(
        user_id=current_user.id,
        sequence_number=seq_num,
        source_resume_id=payload.source_resume_id,
        cvanalysis_record_id=cvanalysis_record_id,
        job_title=payload.job_title,
        company_name=payload.company_name,
        jd_text=jd_text,
        resume_data=clean_resume_data,
        resume_template="modern",
        cover_letter_text=cover_letter.get("body", ""),
        cover_letter_template="standard",
        ai_powered=ai_powered,
        groq_model=groq_model if ai_powered else None,
        created_at=datetime.utcnow(),
    )
    db.add(doc)
    await db.commit()
    await db.refresh(doc)

    result = _fmt(doc)
    # Surface any generation-time warnings (e.g. no Groq key, bad JSON)
    # without failing the request — the document is still saved and
    # editable either way.
    result["resumeWarning"] = resume_data.get("ai_error")
    result["coverLetterWarning"] = cover_letter.get("ai_error")
    return result


# ── Manual / resume.io-style creation ───────────────────────────────────

@router.post("/manual")
async def create_manual(
    payload: ManualCreateRequest,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    clean_resume_data = dict(EMPTY_RESUME_DATA)
    clean_resume_data.update({k: v for k, v in (payload.resume_data or {}).items() if k in EMPTY_RESUME_DATA})

    seq_num = await next_sequence_number(db, ApplicationDocument, current_user.id)
    doc = ApplicationDocument(
        user_id=current_user.id,
        sequence_number=seq_num,
        source_resume_id=payload.source_resume_id,
        cvanalysis_record_id=None,
        job_title=payload.job_title,
        company_name=payload.company_name,
        jd_text="",
        resume_data=clean_resume_data,
        resume_template=payload.resume_template or "modern",
        cover_letter_text=payload.cover_letter_text or "",
        cover_letter_template=payload.cover_letter_template or "standard",
        ai_powered=False,
        created_at=datetime.utcnow(),
    )
    db.add(doc)
    await db.commit()
    await db.refresh(doc)
    return _fmt(doc)


# ── CRUD ─────────────────────────────────────────────────────────────────

@router.get("/documents")
async def list_documents(
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    r = await db.execute(
        select(ApplicationDocument)
        .where(ApplicationDocument.user_id == current_user.id)
        .order_by(ApplicationDocument.created_at.desc())
    )
    return [_fmt(d) for d in r.scalars().all()]


@router.get("/documents/{doc_id}")
async def get_document(
    doc_id: int,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    doc = await _get_owned_doc(db, doc_id, current_user.id)
    return _fmt(doc)


@router.put("/documents/{doc_id}")
async def update_document(
    doc_id: int,
    payload: UpdateRequest,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Manual-edit path — used after AI generation too, so a person can
    tweak wording, fix a date, or reorder bullets without regenerating."""
    doc = await _get_owned_doc(db, doc_id, current_user.id)

    if payload.job_title is not None:
        doc.job_title = payload.job_title
    if payload.company_name is not None:
        doc.company_name = payload.company_name
    if payload.resume_data is not None:
        merged = dict(doc.resume_data or EMPTY_RESUME_DATA)
        merged.update({k: v for k, v in payload.resume_data.items() if k in EMPTY_RESUME_DATA})
        doc.resume_data = merged
    if payload.cover_letter_text is not None:
        doc.cover_letter_text = payload.cover_letter_text
    if payload.resume_template is not None:
        doc.resume_template = payload.resume_template
    if payload.cover_letter_template is not None:
        doc.cover_letter_template = payload.cover_letter_template
    doc.updated_at = datetime.utcnow()

    await db.commit()
    await db.refresh(doc)
    return _fmt(doc)


@router.delete("/documents/{doc_id}")
async def delete_document(
    doc_id: int,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    doc = await _get_owned_doc(db, doc_id, current_user.id)
    await db.delete(doc)
    await db.commit()
    return {"message": "Deleted"}


# ── DOCX BUILDERS ─────────────────────────────────────────────────────────
# Same python-docx approach as routers/jdcreator.py's _build_docx, reused
# here so the whole platform generates Word documents the same way rather
# than each module inventing its own styling code.

def _safe_filename(name: str, fallback: str) -> str:
    cleaned = "".join(c for c in (name or "") if c.isalnum() or c in (" ", "_", "-")).strip().replace(" ", "_")
    return cleaned or fallback


def _build_resume_docx(doc: ApplicationDocument) -> bytes:
    import docx
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    data = doc.resume_data or {}
    d = docx.Document()

    normal = d.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    for section in d.sections:
        section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = Pt(36)

    # ── Header: name, headline, contact line ─────────────────────────
    name_p = d.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_p.add_run(data.get("full_name") or "Your Name")
    name_run.bold = True
    name_run.font.size = Pt(20)
    name_run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)

    if data.get("headline"):
        h_p = d.add_paragraph()
        h_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        h_run = h_p.add_run(data["headline"])
        h_run.font.size = Pt(12)
        h_run.font.color.rgb = RGBColor(0x0D, 0x94, 0x88)

    contact_parts = [v for v in [
        data.get("email"), data.get("phone"), data.get("location"),
        data.get("linkedin"), data.get("portfolio"),
    ] if v]
    if contact_parts:
        c_p = d.add_paragraph()
        c_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        c_run = c_p.add_run("  |  ".join(contact_parts))
        c_run.font.size = Pt(9.5)
        c_run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    def heading(text):
        d.add_paragraph()
        p = d.add_paragraph()
        r = p.add_run(text.upper())
        r.bold = True
        r.font.size = Pt(11.5)
        r.font.color.rgb = RGBColor(0x0D, 0x94, 0x88)
        # thin rule under the heading
        pPr = p._p.get_or_add_pPr()
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        pbdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:color"), "0D9488")
        pbdr.append(bottom)
        pPr.append(pbdr)

    if data.get("summary"):
        heading("Professional Summary")
        d.add_paragraph(data["summary"])

    if data.get("core_skills"):
        heading("Core Skills")
        d.add_paragraph(" \u2022 ".join(data["core_skills"]))

    if data.get("experience"):
        heading("Experience")
        for job in data["experience"]:
            row = d.add_paragraph()
            title_run = row.add_run(f"{job.get('job_title', '')} \u2014 {job.get('company', '')}")
            title_run.bold = True
            title_run.font.size = Pt(11)
            dates = " / ".join([v for v in [job.get("start_date"), job.get("end_date")] if v])
            meta_bits = [v for v in [job.get("location"), dates] if v]
            if meta_bits:
                meta_run = row.add_run("   " + "  |  ".join(meta_bits))
                meta_run.italic = True
                meta_run.font.size = Pt(9.5)
                meta_run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
            for bullet in job.get("bullets") or []:
                bp = d.add_paragraph(style="List Bullet")
                bp.add_run(bullet)

    if data.get("education"):
        heading("Education")
        for edu in data["education"]:
            row = d.add_paragraph()
            r = row.add_run(f"{edu.get('degree', '')} \u2014 {edu.get('institution', '')}")
            r.bold = True
            r.font.size = Pt(10.5)
            meta_bits = [v for v in [edu.get("location"), edu.get("year")] if v]
            if meta_bits:
                mr = row.add_run("   " + "  |  ".join(meta_bits))
                mr.italic = True
                mr.font.size = Pt(9.5)
            if edu.get("details"):
                d.add_paragraph(edu["details"])

    if data.get("certifications"):
        heading("Certifications")
        d.add_paragraph(" \u2022 ".join(data["certifications"]))

    if data.get("projects"):
        heading("Projects")
        for proj in data["projects"]:
            row = d.add_paragraph()
            r = row.add_run(proj.get("name", ""))
            r.bold = True
            if proj.get("description"):
                d.add_paragraph(proj["description"])

    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


def _build_cover_letter_docx(doc: ApplicationDocument, candidate_name: str) -> bytes:
    import docx
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    d = docx.Document()
    normal = d.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    for section in d.sections:
        section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = Pt(48)

    name_p = d.add_paragraph()
    name_run = name_p.add_run(candidate_name or (doc.resume_data or {}).get("full_name") or "")
    name_run.bold = True
    name_run.font.size = Pt(13)
    name_run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)

    date_p = d.add_paragraph()
    date_run = date_p.add_run(datetime.utcnow().strftime("%d %B %Y"))
    date_run.font.size = Pt(10)
    date_run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    if doc.company_name or doc.job_title:
        d.add_paragraph()
        sub = d.add_paragraph()
        sub_run = sub.add_run(f"Re: Application for {doc.job_title or 'the advertised role'}"
                               + (f" at {doc.company_name}" if doc.company_name else ""))
        sub_run.italic = True
        sub_run.font.size = Pt(10.5)

    d.add_paragraph()
    for para in (doc.cover_letter_text or "").split("\n\n"):
        if para.strip():
            d.add_paragraph(para.strip())

    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


@router.get("/documents/{doc_id}/download/resume")
async def download_resume(
    doc_id: int,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    doc = await _get_owned_doc(db, doc_id, current_user.id)
    file_bytes = _build_resume_docx(doc)
    name = _safe_filename((doc.resume_data or {}).get("full_name") or doc.job_title, f"Resume_{doc.id}")
    return Response(
        content=file_bytes,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename=Resume_{name}.docx"},
    )


@router.get("/documents/{doc_id}/download/cover-letter")
async def download_cover_letter(
    doc_id: int,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    doc = await _get_owned_doc(db, doc_id, current_user.id)
    file_bytes = _build_cover_letter_docx(doc, current_user.name or "")
    name = _safe_filename(doc.company_name or doc.job_title, f"CoverLetter_{doc.id}")
    return Response(
        content=file_bytes,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename=CoverLetter_{name}.docx"},
    )
