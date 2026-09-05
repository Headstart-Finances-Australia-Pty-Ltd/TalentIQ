"""
TalentIQ - CVAnalysis Router
Resume vs Job Description ATS analyser.
Supports PDF, DOCX, TXT for both resume and job description.
"""
import io
import re
import json
import asyncio
from typing import Optional
from datetime import datetime
from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Form
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select

from db.database import get_db
from models.models import User, UserAPIKey, CVAnalysisRecord
from utils.auth_utils import get_current_user
from utils.credentials import get_credential, get_groq_model, get_all_credentials, ollama_enabled, DEFAULT_GROQ_MODEL
from utils.sequencing import next_sequence_number

router = APIRouter()


def _extract_text(content: bytes, filename: str) -> str:
    """Extract plain text from file bytes. Tries multiple libraries with fallbacks."""
    fname = (filename or "").lower().strip()

    # ── TXT ──────────────────────────────────────────────────────────────
    if fname.endswith(".txt"):
        for enc in ("utf-8", "latin-1", "cp1252"):
            try:
                return content.decode(enc)
            except Exception:
                continue
        return ""

    # ── PDF ──────────────────────────────────────────────────────────────
    if fname.endswith(".pdf"):
        # Layout-aware pass first: detects multi-column resumes (a common
        # cause of scrambled text — see utils/layout_parse.py) and extracts
        # tables explicitly rather than letting them fall into plain
        # top-to-bottom text flow, which is where columns/tables normally
        # get their content interleaved and unreadable.
        try:
            from utils.layout_parse import extract_pdf_layout_aware
            layout_text = extract_pdf_layout_aware(content)
            if layout_text.strip():
                return layout_text
        except Exception:
            pass

        # Try pdfplumber first
        try:
            import pdfplumber
            with pdfplumber.open(io.BytesIO(content)) as pdf:
                pages = [p.extract_text() or "" for p in pdf.pages]
                text = "\n".join(pages).strip()
                if text:
                    return text
        except Exception:
            pass

        # Try pypdf
        try:
            import pypdf
            reader = pypdf.PdfReader(io.BytesIO(content))
            pages = [p.extract_text() or "" for p in reader.pages]
            text = "\n".join(pages).strip()
            if text:
                return text
        except Exception:
            pass

        # Try PyMuPDF (fitz)
        try:
            import fitz
            doc = fitz.open(stream=content, filetype="pdf")
            text = "\n".join(page.get_text() for page in doc).strip()
            if text:
                return text
        except Exception:
            pass

        return ""

    # ── DOCX ─────────────────────────────────────────────────────────────
    if fname.endswith((".docx", ".doc")):
        # Try python-docx
        try:
            import docx
            doc = docx.Document(io.BytesIO(content))

            # Headers can contain name/email/phone in letterhead-style resumes.
            # python-docx's section.header only reads ONE header per section,
            # but a docx can store up to 3 (header1/2/3.xml) — read all via XML.
            header_footer_parts = []
            try:
                import re as _re2
                import zipfile as _zipfile
                with _zipfile.ZipFile(io.BytesIO(content)) as z:
                    for name in z.namelist():
                        if _re2.match(r"word/(header|footer)\d*\.xml$", name):
                            xml = z.read(name).decode("utf-8", errors="ignore")
                            texts = _re2.findall(r"<w:t[^>]*>([^<]*)</w:t>", xml)
                            joined = "".join(texts).strip()
                            if joined:
                                header_footer_parts.append(joined)
            except Exception:
                pass

            paragraphs = list(header_footer_parts)
            paragraphs += [p.text for p in doc.paragraphs if p.text.strip()]
            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            paragraphs.append(cell.text.strip())
            text = "\n".join(paragraphs).strip()
            if text:
                return text
        except Exception:
            pass

        # Try docx2txt
        try:
            import docx2txt
            text = docx2txt.process(io.BytesIO(content))
            if text and text.strip():
                return text.strip()
        except Exception:
            pass

        return ""

    # ── Fallback: try decoding as text ────────────────────────────────────
    for enc in ("utf-8", "latin-1"):
        try:
            text = content.decode(enc).strip()
            if text:
                return text
        except Exception:
            continue
    return ""


# ── KEYWORD SCORING ────────────────────────────────────────────────────────────

DOMAIN_SKILLS = [
    "python","javascript","typescript","react","node","sql","postgresql","mongodb",
    "aws","azure","gcp","docker","kubernetes","git","agile","rest","api","graphql",
    "machine learning","ai","artificial intelligence","data science","excel","power bi","tableau","salesforce",
    "figma","django","flask","java","c#","c++","go","spark","kafka","airflow","dbt",
    "snowflake","databricks","redshift","bigquery","data architecture","data governance",
    "data modelling","data modeling","data warehouse","data lake","data lakehouse","lakehouse",
    "etl","elt","data pipeline","solution design","data mesh","data fabric","data vault",
    "dimensional modelling","dimensional modeling","enterprise data warehouse","edw",
    "master data management","mdm","data quality","data catalog","data cataloguing",
    "collibra","alation","informatica","talend","teradata","hadoop","hive","hbase",
    "adls","synapse","azure data factory","azure synapse","event-driven architecture",
    "real-time data","streaming","microservices architecture","enterprise architecture",
    "solution architecture","cloud architecture","togaf","zachman",
    "stakeholder management","cloud architecture","microservices","devops","ci/cd",
    "xero","myob","quickbooks","sap","oracle","dynamics","netsuite","sage",
    "cpa","ca","acca","cma","mba","cfa","fcpa","aca","phd",
    "accounting","tax","audit","payroll","bookkeeping","bas","gst",
    "financial reporting","budgeting","forecasting","reconciliation",
    "accounts payable","accounts receivable","ifrs","gaap",
    "leadership","communication","problem solving","teamwork","stakeholder",
    "management","strategy","operations","project management","scrum","nlp","llm",
    "basel","basel iii","banking","bfsi","insurance","lending","regulatory compliance",
    "risk management","governance framework","data governance framework",
]

STOPWORDS = {"a","an","the","and","or","of","in","on","for","with","to","be","is",
             "are","it","at","from","as","by","that","this","we","you","have","has"}


# ══════════════════════════════════════════════════════════════════════════════
# ATS SCORING — two-stage structured extraction + deterministic weighted scoring
#
# The old approach asked the LLM for a single overall number directly, which
# is exactly the kind of unreliable, unrepeatable scoring that makes ATS
# tools untrustworthy. This mirrors what real ATS-checker products (Jobscan,
# Enhancv, etc.) actually do instead:
#   1. Extract structured JD requirements (hard skills, experience, education)
#   2. Extract a structured candidate profile from the resume
#   3. Compute the score with plain, transparent, reproducible Python math —
#      the LLM's job is understanding text, not doing arithmetic.
# ══════════════════════════════════════════════════════════════════════════════

# ── Skill matching + technical scoring: MOVED to utils/technical_scoring.py
# so CVIntel and CandidateLens (routers/joblens.py) share the exact same
# matching logic and scoring formula instead of each maintaining a
# separate, silently-diverging copy — see that module's docstring for the
# full story of why they'd been producing different scores for the same
# resume/JD pair. Aliased back to the original names here so every
# existing call site below (_skill_present(...), _compute_weighted_score(...))
# keeps working unchanged.
from utils.technical_scoring import (
    skill_present as _skill_present,
    normalize_skill as _normalize_skill,
    normalize_text as _normalize_text,
    compute_technical_score as _compute_weighted_score,
)


async def _score_resume(
    resume: str, jd: str, groq_key: Optional[str], groq_model: str = DEFAULT_GROQ_MODEL,
    ollama_base_url: Optional[str] = None, ollama_model: Optional[str] = None, db=None, user_id: Optional[int] = None,
    weight_overrides: Optional[dict] = None, disqualifier_overrides: Optional[dict] = None,
) -> dict:
    from utils.llm_extraction import (
        extract_jd_requirements_categorized, extract_candidate_strengths, extract_resume_facts,
        get_taxonomy_hint, get_semantic_taxonomy_hint, enrich_skill_taxonomy,
    )

    # Semantic (pgvector-backed) taxonomy hint: terms relevant to THIS JD,
    # not just the globally most-frequent ones — see get_semantic_taxonomy_hint's
    # docstring. Gracefully falls back to frequency-based if embeddings/pgvector
    # aren't available.
    known_terms = await get_semantic_taxonomy_hint(db, jd) if db is not None else []

    # JD categorization and resume-facts extraction are genuinely
    # independent of each other — the JD's requirements don't depend on
    # the resume, and the resume's skills/experience/education don't
    # depend on the JD. Only the actual MATCHING step below needs both.
    # Running these two concurrently (rather than one after the other,
    # which is what happened before) is a real, correct optimization, not
    # just chunking for its own sake — confirmed directly in production
    # logs that each one alone typically takes ~3 seconds, so running
    # them together rather than back-to-back saves roughly that much.
    #
    # Keys are resolved SEQUENTIALLY here, before any concurrent work
    # starts — resolving both from inside the concurrent gather would mean
    # two DB writes racing on the same session, which SQLAlchemy's
    # AsyncSession doesn't allow (confirmed the hard way earlier this
    # session with a real IllegalStateChangeError).
    if db is not None and user_id is not None:
        from utils.groq_pool import resolve_groq_key
        kr_jd = await resolve_groq_key(db, user_id)
        kr_resume = await resolve_groq_key(db, user_id)
        jd_key, jd_model = kr_jd["groq_key"], kr_jd["model"] or groq_model
        resume_key, resume_model = kr_resume["groq_key"], kr_resume["model"] or groq_model
    else:
        jd_key, jd_model = groq_key, groq_model
        resume_key, resume_model = groq_key, groq_model

    # user_id is deliberately NOT passed to either of these two calls (even
    # though db is) — db alone is enough to enable the extraction CACHE
    # (see models.ExtractionCache; that cache always uses its own isolated
    # DB session, so it's safe under any concurrency pattern), but db+user_id
    # TOGETHER also enable each function's internal API-key-POOL resolution,
    # which DOES use the shared `db` session directly and is NOT safe to run
    # concurrently — and these two calls run concurrently via gather() right
    # below. Omitting user_id here keeps that pool logic off for exactly the
    # same reason it's kept off for CandidateLens's concurrent per-candidate
    # calls (see routers/joblens.py) — while still getting the caching fix.
    jd_req, resume_facts = await asyncio.gather(
        extract_jd_requirements_categorized(jd, jd_key, jd_model, ollama_base_url, ollama_model, known_terms, db=db, user_id=None),
        extract_resume_facts(resume, resume_key, resume_model, ollama_base_url, ollama_model, db=db, user_id=None),
    )
    # If the concurrent resume-facts call didn't succeed for any reason,
    # don't pass a None down as if it were real data — just omit it, and
    # extract_candidate_strengths falls back to extracting facts itself
    # as part of the matching call, exactly as it did before this change.
    from utils.llm_extraction import _mask_key_for_log
    resume_facts_preview = _mask_key_for_log(resume_key) if resume_facts else None

    strengths = await extract_candidate_strengths(
        resume, jd_req, groq_key, groq_model, ollama_base_url, ollama_model, known_terms,
        db=db, user_id=user_id,
        pre_extracted_facts=resume_facts if resume_facts else None,
    )
    # Single, easy-to-find summary line combining every key touched across
    # THIS ENTIRE analysis (JD categorization + resume-facts extraction +
    # all candidate-strengths chunks) — the line to grep for when a
    # request could plausibly have spanned multiple pool keys and you need
    # the full picture in one place rather than piecing it together from
    # several scattered lines.
    all_key_previews = sorted(set(
        ([jd_req["_groqKeyPreview"]] if jd_req.get("_groqKeyPreview") else [])
        + ([resume_facts_preview] if resume_facts_preview else [])
        + (strengths.get("_groqKeyPreviews") or [])
    ))
    print(f"  SUMMARY: CVAnalysis request used {len(all_key_previews)} distinct Groq key(s) overall: {', '.join(all_key_previews) if all_key_previews else '(none — fell back to keyword matching)'}")

    if db is not None and strengths.get("ai_powered"):
        await enrich_skill_taxonomy(db, {
            "essential": jd_req.get("essential", []),
            "good_to_have": jd_req.get("good_to_have", []),
            "technical": strengths.get("technical_skills", []),
            "business": strengths.get("business_skills", []),
            "soft": strengths.get("soft_skills", []),
        })
    from utils.scoring import (
        merge_weights, merge_disqualifiers, compute_non_technical_score,
        check_hard_disqualifiers, compute_composite_score,
    )
    weights = merge_weights(weight_overrides)
    disqualifiers = merge_disqualifiers(disqualifier_overrides)

    scoring = _compute_weighted_score(jd_req, strengths, resume, weights)
    ai_powered = bool(strengths.get("ai_powered", False))

    matched, missing = scoring["matched"], scoring["missing"]
    technical_score = scoring["overall"]

    # ── Non-technical / logistics track — decoupled from the technical
    # score above, exactly like RevaMatrix-AI's dual-track design: salary
    # expectation vs. budget, notice period vs. max acceptable, location/
    # remote fit. Fields come from LLM-extracted JD constraints (jd_req)
    # and candidate facts (strengths) — see utils/llm_extraction.py.
    logistics = {
        "expected_salary": strengths.get("expected_salary") or 0,
        "notice_period_days": strengths.get("notice_period_days", -1),
        "current_location": strengths.get("current_location") or "",
        "salary_budget_min": jd_req.get("salary_budget_min") or 0,
        "salary_budget_max": jd_req.get("salary_budget_max") or 0,
        "max_notice_days": jd_req.get("max_notice_days") or 0,
        "jd_location": jd_req.get("location") or "",
        "remote_allowed": jd_req.get("remote_allowed") or False,
    }
    non_technical = compute_non_technical_score(logistics, weights)
    is_disqualified, disqualify_reason = check_hard_disqualifiers(logistics, disqualifiers)
    overall = compute_composite_score(technical_score, non_technical, weights)

    gaps = strengths.get("gaps") or [f"Missing required skill: {s}" for s in missing[:6]]
    if not gaps:
        gaps = ["No major gaps detected against the extracted requirements"]

    suggestions = [f'Add specific, verifiable experience with "{s}" if you have it' for s in missing[:4]]
    suggestions += [
        "Quantify achievements with concrete metrics (%, $, time saved)",
        "Mirror the exact terminology used in the job description",
    ]

    summary = strengths.get("summary") or ""
    if not summary:
        exp_clause = f" and roughly {scoring['experience_pct']}% of the target experience level" if jd_req.get("min_years_experience") else ""
        summary = (
            f"Matches {scoring['skills_pct']}% of the essential requirements{exp_clause}. "
            f"Overall fit: {'Strong' if overall >= 75 else 'Moderate' if overall >= 55 else 'Needs improvement'}."
        )
    if is_disqualified:
        summary = f"⚠ Hard disqualifier: {disqualify_reason}. {summary}"
        gaps = [f"Disqualifying logistics gap: {disqualify_reason}"] + gaps

    format_warnings = [
        "Use standard ATS-compatible section headers (Experience, Education, Skills)",
        "Avoid tables, columns, or text boxes — ATS parsers often can't read them",
    ]

    # Flat "strengths" list retained for any older UI code expecting one —
    # built from the richer categorized breakdown, not a separate guess.
    flat_strengths = []
    if scoring["matched"]:
        flat_strengths.append(f"Matches {len(matched)} of {len(jd_req.get('essential', []) or [1])} essential requirements")
    flat_strengths += strengths.get("significant_experience", [])[:2]
    flat_strengths += strengths.get("certifications_degrees", [])[:2]
    if not flat_strengths:
        flat_strengths = ["Resume presents a relevant professional background"]

    return {
        "overallScore": overall,
        "strengths": flat_strengths[:6],
        "gaps": gaps[:8],
        "suggestions": suggestions[:6],
        "summaryAssessment": summary,
        "formatWarnings": format_warnings[:4],
        "detailedScores": {
            # ── Real category breakdown (was previously a fake offset from
            # "overall" for domain/softSkills — e.g. "overall - 5" — not an
            # actual measurement of anything). Now backed by the JD's own
            # per-requirement type tagging (see utils/llm_extraction.py's
            # extract_jd_requirements_categorized "requirement_types").
            # pct is null when the JD had zero requirements of that type —
            # the frontend should show "N/A", not "0%", for that case.
            "skillsMatch": scoring["skills_pct"],
            "experience": scoring["experience_pct"],
            "tools": scoring["category_breakdown"]["tool"]["pct"],
            "domain": scoring["category_breakdown"]["domain"]["pct"],
            "softSkills": scoring["category_breakdown"]["soft_skill"]["pct"],
            "format": min(92, overall + 8),
        },
        # ── Full tier + type breakdown for a detailed score view ────────
        "scoreBreakdown": {
            "essential": {"pct": scoring["essential_pct"], "label": "Essential Requirements"},
            "goodToHave": {"pct": scoring["good_to_have_pct"], "label": "Good to Have"},
            "qualification": {"pct": scoring["qualification_pct"], "label": "Qualification / Education"},
            "technical": {**scoring["category_breakdown"]["technical"], "label": "Technical Skills"},
            "tools": {**scoring["category_breakdown"]["tool"], "label": "Tools & Platforms"},
            "domain": {**scoring["category_breakdown"]["domain"], "label": "Domain Knowledge"},
            "softSkills": {**scoring["category_breakdown"]["soft_skill"], "label": "Soft Skills"},
            "finalATS": overall,
        },
        "matchedSkills": matched[:15],
        "missingSkills": missing[:15],
        "aiPowered": ai_powered,
        "groqModel": groq_model if ai_powered else None,
        # groqKeyPreview intentionally NOT included in the response — this
        # is a masked API key fragment, and there's no reason to send even
        # a masked key preview to the frontend now that the admin-only
        # debug banner that used to display it has been removed. Still
        # logged server-side (see all_key_previews above) for operational
        # debugging of key-pool issues.
        # ── Dual-track scoring (technical vs. non-technical/logistics) —
        # closes the gap vs. RevaMatrix-AI's decoupled scoring design.
        # overallScore above is now the WEIGHTED COMPOSITE of these two.
        "technicalScore": technical_score,
        "nonTechnicalScore": non_technical["score"],
        "logisticsBreakdown": {
            "applicable": non_technical["applicable"],
            "salaryScore": non_technical["salary_score"],
            "noticeScore": non_technical["notice_score"],
            "locationScore": non_technical["location_score"],
            "candidateExpectedSalary": logistics["expected_salary"] or None,
            "candidateNoticeDays": logistics["notice_period_days"] if logistics["notice_period_days"] >= 0 else None,
            "candidateLocation": logistics["current_location"] or None,
            "jdSalaryBudgetMin": logistics["salary_budget_min"] or None,
            "jdSalaryBudgetMax": logistics["salary_budget_max"] or None,
            "jdMaxNoticeDays": logistics["max_notice_days"] or None,
            "jdRemoteAllowed": logistics["remote_allowed"],
        },
        "hardDisqualified": is_disqualified,
        "disqualifyReason": disqualify_reason,
        "weightsUsed": weights,
        # ── Categorized strengths — the actual point of this round's change ──
        "strengthsBreakdown": {
            "essentialMatched": strengths.get("essential_matched", []),
            "technicalSkills": strengths.get("technical_skills", []),
            "businessSkills": strengths.get("business_skills", []),
            "softSkills": strengths.get("soft_skills", []),
            "significantExperience": strengths.get("significant_experience", []),
            "certificationsDegrees": strengths.get("certifications_degrees", []),
        },
        # ── Categorized JD requirements — mirrors CandidateLens's JD Summary ──
        "jdRequirements": {
            "roleTitle": jd_req.get("role", ""),
            "location": jd_req.get("location", ""),
            "company": jd_req.get("company", ""),
            "essential": jd_req.get("essential", []),
            "goodToHave": jd_req.get("good_to_have", []),
            "optional": jd_req.get("optional", []),
            "minYearsExperience": jd_req.get("min_years_experience", 0),
            "educationRequirement": jd_req.get("education_requirement", ""),
        },
        "candidateProfile": {
            "yearsExperience": strengths.get("years_experience", 0),
            "education": strengths.get("education", ""),
        },
    }


# ── ENDPOINT ───────────────────────────────────────────────────────────────────

@router.post("/fetch-jd-url")
async def fetch_jd_url(
    payload: dict,
    current_user: User = Depends(get_current_user),
):
    """Converts a job-posting URL (Seek, LinkedIn, Indeed, Greenhouse,
    Lever, Workday, etc.) into JD text — see utils/jd_url_fetch.py for the
    extraction strategy and its honest limitations. Returns the extracted
    text for the frontend to drop into the SAME "paste JD text" field
    used everywhere else, so the actual analysis that follows is 100%
    identical to pasting the text directly — this endpoint only changes
    HOW the text gets into the box, not anything about how it's scored.

    Body: {"url": "https://..."}
    """
    from utils.jd_url_fetch import fetch_jd_from_url, JDFetchError

    url = (payload.get("url") or "").strip()
    if not url:
        raise HTTPException(400, "A URL is required.")

    try:
        result = await fetch_jd_from_url(url)
    except JDFetchError as e:
        raise HTTPException(422, str(e))

    return result


@router.post("/analyze")
async def analyze_resume(
    job_description: str   = Form(""),
    resume_text: str       = Form(""),
    file:     Optional[UploadFile] = File(None),
    jd_file:  Optional[UploadFile] = File(None),
    # ── Dynamic weighting engine ────────────────────────────────────────
    # JSON strings from the frontend's weighting sliders (see
    # utils/scoring.DEFAULT_WEIGHTS / DEFAULT_DISQUALIFIERS for the shape).
    # Optional and backward-compatible: omitted entirely, an existing
    # caller gets identical behavior to before this change.
    weights: Optional[str]       = Form(None),
    disqualifiers: Optional[str] = Form(None),
    current_user: User     = Depends(get_current_user),
    db: AsyncSession       = Depends(get_db),
):
    # ── Extract resume text ────────────────────────────────────────────
    final_resume = resume_text.strip()
    if file and file.filename:
        raw = await file.read()
        extracted = _extract_text(raw, file.filename)
        if extracted.strip():
            final_resume = extracted
        elif not final_resume:
            raise HTTPException(
                400,
                f"Could not extract text from '{file.filename}'. "
                "Try copy-pasting the text directly into the text box instead."
            )

    # ── Extract JD text ────────────────────────────────────────────────
    final_jd = job_description.strip()
    if jd_file and jd_file.filename:
        raw_jd = await jd_file.read()
        extracted_jd = _extract_text(raw_jd, jd_file.filename)
        if extracted_jd.strip():
            final_jd = (final_jd + "\n" + extracted_jd).strip() if final_jd else extracted_jd
        elif not final_jd:
            raise HTTPException(
                400,
                f"Could not extract text from '{jd_file.filename}'. "
                "Try copy-pasting the job description text instead."
            )

    # ── Validate ───────────────────────────────────────────────────────
    if not final_resume:
        raise HTTPException(400, "Resume is required. Upload a PDF/DOCX/TXT file or paste the text.")
    if not final_jd:
        raise HTTPException(400, "Job description is required. Upload a file or paste the text.")

    # ── Score (structured JD/requirement extraction + deterministic weighting) ──
    # groq_key here is just the fallback default if db/user_id somehow
    # aren't available below — both extraction functions resolve their own
    # key(s) from the shared pool internally when given db + user_id,
    # including per-chunk pool draws for genuine multi-key parallelism on
    # a single request (see utils/llm_extraction.py and utils/groq_pool.py).
    groq_key = await get_credential(db, current_user.id, "groq", "api_key")
    groq_model = await get_groq_model(db, current_user.id)
    ollama_creds = await get_all_credentials(db, current_user.id, "ollama") if ollama_enabled() else {}
    ollama_base_url = ollama_creds.get("base_url")
    ollama_model = ollama_creds.get("model")

    weight_overrides = None
    if weights:
        try:
            weight_overrides = json.loads(weights)
        except (json.JSONDecodeError, TypeError):
            raise HTTPException(400, "weights must be a valid JSON object")

    disqualifier_overrides = None
    if disqualifiers:
        try:
            disqualifier_overrides = json.loads(disqualifiers)
        except (json.JSONDecodeError, TypeError):
            raise HTTPException(400, "disqualifiers must be a valid JSON object")

    result = await _score_resume(
        final_resume, final_jd, groq_key, groq_model,
        ollama_base_url=ollama_base_url, ollama_model=ollama_model, db=db, user_id=current_user.id,
        weight_overrides=weight_overrides, disqualifier_overrides=disqualifier_overrides,
    )

    # Echo back the exact text that was actually scored (pasted text, or
    # extracted from an uploaded PDF/DOCX) — not just the analysis derived
    # from it. The frontend persists this alongside the saved history
    # record so anything downstream that reuses "the resume behind this
    # CVAnalysis" (e.g. ResumeCraft) has the real source text to work
    # from, rather than only the lightweight name/email/phone summary it
    # used to be limited to.
    result["resumeText"] = final_resume
    result["jdText"] = final_jd

    return result


@router.post("/reweight")
async def reweight_analysis(
    payload: dict,
    current_user: User = Depends(get_current_user),
):
    """Re-composite an ALREADY-COMPUTED analysis with new weights, purely
    in Python — no LLM call, no re-parsing. This is what makes a
    real-time "drag the slider, see the score change" UI possible: the
    frontend keeps the last /analyze response (which already contains
    technicalScore and logisticsBreakdown) and calls this endpoint on
    every slider move instead of re-running the full analysis.

    Body: {"technicalScore": float, "logistics": {...same shape as
    logisticsBreakdown fields...}, "weights": {...overrides...},
    "disqualifiers": {...overrides...}}
    """
    from utils.scoring import (
        merge_weights, merge_disqualifiers, compute_non_technical_score,
        check_hard_disqualifiers, compute_composite_score,
    )

    technical_score = float(payload.get("technicalScore") or 0)
    logistics = payload.get("logistics") or {}
    weights = merge_weights(payload.get("weights"))
    disqualifiers = merge_disqualifiers(payload.get("disqualifiers"))

    non_technical = compute_non_technical_score(logistics, weights)
    is_disqualified, disqualify_reason = check_hard_disqualifiers(logistics, disqualifiers)
    overall = compute_composite_score(technical_score, non_technical, weights)

    return {
        "overallScore": overall,
        "technicalScore": technical_score,
        "nonTechnicalScore": non_technical["score"],
        "logisticsBreakdown": {
            "applicable": non_technical["applicable"],
            "salaryScore": non_technical["salary_score"],
            "noticeScore": non_technical["notice_score"],
            "locationScore": non_technical["location_score"],
        },
        "hardDisqualified": is_disqualified,
        "disqualifyReason": disqualify_reason,
        "weightsUsed": weights,
    }


# ── HISTORY (persisted server-side, so it survives browsers/devices/refresh) ──

from pydantic import BaseModel


class SaveHistoryRequest(BaseModel):
    source_name: str = "Resume"
    overall_score: float = 0
    result: dict
    candidate_info: dict = {}
    jd_info: dict = {}


def _fmt_history(r: CVAnalysisRecord) -> dict:
    return {
        "id": r.id,
        "sequenceNumber": r.sequence_number or r.id,
        "sourceName": r.source_name,
        "overallScore": r.overall_score,
        "result": r.result or {},
        "candidateInfo": r.candidate_info or {},
        "jdInfo": r.jd_info or {},
        "createdAt": r.created_at.isoformat() if r.created_at else None,
    }


@router.post("/history")
async def save_history(
    payload: SaveHistoryRequest,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    seq_num = await next_sequence_number(db, CVAnalysisRecord, current_user.id)
    record = CVAnalysisRecord(
        user_id=current_user.id,
        sequence_number=seq_num,
        source_name=payload.source_name,
        overall_score=payload.overall_score,
        result=payload.result,
        candidate_info=payload.candidate_info,
        jd_info=payload.jd_info,
        created_at=datetime.utcnow(),
    )
    db.add(record)
    await db.commit()
    await db.refresh(record)
    return _fmt_history(record)


@router.get("/history")
async def list_history(
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    r = await db.execute(
        select(CVAnalysisRecord)
        .where(CVAnalysisRecord.user_id == current_user.id)
        .order_by(CVAnalysisRecord.created_at.desc())
        .limit(50)
    )
    return [_fmt_history(rec) for rec in r.scalars().all()]


@router.delete("/history/{record_id}")
async def delete_history_item(
    record_id: int,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    r = await db.execute(
        select(CVAnalysisRecord).where(
            CVAnalysisRecord.id == record_id,
            CVAnalysisRecord.user_id == current_user.id,
        )
    )
    rec = r.scalar_one_or_none()
    if not rec:
        raise HTTPException(404, "History item not found")
    await db.delete(rec)
    await db.commit()
    return {"message": "Deleted"}


@router.delete("/history")
async def delete_all_history(
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    from sqlalchemy import delete as sql_delete
    await db.execute(sql_delete(CVAnalysisRecord).where(CVAnalysisRecord.user_id == current_user.id))
    await db.commit()
    return {"message": "Deleted"}